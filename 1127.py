# -*- coding: utf-8 -*-

# ==== stdlib ====
import os
import re
import io
import sys
import json
import uuid
import time
import requests
import sqlite3
import threading
import traceback
import contextlib
import signal
import functools
from datetime import datetime

# ==== third-party ====
import streamlit as st
import pandas as pd
import altair as alt

# ======== 基本路径设置 ========
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
DB_PATH = os.path.join(BASE_DIR, "kb.db")

PROJECT_DIR = os.path.join(BASE_DIR, "uploads")
os.makedirs(PROJECT_DIR, exist_ok=True)

SEM_INDEX_ROOT = os.path.join(BASE_DIR, "semantic_index")
os.makedirs(SEM_INDEX_ROOT, exist_ok=True)

# ---------- 领域归一化 & 索引路径 ----------
def _norm_domain_key(raw: str | None) -> str:
    s = (raw or "").strip()
    if not s:
        s = "未分类"

    for ch in r'\\/:"*?<>|':
        s = s.replace(ch, "_")
    return s

def _index_paths_common(domain_key: str, kb_type: str):
    base_dir = os.path.join(BASE_DIR, "semantic_index", domain_key, kb_type)
    os.makedirs(base_dir, exist_ok=True)
    idx_path = os.path.join(base_dir, "index.faiss")
    map_path = os.path.join(base_dir, "mapping.json")
    vec_path = os.path.join(base_dir, "vectors.npy")
    return idx_path, map_path, vec_path

# ---------- 语义索引路径:按“领域 → 类型”归类 ----------
def _index_paths(project_id: int):
    domain_raw = None
    try:
        if "cur" in globals():
            row = execute_write(
                "SELECT IFNULL(domain,'') FROM items WHERE id=?",
                (int(project_id),)
            ).fetchone()
            if row:
                domain_raw = (row[0] or "").strip()
    except Exception as e:
        log_exception("_index_paths: failed to fetch domain from database", e, level="WARNING")
        domain_raw = None

    domain_key = _norm_domain_key(domain_raw)
    kb_type = "bilingual"
    return _index_paths_common(domain_key, kb_type)

# ======== 轻量日志机制 ========
LOG_DIR = os.path.join(BASE_DIR, "logs")
os.makedirs(LOG_DIR, exist_ok=True)
LOG_FILE = os.path.join(LOG_DIR, "app.log")

def log_event(level: str, message: str, **extra):
    try:
        ts = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        record = {
            "time": ts,
            "level": (level or "INFO").upper(),
            "message": str(message),
        }
        if extra:
            record["extra"] = extra
        with open(LOG_FILE, "a", encoding="utf-8") as f:
            f.write(json.dumps(record, ensure_ascii=False) + "\n")
    except Exception as e:
        try:
            print(
                "[log_event failed] level=%s message=%s extra=%s err=%s"
                % (level, message, extra, e),
                file=sys.stderr,
            )
        except Exception:
            return

def log_exception(context: str, err: Exception, level: str = "ERROR"):
    """Centralized exception logger to avoid swallowing errors silently."""
    log_event(level, f"{context}: {err}", trace=traceback.format_exc())

# ========== P1 Helper Functions ==========

def safe_api_call(func, *args, context: str = "API调用", fallback=None, **kwargs):
    """
    P1-1: 安全的API调用包装器，消除重复的异常处理代码。
    
    Args:
        func: 要调用的函数
        *args: 位置参数
        context: 日志上下文
        fallback: 失败时的返回值
        **kwargs: 关键字参数
        
    Returns:
        func 的返回值，或 fallback 如果出错
    """
    try:
        return func(*args, **kwargs)
    except Exception as e:
        log_exception(f"API调用失败: {context}", e, level="ERROR")
        st.error(f"❌ {context}时出错: {e}")
        return fallback if fallback is not None else ([] if "抽取" in context else None)

@contextlib.contextmanager
def timeout(seconds: int = 10):
    """
    P1-3: 超时上下文管理器，防止应用挂起。
    支持 Windows 和 Unix 系统。
    
    Args:
        seconds: 超时秒数
        
    Raises:
        TimeoutError: 如果超时
    """
    # Windows 不支持 SIGALRM，使用 threading.Timer
    if not hasattr(signal, 'SIGALRM'):
        timer = None
        timeout_occurred = threading.Event()
        
        def _timeout_handler():
            timeout_occurred.set()
        
        timer = threading.Timer(seconds, _timeout_handler)
        timer.start()
        try:
            yield
            if timeout_occurred.is_set():
                raise TimeoutError(f"操作超时 ({seconds}秒)")
        finally:
            timer.cancel()
    else:
        # Unix/Linux 系统使用 signal
        def handler(signum, frame):
            raise TimeoutError(f"操作超时 ({seconds}秒)")
        
        old_handler = signal.signal(signal.SIGALRM, handler)
        signal.alarm(seconds)
        try:
            yield
        finally:
            signal.alarm(0)
            signal.signal(signal.SIGALRM, old_handler)

def cached_query(ttl: int = 300):
    """
    P1-2: 缓存装饰器，避免重复查询相同数据。
    
    Args:
        ttl: 缓存有效期（秒）
    """
    def decorator(func):
        cache = {'data': None, 'timestamp': 0}
        
        @functools.wraps(func)
        def wrapper(*args, **kwargs):
            now = time.time()
            # 检查缓存是否有效
            if cache['data'] is not None and (now - cache['timestamp']) < ttl:
                return cache['data']
            
            # 缓存过期，重新查询
            result = func(*args, **kwargs)
            cache['data'] = result
            cache['timestamp'] = now
            return result
        
        wrapper.invalidate = lambda: cache.update({'data': None, 'timestamp': 0})
        return wrapper
    return decorator

def log_performance(threshold_ms: int = 1000):
    """
    P2-2: 性能监控装饰器。
    
    Args:
        threshold_ms: 超过该时间（毫秒）时记录警告
    """
    def decorator(func):
        @functools.wraps(func)
        def wrapper(*args, **kwargs):
            start = time.time()
            result = func(*args, **kwargs)
            elapsed_ms = (time.time() - start) * 1000
            
            if elapsed_ms > threshold_ms:
                log_event(
                    "WARNING",
                    f"slow_function: {func.__name__}",
                    elapsed_ms=f"{elapsed_ms:.1f}ms"
                )
            
            return result
        return wrapper
    return decorator

# ==== third-party ====
try:
    from docx import Document 
except Exception:
    Document = None

# ========== 页面设置 ==========
st.set_page_config(page_title="齐译知识增强型翻译支持与管理系统 V1.0", layout="wide")

# ========== kb_dynamic  ==========
KBEmbedder = None
recommend_for_segment = None
build_prompt_strict = None
build_prompt_soft = None
try:
    from kb_dynamic import (
        KBEmbedder as _KBEmbedder,
        recommend_for_segment as _recommend_for_segment,
        build_prompt_strict as _build_prompt_strict,
        build_prompt_soft as _build_prompt_soft,
    )
    KBEmbedder = _KBEmbedder
    recommend_for_segment = _recommend_for_segment
    build_prompt_strict = _build_prompt_strict
    build_prompt_soft = _build_prompt_soft
except Exception as e:
    log_exception("Failed to import kb_dynamic; semantic features disabled", e, level="WARNING")

# ========== 工具函数 ==========
# ========== 统一 Session 状态管理器 ==========
class SessionManager:
    """集中管理所有 Streamlit Session 状态，解决状态混乱问题"""
    
    # ===== Session Key 常量（集中定义，易维护） =====
    # 工作区相关
    WORKSPACE_ACTIVATED = "workspace_activated_{pid}"
    WORKSPACE_SRC = "workspace_src_{pid}"
    WORKSPACE_DRAFT = "workspace_draft_{pid}"
    WORKSPACE_TERMS = "workspace_terms_{pid}"
    
    # 术语相关
    TERM_FINAL_MAP = "term_final_map_{pid}"
    TERM_FINAL_DF = "term_final_df_{pid}"
    TERM_SCAN = "term_scan_{pid}"
    TERM_SCAN_DF = "term_scan_df_{pid}"
    TERM_SCAN_STATUS = "term_scan_status_{pid}"
    TERM_SCAN_TIME = "term_scan_time_{pid}"
    
    # 搜索/查询相关
    FS_SEARCH_KW_RUN = "fs_search_kw_run_{pid}"
    LANG = "lang_{pid}"
    USE_SEM = "use_sem_{pid}"
    SCOPE = "scope_{pid}"
    EXTRA_PROMPT = "extra_prompt_{pid}"
    LAST_TRANSLATION = "last_translation_{pid}"
    POLY_SEL = "poly_sel_{pid}_{idx}"
    TRANS_FILE_SEL = "trans_file_sel_{pid}"
    FILE_SEL = "file_sel_{pid}"
    
    # 全局状态
    CORPUS_REFS = "corpus_refs"
    COR_USE_REF = "cor_use_ref"
    DB_CONN = "db_conn"
    KB_EMBEDDER = "kb_embedder"
    SCOPE_NEWPROJ = "scope_newproj"
    USE_SEMANTIC_GLOBAL = "use_semantic_global"
    TRANS_MGMT_TAB_PENDING = "trans_mgmt_tab_pending"
    TRANS_MGMT_TAB = "trans_mgmt_tab"
    ACTIVE_TRANSLATION_PID = "active_translation_pid"
    ACTIVE_TRANSLATION_SRC = "active_translation_src"
    CORPUS_TARGET_PROJECT = "corpus_target_project"
    CORPUS_TARGET_LABEL = "corpus_target_label"
    
    @staticmethod
    def _build_key(template: str, **kwargs) -> str:
        """构建session key，支持参数化"""
        try:
            return template.format(**kwargs)
        except KeyError as e:
            log_event("ERROR", f"SessionManager key format error: {e}", template=template, kwargs=kwargs)
            return template
    
    @staticmethod
    def get(key_template: str, default=None, **kwargs):
        """获取session值，自动null检查和类型感知默认值"""
        key = SessionManager._build_key(key_template, **kwargs)
        value = st.session_state.get(key)
        
        # 如果值为None，返回类型感知的默认值
        if value is None:
            if default is not None:
                return default
            # 根据key推断默认类型
            if "list" in key.lower() or "items" in key.lower():
                return []
            if "dict" in key.lower() or "map" in key.lower():
                return {}
            if "enabled" in key.lower() or "activated" in key.lower():
                return False
            if "text" in key.lower() or "prompt" in key.lower():
                return ""
        return value
    
    @staticmethod
    def set(key_template: str, value, **kwargs):
        """设置session值"""
        key = SessionManager._build_key(key_template, **kwargs)
        st.session_state[key] = value
    
    @staticmethod
    def pop(key_template: str, default=None, **kwargs):
        """删除并返回session值"""
        key = SessionManager._build_key(key_template, **kwargs)
        return st.session_state.pop(key, default)
    
    @staticmethod
    def pop_group(key_templates: list[str], **kwargs):
        """批量删除相关联的keys（解决删除5-10个相关key的问题）"""
        for template in key_templates:
            SessionManager.pop(template, **kwargs)
    
    @staticmethod
    def clear_workspace(pid):
        """一键清除某项目的整个工作区状态"""
        SessionManager.pop_group([
            SessionManager.WORKSPACE_ACTIVATED,
            SessionManager.WORKSPACE_SRC,
            SessionManager.WORKSPACE_DRAFT,
            SessionManager.WORKSPACE_TERMS,
            SessionManager.TERM_FINAL_MAP,
            SessionManager.TERM_FINAL_DF,
            SessionManager.TERM_SCAN,
            SessionManager.TERM_SCAN_DF,
            SessionManager.TERM_SCAN_STATUS,
            SessionManager.TERM_SCAN_TIME,
            SessionManager.FS_SEARCH_KW_RUN,
            SessionManager.LANG,
            SessionManager.USE_SEM,
            SessionManager.SCOPE,
            SessionManager.EXTRA_PROMPT,
            SessionManager.LAST_TRANSLATION,
        ], pid=pid)
    
    @staticmethod
    def clear_term_scan(pid):
        """清除术语扫描相关状态，用于重新扫描"""
        SessionManager.pop_group([
            SessionManager.TERM_SCAN,
            SessionManager.TERM_SCAN_DF,
            SessionManager.TERM_SCAN_STATUS,
            SessionManager.TERM_SCAN_TIME,
            SessionManager.TERM_FINAL_MAP,
            SessionManager.TERM_FINAL_DF,
        ], pid=pid)

# ===== 向后兼容：保留旧的函数接口（逐步迁移） =====
def make_sk(prefix: str):
    return lambda name, id=None: f"{prefix}_{name}_{id}" if id else f"{prefix}_{name}"
sk = make_sk("global")

def render_table(df, *, key=None, hide_index=True, editable=False):
    try:
        if editable:
            return st.data_editor(
                df,
                hide_index=hide_index,
                key=key,
            )
        if key is not None:
            return st.data_editor(
                df,
                hide_index=hide_index,
                disabled=True,
                key=key,
            )
        return st.dataframe(df, hide_index=hide_index)
    except TypeError:
        return st.data_editor(
            df,
            hide_index=hide_index,
            disabled=not editable,
            key=key,
        )
# ========== 文本高亮函数 ==========
def highlight_terms(text: str, term_pairs: list):
    if not term_pairs:
        return text

    import re
    safe = text

    for s, t in term_pairs:
        if not s:
            continue
        # source term 黄色
        safe = re.sub(
            re.escape(s),
            fr"<span style='background: #fff3b0'>{s}</span>",
            safe,
            flags=re.IGNORECASE
        )
        #  target term 淡绿
        if t:
            safe = re.sub(
                re.escape(t),
                fr"<span style='background: #d4f6d4'>{t}</span>",
                safe,
                flags=re.IGNORECASE
            )

    return safe

def render_index_manager(st, conn, cur):
    st.title("Semantic Index Manager")

    # === 1. 项目列表 ===
    rows = execute_write(
        """
        SELECT i.id,
               IFNULL(i.title,''),
               IFNULL(i.domain,''),
               COUNT(DISTINCT c.id)                                           AS corpus_cnt,
               SUM(CASE WHEN c.note LIKE 'from trans_ext%%' THEN 1 ELSE 0 END) AS hist_cnt
        FROM items i
        LEFT JOIN corpus c ON c.project_id = i.id
        GROUP BY i.id, i.title, i.domain
        ORDER BY i.id DESC
        LIMIT 500;
        """
    ).fetchall()

    if not rows:
        st.info("No projects yet. Please create one first.")
        return

    import pandas as pd

    df_proj = pd.DataFrame(
        [
            {
                "Project ID": r[0],
                "Project Name": r[1],
                "Domain": r[2],
                "Corpus Count": r[3],
                "From Translation History": r[4] or 0,
            }
            for r in rows
        ]
    )
    st.dataframe(df_proj, width='stretch')

    # === 2. 选择项目 ===
    proj_options = {f"[{r[0]}] {r[1] or '(Untitled)'}": r[0] for r in rows}
    proj_label = st.selectbox(
        "Select a project to view/rebuild index",
        ["(Select)"] + list(proj_options.keys()),
    )
    pid_sel = proj_options.get(proj_label)

    if not pid_sel:
        st.info("Please select a project to view index details.")
        return

    idx_total = idx_corpus = idx_hist = idx_other = 0
    try:
        mode, index_obj, mapping, vecs = _load_index(int(pid_sel))
        if isinstance(mapping, list):
            idx_total = len(mapping)
            for m in mapping:
                src_tag = (m.get("source") or "").lower()
                if src_tag == "history":
                    idx_hist += 1
                elif src_tag in ("", "corpus"):
                    idx_corpus += 1
                else:
                    idx_other += 1
    except Exception as e:
        st.warning(f"Failed to read index: {e}")

    st.write(f"- Current index count: **{idx_total}**")
    st.write(f"- From corpus: **{idx_corpus}**")
    st.write(f"- From translation history: **{idx_hist}**")
    st.write(f"- Other/unknown: **{idx_other}**")

    row_sel = [r for r in rows if r[0] == pid_sel][0]
    st.write(
        f"- Corpus rows in DB: **{row_sel[3]}** "
        f"(from history: **{row_sel[4] or 0}**)")

    c1, c2 = st.columns(2)

    with c1:
        if st.button("Rebuild current project index", type="primary", key=f"rebuild_idx_{pid_sel}"):
            res = rebuild_project_semantic_index(pid_sel)
            if res.get("ok"):
                st.success(f"Index rebuilt: +{res['added']} new, total {res['total']}.")
            else:
                st.error("Rebuild failed. Check logs.")

    with c2:
        if st.button("Batch rebuild all listed project indexes", key="rebuild_all_idx"):
            ok_cnt = fail_cnt = 0
            for r in rows:
                pid = r[0]
                res = rebuild_project_semantic_index(pid)
                if res.get("ok"):
                    ok_cnt += 1
                else:
                    fail_cnt += 1
            st.success(f"Batch rebuild finished: {ok_cnt} ok, {fail_cnt} failed.")

# 缓存函数定义在外部，避免每次调用都创建新缓存
@cached_query(ttl=300)
def _fetch_terms_from_db_cached(pid: int, use_dynamic: bool):
    """缓存术语查询结果（5分钟有效期）"""
    if use_dynamic:
        rows = execute_write(
            """
            SELECT source_term, target_term, domain, project_id,
                   CASE WHEN project_id = ? THEN 'static' ELSE 'dynamic' END as origin
            FROM term_ext
            WHERE project_id = ? OR project_id IS NULL
            ORDER BY CASE WHEN project_id = ? THEN 0 ELSE 1 END
            """,
            (pid, pid, pid),
        ).fetchall()
    else:
        rows = execute_write(
            """
            SELECT source_term, target_term, domain, project_id, 'static' as origin
            FROM term_ext
            WHERE project_id = ?
            """,
            (pid,),
        ).fetchall()
    return rows

def _normalize_term(term: str) -> str:
    """
    规范化术语用于匹配：
    1. 去除前后空格
    2. 转小写
    3. 移除多余空格和特殊字符（包括中文空格）
    """
    if not term:
        return ""
    t = str(term).strip().lower()
    # 移除多余空格（包括中英文空格、制表符等）
    t = re.sub(r'[\s\u3000\u00A0]+', ' ', t)
    # 对于中文术语，也可以考虑移除空格
    if all(ord(c) >= 0x4E00 and ord(c) <= 0x9FFF or c == ' ' for c in t.replace(' ', '')):
        # 纯中文术语，移除所有空格
        t = t.replace(' ', '')
    return t

def _find_term_in_map(src_term: str, term_map: dict[str, str]) -> str | None:
    """
    从术语图中查找源术语的目标术语，支持多种匹配策略。
    
    Args:
        src_term: 源术语
        term_map: {源术语: 目标术语} 的字典
    
    Returns:
        匹配的目标术语，或 None
    """
    if not src_term or not term_map:
        return None
    
    src_term = src_term.strip()
    if not src_term:
        return None
    
    # 策略1：精确匹配
    if src_term in term_map:
        return term_map[src_term]
    
    # 策略2：不区分大小写精确匹配
    src_lower = src_term.lower()
    for k, v in term_map.items():
        if k.lower() == src_lower:
            return v
    
    # 策略3：规范化后匹配
    src_norm = _normalize_term(src_term)
    if not src_norm:
        return None
    
    for k, v in term_map.items():
        if _normalize_term(k) == src_norm:
            return v
    
    # 策略4：子串匹配（对于复合词）
    for k, v in term_map.items():
        if src_norm in _normalize_term(k) or _normalize_term(k) in src_norm:
            # 优先选择更短的匹配（更精确）
            if len(k) < len(src_term) * 1.5 and len(src_term) < len(k) * 1.5:
                return v
    
    return None

def get_terms_for_project(cur: sqlite3.Cursor, pid: int, use_dynamic: bool = True) -> tuple[dict[str, str], list[dict]]:
    """
    获取项目术语映射表（支持动态补充其他项目的术语）。

    Args:
        cur: 数据库游标
        pid: 项目 ID
        use_dynamic: 是否包含同领域的动态术语

    Returns:
        (term_map, term_meta)
        - term_map: {源术语: 目标术语}
        - term_meta: [{source_term, target_term, domain, origin, project_id}, ...]
    """
    # ✅ 优化: 单次查询替代N+1查询 + 缓存避免重复
    rows = _fetch_terms_from_db_cached(pid, use_dynamic)

    # 3) 去重和组织数据
    dedup: dict[tuple[str, str], tuple[str, str, str, str, int | None]] = {}

    for row in rows:
        s, t, d, pid_term, origin = row
        if not s:
            continue
        
        s_raw = (s or "").strip()
        t_raw = (t or "").strip()
        d_raw = (d or "").strip()
        key = (s_raw.lower(), d_raw.lower())
        
        # 保留第一个出现的（静态优先）
        if key not in dedup:
            dedup[key] = (s_raw, t_raw, d_raw, origin, pid_term if pid_term is not None else None)

    # 4) 拼term_map 和 term_meta
    term_map: dict[str, str] = {}
    term_meta: list[dict] = []

    for (_s_lc, _d_lc), (s_raw, t_raw, d_raw, origin, pid_term) in dedup.items():
        if not s_raw:
            continue
        term_map[s_raw] = t_raw
        term_meta.append(
            {
                "source_term": s_raw,
                "target_term": t_raw,
                "domain": d_raw,
                "origin": origin,
                "project_id": pid_term,
            }
        )

    return term_map, term_meta

# ======= 轻量术语候选=======

def register_project_file(cur, conn, project_id: int, file_name: str, data_bytes: bytes) -> str | None:
    """
    P2-1: 添加类型注解
    注册项目文件。
    """
    if not project_id or not data_bytes:
        log_event("WARNING", "register_project_file: missing project_id or data", project_id=project_id, data_len=len(data_bytes) if data_bytes else 0)
        return None
    
    try:
        safe_name = os.path.basename(file_name) or f"project_{project_id}_file"
        proj_dir = os.path.join(PROJECT_DIR, f"project_{project_id}")
        os.makedirs(proj_dir, exist_ok=True)
        stamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        uniq_name = f"{stamp}_{uuid.uuid4().hex[:6]}_{safe_name}"
        full_path = os.path.join(proj_dir, uniq_name)
        
        # 先写入文件，验证成功后再记录到数据库
        with open(full_path, "wb") as f:
            f.write(data_bytes)
        
        # 验证文件已写入
        if not os.path.exists(full_path) or os.path.getsize(full_path) != len(data_bytes):
            log_event("ERROR", "register_project_file: file write verification failed", path=full_path, expected_size=len(data_bytes), actual_size=os.path.getsize(full_path) if os.path.exists(full_path) else 0)
            return None
        
        execute_write(
            """
            INSERT INTO project_files (project_id, file_path, file_name)
            VALUES (?, ?, ?)
            """,
            (project_id, full_path, safe_name),
        )
        return full_path
    except Exception as e:
        log_exception("register_project_file: failed to register file", e, level="ERROR")
        return None

def fetch_project_files(cur, project_id: int) -> list[tuple]:
    """
    P2-1: 添加类型注解
    获取项目文件列表。
    """
    if not project_id:
        return []
    rows = execute_write(
        """
        SELECT id, IFNULL(file_name,''), IFNULL(file_path,''), IFNULL(uploaded_at,'')
        FROM project_files
        WHERE project_id=?
        ORDER BY id DESC
        """,
        (project_id,),
    ).fetchall()
    items = []
    for fid, name, path, uploaded in rows:
        if not path:
            continue
        display = name or os.path.basename(path) or f"file_{fid}"
        items.append(
            {
                "id": fid,
                "name": display,
                "path": path,
                "uploaded_at": uploaded,
            }
        )
    return items

def ensure_legacy_file_record(cur, conn, project_id, legacy_path):
    if not (project_id and legacy_path):
        return
    exists = execute_write(
        "SELECT 1 FROM project_files WHERE project_id=? AND file_path=?",
        (project_id, legacy_path),
    ).fetchone()
    if exists:
        return
    execute_write(
        """
        INSERT INTO project_files (project_id, file_path, file_name)
        VALUES (?, ?, ?)
        """,
        (project_id, legacy_path, os.path.basename(legacy_path) or None),
    )

def remove_project_file(cur, conn, file_id):
    row = cur.execute("SELECT file_path, project_id FROM project_files WHERE id=?", (file_id,)).fetchone()
    if row:
        path, pid = row
        # 如果该文件是当前项目的 src_path，先清空，避免被自动回写
        if pid:
            try:
                r = cur.execute("SELECT IFNULL(src_path,'') FROM item_ext WHERE item_id=?", (pid,)).fetchone()
                cur_path = (r[0] if r else "") or ""
                if cur_path and path:
                    if os.path.normcase(os.path.normpath(cur_path)) == os.path.normcase(os.path.normpath(path)):
                        execute_write("UPDATE item_ext SET src_path='' WHERE item_id=?", (pid,))
            except Exception as e:
                log_event("WARNING", "clear item_ext src_path failed", path=path, err=str(e))
        if path:
            if os.path.exists(path):
                try:
                    os.remove(path)
                except Exception as e:
                    log_event("ERROR", "remove project file failed", path=path, err=str(e))
                    # 删除失败时，不删除数据库记录，让用户重试
                    return False
            # 文件已不存在，或成功删除，则清理数据库记录
    execute_write("DELETE FROM project_files WHERE id=?", (file_id,))
    return True

def cleanup_project_files(cur, conn, project_id):
    rows = cur.execute("SELECT file_path FROM project_files WHERE project_id=?", (project_id,)).fetchall()
    for (fp,) in rows:
        if fp and os.path.exists(fp):
            try:
                os.remove(fp)
            except Exception as e:
                log_event("WARNING", "remove project file failed", path=fp, err=str(e))
    execute_write("DELETE FROM project_files WHERE project_id=?", (project_id,))

def _ensure_project_ref_map():
    refs = st.session_state.get(CORPUS_REFS)
    if isinstance(refs, dict):
        return refs
    st.session_state[CORPUS_REFS] = {}
    return st.session_state[CORPUS_REFS]

def _ensure_project_switch_map():
    switches = st.session_state.get(COR_USE_REF)
    if isinstance(switches, dict):
        return switches
    st.session_state[COR_USE_REF] = {}
    return st.session_state[COR_USE_REF]

def get_project_ref_ids(project_id: int | None) -> set[int]:
    if not project_id:
        return set()
    ref_map = _ensure_project_ref_map()
    ref_map.setdefault(project_id, set())
    return ref_map[project_id]

def get_project_fewshot_enabled(project_id: int | None) -> bool:
    if not project_id:
        return False
    switch_map = _ensure_project_switch_map()
    return bool(switch_map.get(project_id, False))

def set_project_fewshot_enabled(project_id: int | None, value: bool):
    if not project_id:
        return
    switch_map = _ensure_project_switch_map()
    switch_map[project_id] = bool(value)

def get_project_fewshot_examples(
    cur,
    project_id: int | None,
    *,
    limit: int | None = 5,
    require_enabled: bool = True,
):
    if not project_id:
        return []
    if require_enabled and not get_project_fewshot_enabled(project_id):
        return []

    ref_ids = list(get_project_ref_ids(project_id))
    if not ref_ids:
        return []

    ref_ids = sorted({int(rid) for rid in ref_ids if str(rid).isdigit()}, reverse=True)
    if limit is not None and len(ref_ids) > limit:
        ref_ids = ref_ids[:limit]

    qmarks = ",".join(["?"] * len(ref_ids))
    rows = execute_write(
        f"SELECT id, title, IFNULL(src_text,''), IFNULL(tgt_text,'') FROM corpus WHERE id IN ({qmarks})",
        ref_ids,
    ).fetchall()
    order_map = {rid: idx for idx, rid in enumerate(ref_ids)}
    rows.sort(key=lambda r: order_map.get(r[0], len(order_map)))

    examples = []
    for rid, title, src, tgt in rows:
        src_norm = (src or "").strip()
        tgt_norm = (tgt or "").strip()
        if not (src_norm and tgt_norm):
            continue
        examples.append(
            {
                "id": rid,
                "title": title or f"示例#{rid}",
                "src": src_norm,
                "tgt": tgt_norm,
            }
        )
    return examples

def render_translation_workspace(pid, selected_src_path, cur, conn):
    if not SessionManager.get(SessionManager.WORKSPACE_DRAFT, pid=pid) or not SessionManager.get(
        SessionManager.WORKSPACE_ACTIVATED, pid=pid, default=False
    ):
        return

    st.markdown("##  翻译工作台")

    blocks = SessionManager.get(SessionManager.WORKSPACE_SRC, default=[], pid=pid)
    draft = SessionManager.get(SessionManager.WORKSPACE_DRAFT, default=[], pid=pid)
    terms = SessionManager.get(SessionManager.WORKSPACE_TERMS, default=[], pid=pid)
    poly_terms = [r[0] for r in cur.execute(
        "SELECT source_term FROM poly_term "
        "WHERE (SELECT COUNT(1) FROM poly_sense s WHERE s.term_id=poly_term.id) >= 2 "
        "ORDER BY source_term"
    ).fetchall()]

    if not blocks or not draft:
        st.info("当前暂无草稿，请先点击“进入翻译工作台”生成初稿。")
        return

    edited_blocks = []
    for i, (src, trg) in enumerate(zip(blocks, draft), 1):
        st.markdown(f"### 段落 {i}")
        col1, col2 = st.columns(2)
        with col1:
            st.markdown("**原文**")
            terms_in_block = [t for t in poly_terms if t and t in src]
            if terms_in_block:
                terms_in_block = sorted(terms_in_block, key=len, reverse=True)
                pat = "|".join(re.escape(t) for t in terms_in_block)
                highlighted_src = re.sub(
                    pat,
                    lambda m: (
                        "<span style='background:#ffd6d6;color:#b00020;"
                        "font-weight:600;padding:0 2px;border-radius:3px'>"
                        f"{m.group(0)}</span>"
                    ),
                    src,
                )
            else:
                highlighted_src = src
            st.markdown(
                f"<div style='padding:8px;border:1px solid #ccc;background:#f8f8f8'>"
                f"{highlighted_src}</div>",
                unsafe_allow_html=True,
            )
            if terms_in_block:
                st.caption("红色高亮为一词多义，点击下方词条查看译入语列表")
                sel_key = SessionManager._build_key(SessionManager.POLY_SEL, pid=pid, idx=i)
                for term in terms_in_block:
                    if st.button(f"查看：{term}", key=f"poly_btn_{pid}_{i}_{term}"):
                        st.session_state[sel_key] = term
                sel_term = st.session_state.get(sel_key)
                if sel_term:
                    term_row = cur.execute(
                        "SELECT id FROM poly_term WHERE source_term=?",
                        (sel_term,),
                    ).fetchone()
                    if term_row:
                        senses = cur.execute(
                            "SELECT target_term, IFNULL(example_src,''), IFNULL(example_tgt,'') "
                            "FROM poly_sense WHERE term_id=? ORDER BY id",
                            (term_row[0],),
                        ).fetchall()
                        if senses:
                            df_sense = pd.DataFrame(
                                senses,
                                columns=["译入语", "例句(源)", "例句(译)"],
                            )
                            df_sense.index = range(1, len(df_sense) + 1)
                            st.dataframe(df_sense, width='stretch')
        with col2:
            st.markdown("**译文（可编辑）**")
            new_trg = st.text_area(
                label="编辑后的译文",
                value=trg,
                key=f"edit_{pid}_{i}",
                height=120,
            )
            edited_blocks.append(new_trg)

            if "highlight_terms" in globals():
                highlighted = highlight_terms(new_trg, terms)
                st.markdown("术语高亮：")
                st.markdown(
                    f"<div style='padding:8px;border:1px solid #ccc;"
                    f"background:#f0fff0'>{highlighted}</div>",
                    unsafe_allow_html=True,
                )

    if st.button(" 确认生成最终译文", key=f"confirm_{pid}", type="primary"):
        final_text = "\n\n".join(edited_blocks)
        lang_pair_val = SessionManager.get(SessionManager.LANG, default="中英", pid=pid)
        execute_write(
            """
            INSERT INTO trans_ext (
                project_id, src_path, lang_pair, mode, output_text, created_at
            )
            VALUES (?, ?, ?, ?, ?, datetime('now'))
        """,
            (pid, selected_src_path, lang_pair_val, "工作台模式", final_text),
        )

        st.success("译文已生成并写入历史！")

        SessionManager.pop(SessionManager.WORKSPACE_SRC, None, pid=pid)
        SessionManager.pop(SessionManager.WORKSPACE_DRAFT, None, pid=pid)
        SessionManager.pop(SessionManager.WORKSPACE_TERMS, None, pid=pid)
        SessionManager.pop(SessionManager.WORKSPACE_ACTIVATED, None, pid=pid)

def _save_translation_to_history(pid, src_path, final_text, blocks_src, all_results, cur, conn, locked_terms):
    """
    辅助函数：保存翻译历史到数据库
    在 render_translation_outputs 中调用以确保数据被保存
    """
    # 检查是否已经保存过（通过检查最近的相同内容）
    try:
        # 避免重复保存：检查最近是否有完全相同的记录
        recent = cur.execute("""
            SELECT id, created_at FROM trans_ext 
            WHERE project_id=? AND output_text=?
            ORDER BY id DESC LIMIT 1
        """, (pid, final_text)).fetchone()
        
        if recent:
            # 检查是否是最近创建的（5秒内）
            recent_id, recent_time = recent
            try:
                recent_dt = datetime.strptime(recent_time, "%Y-%m-%d %H:%M:%S")
                time_diff = (datetime.now() - recent_dt).total_seconds()
                if time_diff < 5:
                    # 5秒内的重复，跳过保存
                    return
            except Exception:
                pass
    except Exception as e:
        log_event("WARNING", "check duplicate translation failed", err=str(e))
    
    # 计算分段数
    try:
        if blocks_src:
            seg_count = len(blocks_src)
        else:
            seg_count = len(split_sents(final_text, lang_hint="auto")) if final_text else 1
    except Exception as e:
        log_event("WARNING", "split_sents failed in save_history", err=str(e))
        seg_count = len(blocks_src) if blocks_src else 1
    
    # 获取语言对
    try:
        lang_pair_db = SessionManager.get(SessionManager.LANG, default="自动", pid=pid)
    except Exception:
        lang_pair_db = "自动"
    
    # 保存翻译历史
    try:
        execute_write("""
            INSERT INTO trans_ext (
                project_id, src_path, lang_pair, mode, output_text,
                stats_json, segments, term_hit_total, created_at
            )
            VALUES (?, ?, ?, ?, ?, ?, ?, ?, datetime('now'))
        """, (
            pid,
            src_path,
            lang_pair_db,
            "翻译流程",
            final_text,
            None,
            seg_count,
            None
        ))
        log_event("INFO", "translation history saved", pid=pid)
    except Exception as e:
        log_event("ERROR", "save translation history failed", err=str(e))
    
    # 尝试写入语料库（非关键操作，失败不影响）
    try:
        dom_val = _get_domain_for_proj(cur, int(pid))
    except Exception:
        dom_val = None
    
    try:
        scene_val = _get_scene_for_proj(cur, int(pid))
    except Exception:
        scene_val = None
    
    try:
        if all_results and blocks_src:
            # 获取项目标题
            try:
                proj_title_row = cur.execute("SELECT title FROM items WHERE id=?", (pid,)).fetchone()
                proj_title = (proj_title_row[0] if proj_title_row else None) or f"project_{pid}"
            except Exception:
                proj_title = f"project_{pid}"
            
            write_translation_segments_to_corpus(
                cur, conn, pid,
                title=proj_title,
                lang_pair=lang_pair_db,
                domain=dom_val,
                scene=scene_val,
                blocks_src=blocks_src,
                blocks_tgt=all_results,
            )
    except Exception as e:
        log_event("WARNING", "write corpus failed in save_history", err=str(e))

def render_translation_outputs(
    pid,
    project_title,
    selected_src_path,
    cur,
    conn,
    blocks_src_safe,
    all_results_safe,
    locked_terms,
):
    final_text = "\n\n".join(all_results_safe)

    # ===== 重要：在此处写入翻译历史（确保每次显示结果都会保存） =====
    # 这很关键，因为 Streamlit 页面重新渲染时，按钮状态会变化
    # 需要在显示结果时就保存数据，而不是依赖按钮状态
    _save_translation_to_history(
        pid, selected_src_path, final_text, blocks_src_safe, all_results_safe, cur, conn, locked_terms
    )

    try:
        term_map_report = locked_terms or {}
        df_rep = semantic_consistency_report(
            project_id=pid,
            blocks_src=blocks_src_safe,
            blocks_tgt=all_results_safe,
            term_map=term_map_report,
            topk=3,
            thr=0.70,
        )
        st.dataframe(df_rep, width='stretch')
    except Exception as e:
        st.caption(f"(一致性报告生成失败: {e})")

    proj_title = project_title or f"project_{pid}"
    st.download_button(
        "下载译文(TXT)",
        final_text or "",
        file_name=f"{proj_title}_译文.txt",
        mime="text/plain",
        key=f"dl_txt_{pid}",
    )
    if st.button("进入翻译工作台", key=f"workspace_after_{pid}", type="secondary"):
        SessionManager.set(SessionManager.WORKSPACE_ACTIVATED, True, pid=pid)
        SessionManager.set(SessionManager.WORKSPACE_SRC, blocks_src_safe, pid=pid)
        SessionManager.set(SessionManager.WORKSPACE_DRAFT, all_results_safe, pid=pid)
        locked_terms_ws = SessionManager.get(SessionManager.TERM_FINAL_MAP, default={}, pid=pid)
        if locked_terms_ws:
            SessionManager.set(SessionManager.WORKSPACE_TERMS, list(locked_terms_ws.items()), pid=pid)
        else:
            term_map_all_ws, _ = get_terms_for_project(cur, pid, use_dynamic=True)
            SessionManager.set(SessionManager.WORKSPACE_TERMS, list(term_map_all_ws.items()), pid=pid)

    render_translation_workspace(pid, selected_src_path, cur, conn)
def run_project_translation_ui(
    pid,
    project_title,
    src_path,
    conn,
    cur
):
    st.subheader(f" 项目：{project_title}")

    term_map, term_meta = get_terms_for_project(cur, pid, use_dynamic=True)
    proj_terms_all = term_map  # 给 _detect_hits 用

    ak, model = get_deepseek()
    if not ak:
        st.error("未检测到 DeepSeek Key.请在 `.streamlit/secrets.toml` 配置 [deepseek]")
        st.stop()
    selected_src_path = src_path
    if not selected_src_path or not os.path.exists(selected_src_path):
        st.error("缺少源文件")
        st.stop()

    st.markdown("### 1.段落切分")
    src_text = read_source_file(selected_src_path)
    blocks = split_paragraphs(src_text)
    if not blocks:
        st.error("源文件内容为空，或未识别到有效段落")
        st.stop()

    st.info(f"按段落切分，共 {len(blocks)} 段")

    # === 术语扫描/确认：先定术语表，再进入翻译 ===
    st.markdown("### 2.术语扫描与确认")
    term_scan_key = SessionManager._build_key(SessionManager.TERM_SCAN, pid=pid)
    term_final_key = SessionManager._build_key(SessionManager.TERM_FINAL_MAP, pid=pid)
    term_final_df_key = SessionManager._build_key(SessionManager.TERM_FINAL_DF, pid=pid)
    term_scan_df_key = SessionManager._build_key(SessionManager.TERM_SCAN_DF, pid=pid)
    term_scan_status_key = SessionManager._build_key(SessionManager.TERM_SCAN_STATUS, pid=pid)
    term_scan_time_key = SessionManager._build_key(SessionManager.TERM_SCAN_TIME, pid=pid)

    # 现有术语(系统表)映射，供“备选1”使用
    sys_term_map = dict(term_map)
    sys_term_map_lower = {str(k).lower(): v for k, v in sys_term_map.items()}

    with st.expander(" 术语扫描与确认（先建表，后翻译）", expanded=True):
        st.markdown(
            "- 步骤1：用 DeepSeek 抽取全文术语\n"
            "- 步骤2：对照系统术语表填入“备选1”，DeepSeek 推荐为“备选2”，可手填“备选3”并确定“最终译”\n"
            "- 步骤3：锁定术语表，作为翻译约束；如需新增术语，先更新表再翻译"
        )

        c_scan, c_del = st.columns([1, 1])
        with c_scan:
            scan_click = st.button("① DeepSeek 抽取术语", key=f"scan_terms_{pid}")
        with c_del:
            del_click = st.button("删除已抽取术语", key=f"clear_scan_terms_{pid}")

        if del_click:
            st.session_state.pop(term_scan_key, None)
            st.session_state.pop(term_final_key, None)
            st.session_state.pop(term_final_df_key, None)
            st.session_state.pop(term_scan_status_key, None)
            st.session_state.pop(term_scan_time_key, None)
            st.success("已删除本次抽取的术语")
            st.rerun()

        if scan_click:
            st.session_state[term_scan_status_key] = "running"
            st.session_state[term_scan_time_key] = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
            with st.spinner("调用 DeepSeek 抽取术语中…"):
                res_scan = safe_api_call(
                    ds_extract_terms,
                    src_text, ak, model,
                    context="术语抽取",
                    src_lang="zh", tgt_lang="en"
                )
                if res_scan is None:
                    st.session_state[term_scan_status_key] = "error"
                else:
                    st.session_state[term_scan_key] = res_scan or []
                    # 清理旧的锁定表，避免用旧表
                    st.session_state.pop(term_final_key, None)
                    st.session_state.pop(term_final_df_key, None)
                    if res_scan:
                        st.session_state[term_scan_status_key] = f"ok:{len(res_scan)}"
                    else:
                        st.session_state[term_scan_status_key] = "empty"
                    st.session_state.pop(term_scan_df_key, None)
                    st.success(f"已抽取 {len(res_scan or [])} 条候选术语")

        scan_rows = st.session_state.get(term_scan_key, [])
        df_init = None

        if st.session_state.get(term_scan_df_key) is not None:
            try:
                df_init = st.session_state[term_scan_df_key]
            except Exception:
                df_init = None

        if scan_rows and df_init is None:
            rows_df = []
            for o in scan_rows:
                src_term = (o.get("source_term") or o.get("source") or "").strip()
                tgt_ds = (o.get("target_term") or o.get("target") or "").strip()
                tgt_ds_alt = (o.get("target_term_alt") or o.get("alt") or "").strip()
                example = (o.get("example") or "").strip()
                if not src_term:
                    continue
                try:
                    term_count = len(re.findall(re.escape(src_term), src_text, flags=re.IGNORECASE))
                except Exception:
                    term_count = 0
                
                # 使用改进的匹配函数查找库中的术语
                alt1 = _find_term_in_map(src_term, sys_term_map)
                final_guess = alt1 or tgt_ds or tgt_ds_alt
                rows_df.append({
                    "删除": False,
                    "使用": True,
                    "序号": 0,
                    "源术语": src_term,
                    "出现次数": term_count,
                    "例句": example,
                    "备选1_系统术语": alt1 or "",
                    "备选2_DeepSeek推荐": tgt_ds,
                    "备选3_DeepSeek备选": tgt_ds_alt,
                    "最终译": final_guess or "",
                })
            if rows_df:
                df_init = pd.DataFrame(rows_df)
                df_init["序号"] = range(1, len(df_init) + 1)

        if scan_rows:
            try:
                term_pairs = [(o.get("source_term") or o.get("source") or "", "") for o in scan_rows]
                hl = highlight_terms(src_text, term_pairs)
                st.markdown("**原文术语高亮预览**")
                st.markdown(hl, unsafe_allow_html=True)
            except Exception:
                st.markdown("**原文术语高亮预览**")
                st.markdown(src_text)

        if df_init is not None and not df_init.empty:
            if "序号" not in df_init.columns:
                df_init.insert(1, "序号", range(1, len(df_init) + 1))
            col_order = ["删除", "使用", "序号"] + [c for c in df_init.columns if c not in ("删除", "使用", "序号")]
            df_init = df_init[col_order]
            edited_df = st.data_editor(
                df_init,
                hide_index=True,
                num_rows="dynamic",
                column_config={
                    "源术语": st.column_config.TextColumn(disabled=True),
                    "使用": st.column_config.CheckboxColumn(help="勾选后才会锁定并用于翻译"),
                    "序号": st.column_config.NumberColumn(disabled=True),
                    "出现次数": st.column_config.NumberColumn(disabled=True),
                    "例句": st.column_config.TextColumn(disabled=True),
                    "备选1_系统术语": st.column_config.TextColumn(disabled=True),
                    "备选2_DeepSeek推荐": st.column_config.TextColumn(disabled=True),
                    "备选3_DeepSeek备选": st.column_config.TextColumn(disabled=True),
                    "最终译": st.column_config.TextColumn(help="作为翻译硬约束，必须唯一"),
                    "删除": st.column_config.CheckboxColumn(help="勾选后点击删除"),
                },
                key=f"term_editor_{pid}",
            )
            st.session_state[term_scan_df_key] = edited_df.copy()

            # 兼容部分版本 data_editor 勾选无效，提供多选删除入口
            delete_options = edited_df["源术语"].tolist()
            default_selected = [
                str(r.get("源术语") or "").strip()
                for _, r in edited_df.iterrows()
                if bool(r.get("删除"))
            ]
            selected_terms = st.multiselect(
                "选择要删除的术语（如果勾选无效，请用这里选择）",
                options=delete_options,
                default=default_selected,
                key=f"term_delete_sel_{pid}",
            )

            col_del, col_lock, col_reset = st.columns(3)
            with col_del:
                if st.button("删除所选术语", key=f"delete_terms_{pid}"):
                    del_terms = set(
                        str(r.get("源术语") or "").strip()
                        for _, r in edited_df.iterrows()
                        if bool(r.get("删除"))
                    )
                    del_terms.update([s for s in selected_terms if s])
                    if not del_terms:
                        st.info("未选择要删除的术语")
                    else:
                        edited_df = edited_df[~edited_df["源术语"].isin(del_terms)].copy()
                        edited_df["删除"] = False
                        if "使用" in edited_df.columns:
                            edited_df["使用"] = edited_df["使用"].fillna(True)
                        if "序号" in edited_df.columns:
                            edited_df["序号"] = range(1, len(edited_df) + 1)
                        st.session_state[term_scan_df_key] = edited_df
                        # 同步扫描结果，避免高亮与表不一致
                        scan_rows = [
                            o for o in scan_rows
                            if (o.get("source_term") or o.get("source") or "").strip()
                            not in del_terms
                        ]
                        st.session_state[term_scan_key] = scan_rows
                        st.session_state.pop(term_final_key, None)
                        st.success(f"已删除 {len(del_terms)} 条术语")
                        st.rerun()

            with col_lock:
                if st.button("② 锁定术语表（作为翻译约束）", key=f"lock_terms_{pid}", type="primary"):
                    final_map = {}
                    for _, r in edited_df.iterrows():
                        src = str(r.get("源术语") or "").strip()
                        final = str(r.get("最终译") or "").strip()
                        use_flag = r.get("使用")
                        if src and final and (use_flag is True or use_flag == 1):
                            final_map[src] = final
                    if not final_map:
                        st.error("未选择要使用的术语或“最终译”为空，请确认后再锁定。")
                    else:
                        st.session_state[term_final_key] = final_map
                        st.session_state[term_final_df_key] = edited_df
                        st.success(f"已锁定术语表 {len(final_map)} 条，将在翻译中强制使用。")
            with col_reset:
                if st.button("清空锁定/重新抽取", key=f"reset_terms_{pid}"):
                    st.session_state.pop(term_scan_key, None)
                    st.session_state.pop(term_final_key, None)
                    st.session_state.pop(term_final_df_key, None)
                    st.session_state.pop(term_scan_df_key, None)
                    st.rerun()

        last_status = st.session_state.get(term_scan_status_key, "未开始")
        last_time = st.session_state.get(term_scan_time_key, "-")
        st.caption(f"抽取状态: {last_status} | 时间: {last_time}")

    locked_terms = st.session_state.get(term_final_key) or {}
    if locked_terms:
        st.success(f"已锁定术语 {len(locked_terms)} 条。")
    else:
        st.warning("未锁定术语表：请先完成术语抽取与锁定。")

    # --- 参考例句选择(翻译前、术语选择后) ---
    st.markdown("### 3.例句查询")
    ref_ids = get_project_ref_ids(pid)
    search_kw_key = f"fs_search_kw_{pid}"
    search_kw = st.text_input(
        "",
        key=search_kw_key,
        label_visibility="collapsed",
    ).strip()
    if st.button("搜索", key=f"fs_search_btn_{pid}"):
        SessionManager.set(SessionManager.FS_SEARCH_KW_RUN, search_kw, pid=pid)

    search_kw_run = SessionManager.get(SessionManager.FS_SEARCH_KW_RUN, default="", pid=pid).strip()
    if search_kw_run:
        like_kw = f"%{search_kw_run}%"
        rows = execute_write(
            "SELECT id, title, IFNULL(src_text,''), IFNULL(tgt_text,'') "
            "FROM corpus "
            "WHERE title LIKE ? OR src_text LIKE ? OR tgt_text LIKE ? "
            "ORDER BY id DESC LIMIT 50",
            (like_kw, like_kw, like_kw),
        ).fetchall()
        if not rows:
            st.info("未找到匹配语料")
        else:
            data = []
            for rid, title, src, tgt in rows:
                data.append(
                    {
                        "加入": False,
                        "ID": rid,
                        "标题": title or f"语料#{rid}",
                        "原文": (src or "")[:120],
                        "译文": (tgt or "")[:120],
                    }
                )
            df = pd.DataFrame(data)
            table_key = f"fs_search_table_{pid}_{abs(hash(search_kw_run))}"
            edited = st.data_editor(
                df,
                hide_index=True,
                width='stretch',
                key=table_key,
                column_config={
                    "加入": st.column_config.CheckboxColumn("加入"),
                },
            )
            add_ids = {
                int(row["ID"])
                for _, row in edited.iterrows()
                if row.get("加入") is True
            }
            new_ids = add_ids - ref_ids
            if new_ids:
                ref_ids.update(new_ids)
                st.success(f"已加入 {len(new_ids)} 条参考例句")

    if ref_ids:
        ids_sorted = sorted({int(x) for x in ref_ids}, reverse=True)
        qmarks = ",".join(["?"] * len(ids_sorted))
        rows = execute_write(
            f"SELECT id, title, IFNULL(src_text,''), IFNULL(tgt_text,'') FROM corpus WHERE id IN ({qmarks})",
            ids_sorted,
        ).fetchall()
        order_map = {rid: idx for idx, rid in enumerate(ids_sorted)}
        rows.sort(key=lambda r: order_map.get(r[0], len(order_map)))
        with st.expander("查看 Few-shot 示例预览", expanded=False):
            for rid, title, src, tgt in rows[:10]:
                if not (src and tgt):
                    continue
                st.markdown(f"**[{rid}] {title or f'示例#{rid}'}**")
                st.markdown(f"源文:\n{src}\n\n译文:\n{tgt}\n---")

    st.markdown("### 4.指令")
    extra_prompt_key = SessionManager._build_key(SessionManager.EXTRA_PROMPT, pid=pid)
    extra_prompt = st.text_area(
        "额外提示(可选，将追加到翻译指令中)",
        value=st.session_state.get(extra_prompt_key, ""),
        height=80,
        key=extra_prompt_key,
    )

    lang_pair_val = SessionManager.get(SessionManager.LANG, default="中译英", pid=pid)
    use_semantic  = bool(SessionManager.get(SessionManager.USE_SEM, default=True, pid=pid))
    scope_val = SessionManager.get(
        SessionManager.SCOPE,
        default=st.session_state.get(SCOPE_NEWPROJ, "project"),
        pid=pid,
    )

    # 4) 循环翻译（统一走 translate_block_with_kb 管线）
    fewshot_examples = get_project_fewshot_examples(cur, pid, limit=5, require_enabled=False)
    if fewshot_examples:
        with st.expander(" Few-shot 参考示例(项目级注入)", expanded=False):
            for ex in fewshot_examples:
                st.markdown(
                    f"**{ex['title']}**\\n\\n源文:\\n{ex['src']}\\n\\n译文:\\n{ex['tgt']}\\n---"
                )

    start_translation = st.button("开始翻译", key=f"start_trans_{pid}", type="primary")
    last_translation_key = SessionManager._build_key(SessionManager.LAST_TRANSLATION, pid=pid)
    if not start_translation:
        last_payload = SessionManager.get(SessionManager.LAST_TRANSLATION, pid=pid)
        if last_payload:
            render_translation_outputs(
                pid,
                project_title,
                selected_src_path,
                cur,
                conn,
                last_payload.get("blocks", []),
                last_payload.get("results", []),
                locked_terms,
            )
        return
    if not locked_terms:
        st.error("未锁定术语表，已终止。请先完成术语抽取与锁定。")
        return

    all_results_local = []

    for i, blk in enumerate(blocks, start=1):
        blk = str(blk or "").strip()
        if not blk:
            continue

        res = translate_block_with_kb(
            cur=cur,
            project_id=pid,
            block_text=blk,
            lang_pair=lang_pair_val,
            ak=ak,
            model=model,
            use_semantic=use_semantic,
            scope=scope_val,
            fewshot_examples=fewshot_examples,
            locked_term_map=locked_terms,
            extra_prompt=extra_prompt,
        )
        out_text = res["tgt"]
        term_map_all = res["term_map_all"]
        terms_in_block = res.get("terms_in_block", {})
        terms_corpus = res.get("terms_corpus_dyn", {})
        terms_final = res.get("terms_final", {})
        ref_context = res["ref_context"]
        poly_hint = res.get("poly_hint", "")
        violated = res["violated_terms"]

        with st.expander(f"段落 {i} 术语命中与注入", expanded=False):
            if term_map_all:
                df_all = pd.DataFrame(
                    [
                        {
                            "源术语": s,
                            "目标术语": t,
                            "命中当前段": s in terms_in_block,
                            "命中参考例句": s in terms_corpus,
                            "已注入Prompt": s in terms_final,
                        }
                        for s, t in term_map_all.items()
                    ]
                )
                st.dataframe(df_all, width='stretch')

                corpus_only = [
                    {"源术语": s, "目标术语": t}
                    for s, t in terms_corpus.items()
                ]
                if corpus_only:
                    st.markdown("**仅来自参考例句的术语(当前段未出现)**")
                    st.dataframe(pd.DataFrame(corpus_only), width='stretch')

            if ref_context:
                st.text(ref_context[:1500])

        with st.expander("一词多义提示预览", expanded=False):
            if poly_hint:
                st.text(poly_hint)
            else:
                st.caption("未命中一词多义或无可用提示。")

        all_results_local.append(out_text)
        st.write(f"已完成第{i}段")

        if violated:
            st.warning("可能未遵守术语(粗检): " + ", ".join(violated))

    # ===== 译后:一致性报告 =====
    all_results_safe = list(all_results_local)
    blocks_src_safe  = list(blocks if 'blocks' in locals() else [])
    if len(blocks_src_safe) != len(all_results_safe):
        n = min(len(blocks_src_safe), len(all_results_safe))
        blocks_src_safe  = blocks_src_safe[:n]
        all_results_safe = all_results_safe[:n]

    final_text = "\n\n".join(all_results_safe)

    SessionManager.set(
        SessionManager.LAST_TRANSLATION,
        {
            "blocks": blocks_src_safe,
            "results": all_results_safe,
            "final_text": final_text,
            "src_path": selected_src_path,
        },
        pid=pid,
    )

    try:
        term_map_report = locked_terms or {}
        df_rep = semantic_consistency_report(
            project_id=pid,
            blocks_src=blocks_src_safe,
            blocks_tgt=all_results_safe,
            term_map=term_map_report,
            topk=3,
            thr=0.70,
        )
        st.dataframe(df_rep, width='stretch')
    except Exception as e:
        st.caption(f"(一致性报告生成失败: {e})")

    render_translation_outputs(
        pid,
        project_title,
        selected_src_path,
        cur,
        conn,
        blocks_src_safe,
        all_results_safe,
        locked_terms,
    )

def quick_diagnose_vectors(pid: int):
    try:
        mode, index, mapping, vecs = _load_index(pid)
        if mode == "none":
            dom = None
            try:
                if "cur" in globals():
                    dom = _get_domain_for_proj(cur, int(pid))  # type: ignore[name-defined]
            except Exception:
                dom = None
            dom_key = _norm_domain_key(dom)
            st.warning(
                f"项目 {pid} 尚未建立向量索引(semantic_index/{dom_key}/bilingual 下无索引文件)。"
            )
            return
        msg = f"索引模式: {mode}; 映射条数: {len(mapping)}"

        if mode == "faiss" and index is not None:
            msg += f"; FAISS ntotal: {index.ntotal}"
        elif vecs is not None:
            msg += f"; NPY shape: {getattr(vecs, 'shape', None)}"
        st.info(msg)

        # 抽样验证映射的 corpus_id 是否都能回查到文本
        bad = 0
        for m in mapping[:10]:
            cid = int(m.get("corpus_id") or -1)
            row = cur.execute("SELECT id FROM corpus WHERE id=?", (cid,)).fetchone()
            if not row:
                bad += 1
        if bad:
            st.warning(f"映射中有 {bad} 条 corpus_id 无法回查.请考虑重建索引。")
    except Exception as e:
        st.error(f"向量诊断异常:{e}")

# ---------- 文件读取工具 ----------
def _lazy_docx():
    try:
        import docx  # python-docx
        return docx
    except Exception:
        return None

def _normalize(t: str) -> str:
    if not t: return ""
    t = t.replace("\xa0", " ").replace("\u200b", "")
    t = re.sub(r"[ \t]+", " ", t)
    t = re.sub(r"\n{2,}", "\n", t)
    return t.strip()

def read_docx_tables_info(file_like):
    docx = _lazy_docx()
    if not docx: return {}
    try:
        doc = docx.Document(file_like)
    except Exception:
        return {}
    info = {}
    for ti, tbl in enumerate(doc.tables):
        rows = len(tbl.rows)
        cols = len(tbl.columns) if rows else 0
        prev = []
        for r in tbl.rows[: min(6, rows)]:
            prev.append(tuple(_normalize(c.text) for c in r.cells))
        info[ti] = {"rows": rows, "cols": cols, "preview": prev}
    return info

def extract_pairs_from_docx_table(file_like, table_index=0, src_col=0, tgt_col=1,
                                  ffill=True, drop_empty_both=True, dedup=True):
    docx = _lazy_docx()
    if not docx: return []
    try:
        doc = docx.Document(file_like)
    except Exception:
        return []
    if table_index >= len(doc.tables): return []
    tbl = doc.tables[table_index]
    rows = []
    for r in tbl.rows:
        rows.append([_normalize(c.text) for c in r.cells])
    if not rows: return []
    max_cols = max(len(r) for r in rows)
    if src_col >= max_cols or tgt_col >= max_cols: return []

    if ffill:
        for col in (src_col, tgt_col):
            last = ""
            for i in range(len(rows)):
                val = rows[i][col] if col < len(rows[i]) else ""
                if val: last = val
                else: rows[i][col] = last

    pairs = []
    for r in rows:
        s = r[src_col] if src_col < len(r) else ""
        t = r[tgt_col] if tgt_col < len(r) else ""
        s, t = s.strip(), t.strip()
        if drop_empty_both and (not s and not t):
            continue
        pairs.append((s, t))

    if dedup:
        seen, out = set(), []
        for p in pairs:
            if p in seen: continue
            seen.add(p); out.append(p)
        pairs = out
    return pairs

def read_docx_text(file_like) -> str:
    docx = _lazy_docx()
    if not docx: return ""
    try:
        doc = docx.Document(file_like)

    except Exception:
        return ""
    blocks = []
    for p in doc.paragraphs:
        t = _normalize(p.text)
        if t: blocks.append(t)

    for tbl in doc.tables:
        for r in tbl.rows:
            line = " ".join(_normalize(c.text) for c in r.cells if _normalize(c.text))
            if line: blocks.append(line)
    return "\n".join(blocks)

def _guess_lang(text: str) -> str | None:
    if not text:
        return None
    zh = len(re.findall(r"[\u4e00-\u9fff]", text))
    en = len(re.findall(r"[A-Za-z]", text))
    if zh == 0 and en == 0:
        return None
    if zh >= en * 1.2:
        return "zh"
    if en >= zh * 1.2:
        return "en"
    return None

def _expected_first_from_lp(lp: str | None) -> str | None:
    if not lp:
        return None
    if lp.startswith("中"):
        return "zh"
    if lp.startswith("英"):
        return "en"
    return None

def detect_bilingual_pairs_from_lines(
    lines: list[str],
    expected_first: str | None = None,
    min_pairs: int = 3,
    min_ratio: float = 0.6,
) -> list[tuple[str, str]]:
    filtered = [(line, _guess_lang(line)) for line in lines]
    filtered = [(line, lang) for line, lang in filtered if lang in ("zh", "en")]
    if len(filtered) < min_pairs * 2:
        return []

    if expected_first is None:
        zh_cnt = sum(1 for _, lang in filtered if lang == "zh")
        en_cnt = sum(1 for _, lang in filtered if lang == "en")
        expected_first = "zh" if zh_cnt >= en_cnt else "en"

    def build_pairs(first_lang: str) -> list[tuple[str, str]]:
        pairs = []
        pending = None
        for line, lang in filtered:
            if lang == first_lang:
                pending = line
            elif pending and lang != first_lang:
                pairs.append((pending, line))
                pending = None
        return pairs

    pairs = build_pairs(expected_first)
    other = "en" if expected_first == "zh" else "zh"
    if not pairs:
        pairs_alt = build_pairs(other)
        if len(pairs_alt) > len(pairs):
            pairs = pairs_alt

    if len(pairs) < min_pairs:
        return []
    ratio = (len(pairs) * 2) / max(1, len(filtered))
    if ratio < min_ratio:
        return []
    return pairs

def extract_pairs_from_docx_paragraphs(file_like, expected_first: str | None = None):
    docx = _lazy_docx()
    if not docx:
        return []
    try:
        doc = docx.Document(file_like)
    except Exception:
        return []
    lines = []
    for p in doc.paragraphs:
        t = _normalize(p.text)
        if t:
            lines.append(t)
    return detect_bilingual_pairs_from_lines(lines, expected_first=expected_first)
   
def read_txt(file_like_or_bytes) -> str:
    try:
        if hasattr(file_like_or_bytes, "read"):
            data = file_like_or_bytes.read()
        else:
            data = file_like_or_bytes
        if isinstance(data, bytes):
            try: return data.decode("utf-8")
            except: return data.decode("utf-8", errors="ignore")
        return str(data)
    except Exception:
        return ""

def read_pdf_text(file_like) -> str:
    """
    Robust PDF text extraction with multiple fallbacks.
    1) PyPDF2
    2) pdfplumber (handles many scanned/column PDFs)
    3) pdfminer.six
    返回经过 _normalize 处理的文本串。若全部失败则返回空串。
    """
    # 将输入统一为 bytes，便于多次重读
    try:
        data = file_like.read() if hasattr(file_like, "read") else file_like
    except Exception:
        data = b""
    if isinstance(data, str):
        data = data.encode(errors="ignore")
    bio = io.BytesIO(data or b"")

    # 1) PyPDF2
    try:
        from PyPDF2 import PdfReader
        reader = PdfReader(io.BytesIO(bio.getvalue()))
        txts = []
        for page in reader.pages:
            t = page.extract_text() or ""
            t = _normalize(t)
            if t:
                txts.append(t)
        if txts:
            return "\n".join(txts)
    except Exception:
        pass

    # 2) pdfplumber
    try:
        import pdfplumber
        txts = []
        with pdfplumber.open(io.BytesIO(bio.getvalue())) as pdf:
            for p in pdf.pages:
                t = p.extract_text() or ""
                t = _normalize(t)
                if t:
                    txts.append(t)
        if txts:
            return "\n".join(txts)
    except Exception:
        pass

    # 3) pdfminer.six
    try:
        from pdfminer.high_level import extract_text
        t = extract_text(io.BytesIO(bio.getvalue())) or ""
        t = _normalize(t)
        return t
    except Exception:
        return ""
def _lazy_import_vec():
    try:
        import faiss
    except Exception:
        faiss = None

    try:
        from sentence_transformers import SentenceTransformer
    except Exception:
        SentenceTransformer = None

    try:
        from fastembed import TextEmbedding as FastEmbedModel 
    except Exception:
        FastEmbedModel = None

    try:
        from sklearn.feature_extraction.text import TfidfVectorizer
    except Exception:
        TfidfVectorizer = None

    return np, faiss, SentenceTransformer, FastEmbedModel, TfidfVectorizer, None

def _make_sbert_encoder(
    model_name: str,
    *,
    normalize_embeddings: bool,
    extra_normalize: bool,
    batch_size: int | None = None,
    handle_empty: bool = False,
):
    from sentence_transformers import SentenceTransformer
    model = SentenceTransformer(model_name)

    def encode(texts: list[str]):
        if handle_empty and not texts:
            dim = model.get_sentence_embedding_dimension()
            return np.zeros((0, dim), dtype="float32")

        eff_batch = batch_size or 32  # SentenceTransformer 不接受 None，避免 step=None 报错
        emb = model.encode(
            texts,
            normalize_embeddings=normalize_embeddings,
            batch_size=eff_batch,
            convert_to_numpy=True,
        ).astype("float32")
        if extra_normalize:
            norms = np.linalg.norm(emb, axis=1, keepdims=True) + 1e-12
            emb = (emb / norms).astype("float32")
        return emb

    return encode

# ========== 向量召回(多后端:Sentence-Transformers → Fastembed → TF-IDF)==========
@st.cache_resource(show_spinner=False)
def get_embedder():
    """
    优先 sentence-transformers；若导入/加载失败，自动回退 fastembed，再回退 TF-IDF。
    避免 torch/transformers 兼容问题导致直接报错。
    """
    # 1) sentence-transformers
    try:
        encode_st = _make_sbert_encoder(
            "distiluse-base-multilingual-cased-v1",
            normalize_embeddings=True,
            extra_normalize=True,
            batch_size=32,
            handle_empty=True,
        )
        return "st", encode_st
    except Exception as e:
        log_event("WARNING", "sentence-transformers unavailable, fallback to fastembed/tfidf", err=str(e))

    # 2) fastembed
    try:
        from fastembed import TextEmbedding
        import numpy as _np

        fe_model = TextEmbedding("BAAI/bge-small-en-v1.5")

        def encode_fe(texts: list[str]):
            if not texts:
                return _np.zeros((0, fe_model.dim), dtype="float32")
            vecs = list(fe_model.embed(texts))
            arr = _np.asarray(vecs, dtype="float32")
            arr = arr / (_np.linalg.norm(arr, axis=1, keepdims=True) + 1e-12)
            return arr

        return "fastembed", encode_fe
    except Exception as e:
        log_event("WARNING", "fastembed unavailable, fallback to tfidf", err=str(e))

    # 3) TF-IDF 回退
    try:
        from sklearn.feature_extraction.text import TfidfVectorizer
        import numpy as _np
    except Exception as e:
        st.error(f"向量编码后端不可用: {e}")
        raise RuntimeError("no available embedder") from e

    tf_vec = TfidfVectorizer(min_df=1)

    def encode_tf(texts: list[str]):
        if not texts:
            return _np.zeros((0, 1), dtype="float32")
        X = tf_vec.fit_transform(texts)
        arr = X.toarray().astype("float32")
        arr = arr / (_np.linalg.norm(arr, axis=1, keepdims=True) + 1e-12)
        return arr

    return "tfidf", encode_tf

def _load_index_by_paths(idx_path: str, map_path: str, vec_path: str, *, project_id=None, max_vec_bytes: int = 1_000_000_000):
    """
    统一的索引加载函数，增加 .npy 大小检查，避免直接 np.load 超大文件导致 OOM。
    """
    mapping: list = []
    try:
        if os.path.exists(map_path):
            with open(map_path, "r", encoding="utf-8") as f:
                mapping = json.load(f) or []
    except Exception as e:
        log_exception("load_index_by_paths: failed to load mapping", e, level="WARNING")
        mapping = []

    # 优先读取 FAISS 索引
    try:
        _, faiss, *_ = _lazy_import_vec()
    except Exception:
        faiss = None
    if faiss is not None and os.path.exists(idx_path):
        try:
            index = faiss.read_index(idx_path)
            return "faiss", index, mapping, None
        except Exception as e:
            log_exception("load_index_by_paths: failed to load faiss index", e, level="WARNING")

    # 回退读取 numpy 向量，先检查大小
    if os.path.exists(vec_path):
        try:
            vec_size = os.path.getsize(vec_path)
        except Exception:
            vec_size = None
        if vec_size is not None and max_vec_bytes and vec_size > max_vec_bytes:
            log_event(
                "WARNING",
                "skip loading vector npy to avoid OOM",
                bytes=vec_size,
                limit=max_vec_bytes,
                project_id=project_id,
                path=vec_path,
            )
            return "none", None, [], None
        try:
            import numpy as np
            vecs = np.load(vec_path)
            return "fallback", None, mapping, vecs
        except Exception as e:
            log_exception("load_index_by_paths: failed to load fallback vectors", e, level="WARNING")

    return "none", None, mapping or [], None

def _save_index_by_paths(idx_path: str,
                         map_path: str,
                         vec_path: str,
                         mode: str,
                         index,
                         mapping,
                         *,
                         vecs=None,
                         cleanup: bool = True):
    os.makedirs(os.path.dirname(idx_path), exist_ok=True)

    try:
        with open(map_path, "w", encoding="utf-8") as f:
            json.dump(mapping or [], f, ensure_ascii=False)
    except Exception as e:
        log_exception("save_index_by_paths: failed to persist mapping", e, level="WARNING")

    try:
        _, faiss, *_ = _lazy_import_vec()
    except Exception:
        faiss = None

    if mode == "faiss":
        try:
            if faiss is not None and index is not None:
                faiss.write_index(index, idx_path)
        except Exception as e:
            log_exception("save_index_by_paths: failed to save faiss index", e, level="WARNING")
        if cleanup:
            try:
                if os.path.exists(vec_path):
                    os.remove(vec_path)
            except Exception:
                pass
    elif mode == "fallback":
        try:
            if vecs is not None:
                import numpy as np
                np.save(vec_path, vecs)
        except Exception as e:
            log_exception("save_index_by_paths: failed to save fallback vectors", e, level="WARNING")
        if cleanup:
            try:
                if os.path.exists(idx_path):
                    os.remove(idx_path)
            except Exception:
                pass
    elif mode == "none" and cleanup:
        for p in (idx_path, vec_path):
            try:
                if os.path.exists(p):
                    os.remove(p)
            except Exception:
                pass

def _load_index(project_id: int):
    idx_path, map_path, vec_path = _index_paths(project_id)
    return _load_index_by_paths(idx_path, map_path, vec_path, project_id=project_id)

def _save_index(project_id: int, mode: str, index, mapping, vecs=None):
    idx_path, map_path, vec_path = _index_paths(project_id)
    _save_index_by_paths(
        idx_path,
        map_path,
        vec_path,
        mode,
        index,
        mapping,
        vecs=vecs,
        cleanup=False,
    )

def _index_paths_domain(domain: str, kb_type: str):
    domain_key = _norm_domain_key(domain)
    return _index_paths_common(domain_key, kb_type)

def _load_index_domain(domain: str, kb_type: str):
    idx_path, map_path, vec_path = _index_paths_domain(domain, kb_type)
    return _load_index_by_paths(idx_path, map_path, vec_path)

def _save_index_domain(domain: str, kb_type: str, mode: str, index, mapping, vecs=None):
    idx_path, map_path, vec_path = _index_paths_domain(domain, kb_type)
    _save_index_by_paths(
        idx_path,
        map_path,
        vec_path,
        mode,
        index,
        mapping,
        vecs=vecs,
        cleanup=True,
    )

def _build_vector_index_from_texts(texts, metas, backend, encode, faiss, save_fn):
    new_vecs = encode(texts)
    if hasattr(new_vecs, "toarray"):
        new_vecs = new_vecs.toarray()
    new_vecs = _np.asarray(new_vecs, dtype="float32")
    if new_vecs.ndim == 1:
        new_vecs = new_vecs.reshape(1, -1)
    new_vecs = new_vecs / (_np.linalg.norm(new_vecs, axis=1, keepdims=True) + 1e-12)

    if faiss is not None and backend in ("st", "fastembed"):
        dim = int(new_vecs.shape[1])
        index = faiss.IndexFlatIP(dim)
        index.add(new_vecs)
        save_fn("faiss", index, metas, None)
        return int(index.ntotal)
    vecs = new_vecs
    save_fn("fallback", None, metas, vecs)
    return int(vecs.shape[0])

def build_strategy_index_for_domain(domain: str):
    np, faiss, *_ = _lazy_import_vec()
    backend, encode = get_embedder()

    dom = (domain or "").strip() or "未分类"

    # 1) 确保策略表存在
    execute_write(
        "CREATE TABLE IF NOT EXISTS strategy_texts ("
        "id INTEGER PRIMARY KEY,"
        "domain TEXT,"
        "title TEXT,"
        "content TEXT NOT NULL,"
        "collection TEXT,"
        "source TEXT,"
        "created_at TEXT DEFAULT (datetime('now'))"
        ");"
    )

    # 2) 拉取该领域下的全部策略文本
    rows = execute_write(
        """
        SELECT id,
               IFNULL(domain,''), IFNULL(title,''), content,
               IFNULL(collection,''), IFNULL(source,'')
          FROM strategy_texts
         WHERE IFNULL(domain,'') = ?
         ORDER BY id ASC
        """,
        (dom,)
    ).fetchall()

    texts, metas = [], []
    for sid, d, ttl, content, coll, src in rows:
        txt = (content or "").strip()
        if not txt:
            continue
        texts.append(txt)
        metas.append({
            "strategy_id": sid,
            "domain": d or dom,
            "title": ttl,
            "content_preview": txt[:200],
            "collection": coll,
            "source": src,
            "kb_type": "strategy",
        })

    if not texts:
        _save_index_domain(dom, "strategy", "none", None, [])
        return {"added": 0, "total": 0}

    total = _build_vector_index_from_texts(
        texts,
        metas,
        backend,
        encode,
        faiss,
        lambda mode, index, mapping, vecs: _save_index_domain(
            dom,
            "strategy",
            mode,
            index,
            mapping,
            vecs=vecs,
        ),
    )
    return {"added": len(texts), "total": total}

def build_project_vector_index(project_id: int,
                               use_src: bool = True,
                               use_tgt: bool = True):
    import numpy as np

    _, faiss, *_ = _lazy_import_vec()
    backend, encode = get_embedder()

    pid = int(project_id)

    proj_domain = None
    try:
        row = execute_write(
            "SELECT IFNULL(domain,'') FROM items WHERE id=?",
            (pid,)
        ).fetchone()
        if row:
            proj_domain = (row[0] or "").strip()
    except Exception:
        proj_domain = None

    if not proj_domain:
        proj_domain = "未分类"

    # 1) 读取已存在的索引，准备增量去重
    existing_mode, existing_index, existing_mapping_raw, existing_vecs = _load_index(pid)
    existing_mapping = existing_mapping_raw if isinstance(existing_mapping_raw, list) else []
    existing_total = len(existing_mapping)
    can_reuse_existing = (
        (existing_mode == "faiss" and existing_index is not None) or
        (existing_mode == "fallback" and existing_vecs is not None and hasattr(existing_vecs, "shape") and existing_vecs.shape[0] == existing_total)
    )
    if not can_reuse_existing:
        existing_mapping = []
        existing_total = 0
    existing_keys = {
        (int(m.get("corpus_id") or -1), int(m.get("idx") or -1))
        for m in existing_mapping
        if isinstance(m, dict)
    }

    # 2) 从 DB 读取该领域的语料(不再按 project_id 限制)
    rows = execute_write(
        """
        SELECT c.id,
               IFNULL(c.src_text, ''), IFNULL(c.tgt_text, ''),
               IFNULL(c.title, ''),    IFNULL(c.lang_pair, ''),
               IFNULL(c.project_id, 0), IFNULL(c.domain, '')
        FROM corpus c
        WHERE IFNULL(c.domain, '') = ?
        ORDER BY c.id ASC
        """,
        (proj_domain,)
    ).fetchall()

    texts, metas = [], []
    has_usable_corpus = False

    for cid, s, t, ttl, lp, pj, dom in rows:
        s = (s or "").strip()
        t = (t or "").strip()
        if not s and not t:
            continue
        has_usable_corpus = True

        # 句子切分
        try:
            if "split_sents" in globals():
                src_sents = split_sents(s, lang_hint="zh")
                tgt_sents = split_sents(t, lang_hint="en")
            else:
                src_sents = (s.split("。") if s else [])
                tgt_sents = (t.split(".") if t else [])
        except Exception:
            src_sents = (s.split("。") if s else [])
            tgt_sents = (t.split(".") if t else [])

        n = min(len(src_sents), len(tgt_sents)) if (use_src and use_tgt) else len(src_sents or [])

        for idx in range(n):
            src_j = (src_sents[idx] if idx < len(src_sents) else "").strip()
            tgt_j = (tgt_sents[idx] if idx < len(tgt_sents) else "").strip()
            if not src_j:
                continue
            key = (int(cid), int(idx))
            if key in existing_keys:
                continue

            texts.append(src_j)
            metas.append({
                "corpus_id": cid,
                "idx": idx,
                "src": src_j,
                "tgt": tgt_j,
                "project_id": pj,
                "domain": dom or proj_domain or "",
                "title": ttl,
                "lang_pair": lp or "",
                "kb_type": "bilingual",
            })

    if not texts:
        if not has_usable_corpus and not existing_mapping:
            try:
                _save_index(pid, "none", None, [], vecs=None)
            except Exception as e:
                log_exception(f"Failed to persist empty semantic index placeholder for pid={pid}", e, level="WARNING")
            return {"added": 0, "total": 0}
        return {"added": 0, "total": existing_total}

    def _encode_for_index(raw_texts: list[str]):
        vecs_new = encode(raw_texts)
        if hasattr(vecs_new, "toarray"):
            vecs_new = vecs_new.toarray()
        vecs_new = np.asarray(vecs_new, dtype="float32")
        if vecs_new.ndim == 1:
            vecs_new = vecs_new.reshape(1, -1)
        vecs_new = vecs_new / (np.linalg.norm(vecs_new, axis=1, keepdims=True) + 1e-12)
        return vecs_new

    new_vecs = _encode_for_index(texts)
    mapping_all = existing_mapping + metas

    # 优先使用 FAISS 增量追加
    if faiss is not None and backend in ("st", "fastembed"):
        dim_new = int(new_vecs.shape[1])
        if existing_mode == "faiss" and existing_index is not None and getattr(existing_index, "d", dim_new) == dim_new:
            existing_index.add(new_vecs)
            _save_index(pid, "faiss", existing_index, mapping_all, vecs=None)
            return {"added": len(texts), "total": int(existing_index.ntotal)}
        # 回退到重建（带上已有向量）
        base_vecs = None
        has_full_base = False
        if existing_vecs is not None and hasattr(existing_vecs, "shape") and existing_vecs.shape[0] == existing_total:
            base_vecs = np.asarray(existing_vecs, dtype="float32")
            has_full_base = True
        elif existing_mode == "faiss" and existing_index is not None:
            try:
                base_vecs = existing_index.reconstruct_n(0, existing_index.ntotal)
                has_full_base = base_vecs is not None and base_vecs.shape[0] == existing_total
            except Exception:
                base_vecs = None
                has_full_base = False
        if not has_full_base:
            base_vecs = None
        mapping_to_save = mapping_all if has_full_base else metas
        all_vecs = new_vecs if base_vecs is None else np.vstack([base_vecs, new_vecs])
        index = faiss.IndexFlatIP(int(all_vecs.shape[1]))
        index.add(all_vecs)
        _save_index(pid, "faiss", index, mapping_to_save, vecs=None)
        return {"added": len(texts), "total": int(index.ntotal)}

    # 无 FAISS 或其他后端时，使用 numpy 数组存储
    base_vecs = existing_vecs if (existing_mode == "fallback" and existing_vecs is not None) else None
    has_full_base = bool(base_vecs is not None and hasattr(base_vecs, "shape") and base_vecs.shape[0] == existing_total)
    if not has_full_base:
        base_vecs = None
    mapping_to_save = mapping_all if has_full_base else metas
    vecs_final = new_vecs if base_vecs is None else np.vstack([base_vecs, new_vecs])
    _save_index(pid, "fallback", None, mapping_to_save, vecs=vecs_final)
    return {"added": len(texts), "total": int(vecs_final.shape[0])}

def rebuild_project_semantic_index(project_id: int) -> dict:
    try:
        pid = int(project_id)
    except (TypeError, ValueError):
        return {"ok": False, "added": 0, "total": 0, "msg": f"非法项目ID: {project_id!r}"}

    try:
        res = build_project_vector_index(pid, use_src=True, use_tgt=True)
        return {
            "ok": True,
            "added": int(res.get("added", 0)),
            "total": int(res.get("total", 0)),
            "msg": "索引重建成功",
        }
    except Exception as e:
        return {
            "ok": False,
            "added": 0,
            "total": 0,
            "msg": f"索引重建失败: {e}",
        }

# =========================
# 语义召回(支持范围:project/domain/all)
# 返回: [(score, meta, src_sent, tgt_sent)]
# =========================
def _get_domain_for_proj(cur, project_id: int) -> str | None:
    try:
        row = execute_write(
            "SELECT IFNULL(domain,'') FROM items WHERE id=?",
            (int(project_id),),
        ).fetchone()
    except Exception:
        return None

    if not row:
        return None

    dom = (row[0] or "").strip()
    return dom or None

def semantic_retrieve(project_id: int,
                      query_text: str,
                      topk: int = 20,
                      scope: str = "project",
                      min_char: int = 3):

    q = (query_text or "").strip()
    if len(q) < min_char:
        return []

    def _split(text: str) -> list[str]:
        try:
            if "split_sents" in globals():
                segs = split_sents(text, lang_hint="auto")  # type: ignore
                return [s for s in segs if s and len(s.strip()) >= min_char]
        except Exception as e:
            log_exception("split_sents failed, falling back to regex splitting", e, level="WARNING")
        import re
        segs = re.split(r"(?<=[\.\!\?;。！？；])\s*", text)
        return [s.strip() for s in segs if s and len(s.strip()) >= min_char]

    pieces = _split(q)
    if not pieces:
        return []

    # --- 取向量编码器 ---
    try:
        backend, encode = get_embedder()
    except RuntimeError as e:
        st.error(f"向量模型加载失败: {e}")
        return

    # --- 取索引 & mapping ---
    mode, index, mapping, vecs = _load_index(int(project_id))  # 已在文件里定义
    if mode == "none" or not mapping:
        return []

    # 懒加载 numpy / faiss
    np_mod, faiss, *_ = _lazy_import_vec()  # 已有的工具函数
    np = np_mod  

    cur_domain = None
    if scope == "domain":
        try:
            cur_domain = _get_domain_for_proj(cur, int(project_id))  # 文件里已有
        except Exception as e:
            log_exception(f"semantic_retrieve: failed to get domain for project {project_id}", e, level="WARNING")
            cur_domain = None

    def _scope_ok(meta: dict) -> bool:
        if scope == "project":
            return int(meta.get("project_id", 0) or 0) == int(project_id)
        elif scope == "domain" and cur_domain:
            return (meta.get("domain") or "") == (cur_domain or "")
        else:
            # "all" 或拿不到 domain 时，都不过滤
            return True

    all_hits: list[tuple[float, dict, str, str]] = []

    per_piece_k = max(topk * 3, topk)

    for piece in pieces:
        if not piece:
            continue

        try:
            # 1) 生成查询向量 qv
            qv = encode([piece])
            if hasattr(qv, "toarray"):  # tf-idf 稀疏矩阵
                qv = qv.toarray()
            qv = np.asarray(qv, dtype="float32")
            if qv.ndim == 2:
                qv = qv[0]

            # 查询向量归一化
            q_norm = np.linalg.norm(qv) + 1e-12
            qv = qv / q_norm

            # 2) FAISS 分支
            if mode == "faiss" and index is not None and faiss is not None:
                k = min(per_piece_k, len(mapping))
                if k <= 0:
                    continue
                D, I = index.search(qv.reshape(1, -1), k)
                for score, idx in zip(D[0].tolist(), I[0].tolist()):
                    idx = int(idx)
                    if idx < 0 or idx >= len(mapping):
                        continue
                    meta = mapping[idx] or {}
                    if not isinstance(meta, dict):
                        continue
                    if not _scope_ok(meta):
                        continue
                    src_sent = (meta.get("src") or "").strip()
                    tgt_sent = (meta.get("tgt") or "").strip()
                    if not src_sent and not tgt_sent:
                        continue
                    all_hits.append((float(score), meta, src_sent, tgt_sent))

            # 3) fallback 分支：纯 numpy 相似度
            elif mode == "fallback" and vecs is not None:
                arr = np.asarray(vecs, dtype="float32")
                if arr.ndim != 2 or arr.shape[0] == 0:
                    continue
                sims = arr @ qv.reshape(-1, 1)   
                sims = sims.reshape(-1)
                k = min(per_piece_k, sims.shape[0])
                if k <= 0:
                    continue
                idxs = np.argsort(-sims)[:k]
                for idx in idxs:
                    idx = int(idx)
                    score = float(sims[idx])
                    if idx < 0 or idx >= len(mapping):
                        continue
                    meta = mapping[idx] or {}
                    if not isinstance(meta, dict):
                        continue
                    if not _scope_ok(meta):
                        continue
                    src_sent = (meta.get("src") or "").strip()
                    tgt_sent = (meta.get("tgt") or "").strip()
                    if not src_sent and not tgt_sent:
                        continue
                    all_hits.append((score, meta, src_sent, tgt_sent))

            else:
                # 没有可用索引
                continue

        except Exception as e:
            log_exception("semantic_retrieve: error in vector similarity computation", e, level="DEBUG")
            continue

    if not all_hits:
        return []

    # --- 去重 + 按得分排序，保留前 topk ---
    dedup = {}
    for score, meta, src_sent, tgt_sent in all_hits:
        key = (src_sent, tgt_sent)
        if key not in dedup or score > dedup[key][0]:
            dedup[key] = (score, meta, src_sent, tgt_sent)

    hits = sorted(dedup.values(), key=lambda x: x[0], reverse=True)
    return hits[:topk]

def semantic_consistency_report(project_id: int,
                                blocks_src: list,
                                blocks_tgt: list,
                                term_map: dict,
                                topk: int = 3,
                                thr: float = 0.70):
    hits_all = []

    # 对齐长度，避免两侧长度不一致出错
    n = min(len(blocks_src or []), len(blocks_tgt or []))
    if n == 0:
        return pd.DataFrame([])

    for i, (s, t) in enumerate(zip(blocks_src[:n], blocks_tgt[:n]), 1):
        s = s or ""
        t = t or ""

        # 1) 用“译文”去检索参考译文
        try:
            hits = semantic_retrieve(project_id, t, topk=topk)
        except Exception as e:
            log_exception(f"semantic_consistency_report: failed semantic retrieve for segment {i}", e, level="WARNING")
            hits = []

        # semantic_retrieve 统一返回: (score, meta, src_sent, tgt_sent)
        if hits:
            top_score = float(hits[0][0])
        else:
            top_score = 0.0

        # 2) 术语遵守：源段包含源术语，但译文里没出现目标术语
        violated = []
        for src_term, tgt_term in (term_map or {}).items():
            if not src_term or not tgt_term:
                continue
            if src_term in s and tgt_term not in t:
                violated.append(f"{src_term}->{tgt_term}")

        hits_all.append(
            {
                "段号": i,
                "相似参考得分": round(top_score, 2),
                "低于阈值": (top_score < thr),
                "未遵守术语": ", ".join(violated) if violated else "",
            }
        )

    return pd.DataFrame(hits_all)

# ========== 路径/DB ==========
db_lock = threading.RLock()
# Streamlit session 键名，与 SessionManager.DB_CONN 保持一致
DB_CONN = SessionManager.DB_CONN
# 语料引用映射与开关
CORPUS_REFS = SessionManager.CORPUS_REFS
COR_USE_REF = SessionManager.COR_USE_REF
# 知识库向量化器 session 键
KB_EMBEDDER = SessionManager.KB_EMBEDDER
# 新建项目作用域 session 键
SCOPE_NEWPROJ = SessionManager.SCOPE_NEWPROJ
# 全局语义开关 session 键
USE_SEMANTIC_GLOBAL = SessionManager.USE_SEMANTIC_GLOBAL
# 翻译管理页待切换的 tab（外部触发）键
TRANS_MGMT_TAB_PENDING = SessionManager.TRANS_MGMT_TAB_PENDING
# 翻译管理页当前 tab 键
TRANS_MGMT_TAB = SessionManager.TRANS_MGMT_TAB
# 当前翻译项目 id / 源文件路径
ACTIVE_TRANSLATION_PID = SessionManager.ACTIVE_TRANSLATION_PID
ACTIVE_TRANSLATION_SRC = SessionManager.ACTIVE_TRANSLATION_SRC
# 语料 Few-shot 目标项目信息
CORPUS_TARGET_PROJECT = SessionManager.CORPUS_TARGET_PROJECT
CORPUS_TARGET_LABEL = SessionManager.CORPUS_TARGET_LABEL

def _is_write_sql(sql: str) -> bool:
    if not sql:
        return False
    s = sql.lstrip().lower()
    while s.startswith("("):
        s = s[1:].lstrip()
    if not s:
        return False
    keyword = s.split(None, 1)[0]
    return keyword not in ("select", "pragma")

class LockedConnection:
    def __init__(self, conn, lock):
        self._conn = conn
        self._lock = lock
        self._write_lock_held = False

    def cursor(self):
        return LockedCursor(self._conn.cursor(), self)

    def commit(self):
        if self._write_lock_held:
            try:
                return self._conn.commit()
            finally:
                self._write_lock_held = False
                self._lock.release()
        with self._lock:
            return self._conn.commit()

    def execute(self, *args, **kwargs):
        with self._lock:
            return self._conn.execute(*args, **kwargs)

    def executemany(self, *args, **kwargs):
        with self._lock:
            return self._conn.executemany(*args, **kwargs)

    def __getattr__(self, name):
        return getattr(self._conn, name)

class LockedCursor:
    def __init__(self, cursor, conn):
        self._cursor = cursor
        self._conn = conn

    def execute(self, sql, params=()):
        if _is_write_sql(sql):
            self._conn._lock.acquire()
            self._conn._write_lock_held = True
            return self._cursor.execute(sql, params)
        with self._conn._lock:
            return self._cursor.execute(sql, params)

    def executemany(self, sql, seq):
        if _is_write_sql(sql):
            self._conn._lock.acquire()
            self._conn._write_lock_held = True
            return self._cursor.executemany(sql, seq)
        with self._conn._lock:
            return self._cursor.executemany(sql, seq)

    def __getattr__(self, name):
        return getattr(self._cursor, name)

def get_db():
    """
    获取数据库连接（每个 Streamlit session 一个连接）。
    使用 LockedConnection 确保线程安全。
    """
    if DB_CONN not in st.session_state:
        base_conn = sqlite3.connect(DB_PATH, check_same_thread=False)
        st.session_state[DB_CONN] = LockedConnection(base_conn, db_lock)
    return st.session_state[DB_CONN]

# 模块级连接：Streamlit 每个 session 独立，线程安全由 LockedConnection 保证
# 注意：不要在函数外直接使用 conn/cur，应通过 get_db() 获取
conn = get_db()
cur = conn.cursor()
RAW_CONN = conn._conn if hasattr(conn, "_conn") else conn

class _NullCursor:
    def fetchall(self):
        return []
    def fetchone(self):
        return None
    @property
    def lastrowid(self):
        return None

def execute_write(sql, params=()):
    try:
        with db_lock:
            if not _is_write_sql(sql):
                return cur.execute(sql, params)
            if isinstance(params, (list, tuple)) and params and isinstance(params[0], (list, tuple)):
                cur.executemany(sql, params)
            else:
                cur.execute(sql, params)
            conn.commit()
        return cur
    except Exception as e:
        log_event("ERROR", "db write failed", sql=sql, err=str(e))
        return None

def _get_scene_for_proj(cur, project_id: int | None):
    try:
        if not project_id:
            return None
        row = execute_write(
            "SELECT scene FROM items WHERE id=?",
            (int(project_id),)
        ).fetchone()
        if not row:
            return None
        val = (row[0] or "").strip()
        return val or None
    except Exception as e:
        log_exception(f"_get_domain_for_proj failed for project_id={project_id}", e, level="WARNING")
        return None

def ensure_poly_tables(conn, cur):
    def _table_exists(name: str) -> bool:
        row = cur.execute(
            "SELECT name FROM sqlite_master WHERE type='table' AND name=?",
            (name,),
        ).fetchone()
        return bool(row)

    def _create_new_tables():
        execute_write(
            "CREATE TABLE IF NOT EXISTS poly_term ("
            "id INTEGER PRIMARY KEY,"
            "source_term TEXT NOT NULL,"
            "created_at TEXT DEFAULT (datetime('now'))"
            ");"
        )
        execute_write(
            "CREATE TABLE IF NOT EXISTS poly_sense ("
            "id INTEGER PRIMARY KEY,"
            "term_id INTEGER NOT NULL,"
            "target_term TEXT NOT NULL,"
            "domain TEXT,"
            "example_src TEXT,"
            "example_tgt TEXT,"
            "created_at TEXT DEFAULT (datetime('now')),"
            "FOREIGN KEY(term_id) REFERENCES poly_term(id)"
            ");"
        )
        execute_write("CREATE INDEX IF NOT EXISTS idx_poly_term_source ON poly_term(source_term)")
        execute_write("CREATE INDEX IF NOT EXISTS idx_poly_sense_term ON poly_sense(term_id)")

    if not _table_exists("poly_term"):
        _create_new_tables()
        return

    cols = [r[1] for r in cur.execute("PRAGMA table_info(poly_term)").fetchall()]
    if "source_term" in cols:
        return

    try:
        if not _table_exists("poly_term_legacy"):
            execute_write("ALTER TABLE poly_term RENAME TO poly_term_legacy")
        if _table_exists("poly_sense") and not _table_exists("poly_sense_legacy"):
            execute_write("ALTER TABLE poly_sense RENAME TO poly_sense_legacy")
        _create_new_tables()

        execute_write(
            "INSERT INTO poly_term (source_term, created_at) "
            "SELECT DISTINCT term, datetime('now') FROM poly_term_legacy"
        )
        if _table_exists("poly_sense_legacy"):
            execute_write(
                "INSERT INTO poly_sense (term_id, target_term, domain, example_src, example_tgt) "
                "SELECT t.id, s.keywords_en, COALESCE(s.domain, lt.domain), s.example_src, s.example_tgt "
                "FROM poly_sense_legacy s "
                "JOIN poly_term t ON t.source_term = s.term "
                "LEFT JOIN poly_term_legacy lt ON lt.term = s.term "
                "WHERE IFNULL(s.keywords_en,'') <> ''"
            )
    except Exception as e:
        st.warning(f"一词多义表迁移失败: {e}")

def _split_sentences_zh(text: str):
    return [s.strip() for s in re.split(r"(?<=[。！？；])", text or "") if s.strip()]

def _split_sentences_en(text: str):
    return [s.strip() for s in re.split(r"(?<=[\.\!\?;:])\s+", text or "") if s.strip()]

def _pick_sentence(text: str, token: str, lang: str):
    if not text or not token:
        return ""
    sents = _split_sentences_zh(text) if lang == "zh" else _split_sentences_en(text)
    for s in sents:
        if token in s:
            return s
    return sents[0] if sents else ""

def add_polysemy_manual(cur, conn, term, en_terms, domain=None, example_src="", example_tgt=""):
    source_term = (term or "").strip()
    if not source_term or not en_terms:
        return 0

    cleaned = []
    seen = set()
    for w in en_terms:
        w = (w or "").strip()
        if not w:
            continue
        wl = w.lower()
        if wl in seen:
            continue
        seen.add(wl)
        cleaned.append(w)
    if not cleaned:
        return 0

    row = cur.execute(
        "SELECT id FROM poly_term WHERE source_term=?",
        (source_term,),
    ).fetchone()
    if not row:
        execute_write(
            "INSERT INTO poly_term(source_term) VALUES (?)",
            (source_term,),
        )
        row = cur.execute(
            "SELECT id FROM poly_term WHERE source_term=?",
            (source_term,),
        ).fetchone()
    if not row:
        return 0
    term_id = row[0]

    existing = cur.execute(
        "SELECT target_term FROM poly_sense WHERE term_id=?",
        (term_id,),
    ).fetchall()
    existing_en = {((r[0] or "").strip().lower()) for r in existing if r and r[0]}
    to_add = [w for w in cleaned if w.lower() not in existing_en]
    if not to_add:
        return 0

    for w in to_add:
        execute_write(
            "INSERT INTO poly_sense(term_id, target_term, domain, example_src, example_tgt) "
            "VALUES(?, ?, ?, ?, ?)",
            (
                int(term_id),
                w,
                domain or None,
                example_src or "",
                example_tgt or "",
            ),
        )

    return len(to_add)

def build_poly_hint(cur, project_id, blk, domain, scene, lang_pair, topk: int = 8):
    try:
        rows = cur.execute(
            "SELECT t.id, t.source_term "
            "FROM poly_term t "
            "WHERE (SELECT COUNT(1) FROM poly_sense s WHERE s.term_id=t.id) >= 2 "
            "ORDER BY t.id DESC LIMIT 30"
        ).fetchall()
    except Exception:
        return ""
    if not rows:
        return ""

    hit_terms = [(tid, t) for (tid, t) in rows if t and t in (blk or "")]
    if not hit_terms:
        return ""

    lines = []
    for term_id, term in hit_terms:
        try:
            senses = cur.execute(
                "SELECT target_term, IFNULL(example_src,''), IFNULL(example_tgt,'') "
                "FROM poly_sense WHERE term_id=? ORDER BY id",
                (term_id,),
            ).fetchall()
        except Exception:
            senses = []
        if not senses:
            continue

        parts = []
        for tgt, ex_src, ex_tgt in senses[:3]:
            ex_src = (ex_src or "")[:60]
            ex_tgt = (ex_tgt or "")[:80]
            if ex_src or ex_tgt:
                parts.append(f"{tgt} 例: {ex_src} / {ex_tgt}")
            else:
                parts.append(tgt)
        if parts:
            lines.append(f"- 词={term}: " + "; ".join(parts))

    if not lines:
        return ""

    header = "【一词多义处理提示】"
    hint = header + "\n" + "\n".join(lines)
    if len(hint) > 400:
        hint = hint[:400]
    return hint

def write_translation_segments_to_corpus(
    cur,
    conn,
    project_id,
    title,
    lang_pair,
    domain,
    scene,
    blocks_src,
    blocks_tgt,
):
    if not blocks_src or not blocks_tgt:
        return 0
    n = min(len(blocks_src), len(blocks_tgt))
    inserted = 0
    for i in range(n):
        s = (blocks_src[i] or "").strip()
        t = (blocks_tgt[i] or "").strip()
        if not s and not t:
            continue
        try:
            exists = execute_write(
                "SELECT 1 FROM corpus WHERE project_id=? AND IFNULL(src_text,'')=? AND IFNULL(tgt_text,'')=? LIMIT 1",
                (int(project_id), s, t),
            ).fetchone()
            if exists:
                continue
        except Exception as e:
            log_exception(f"Failed to check duplicate corpus row for project_id={project_id}", e, level="WARNING")

        cols = ["title","project_id","lang_pair","src_text","tgt_text","note","created_at"]
        vals = [title, int(project_id), lang_pair, s, t, "from translation_history", None]
        if _has_col("corpus", "domain"):
            cols.append("domain")
            vals.append(domain or "")
        if _has_col("corpus", "scene"):
            cols.append("scene")
            vals.append(scene or "")
        if _has_col("corpus", "unit"):
            cols.append("unit")
            vals.append("paragraph")
        if _has_col("corpus", "source_type"):
            cols.append("source_type")
            vals.append("translation_history")

        placeholders = ",".join(["?"] * len(cols))
        sql = f"INSERT INTO corpus ({', '.join(cols)}) VALUES ({placeholders})"
        try:
            execute_write(sql, vals)
            inserted += 1
        except Exception:
            continue
    return inserted

def ensure_domain_columns_and_backfill(conn, cur, corpus_table="corpus"):
    # items.domain
    try:
        cols = [r[1] for r in cur.execute("PRAGMA table_info(items)").fetchall()]
        if "domain" not in cols:
            execute_write("ALTER TABLE items ADD COLUMN domain TEXT;")
    except Exception as e:
        log_exception("ensure_domain_columns_and_backfill: add domain column to items failed", e, level="WARNING")

    # corpus.domain
    try:
        cols = [r[1] for r in cur.execute(f"PRAGMA table_info({corpus_table})").fetchall()]
        if "domain" not in cols:
            execute_write(f"ALTER TABLE {corpus_table} ADD COLUMN domain TEXT;")
    except Exception as e:
        log_exception(f"ensure_domain_columns_and_backfill: add domain column to {corpus_table} failed", e, level="WARNING")

    # backfill corpus.domain from items.domain
    try:
        execute_write(f"""
            UPDATE {corpus_table}
            SET domain = (
              SELECT i.domain FROM items i WHERE i.id = {corpus_table}.project_id
            )
            WHERE domain IS NULL AND project_id IS NOT NULL;
        """)
    except Exception as e:
        log_exception(f"ensure_domain_columns_and_backfill: backfill domain for {corpus_table} failed", e)

ensure_domain_columns_and_backfill(conn, cur, corpus_table="corpus")
ensure_poly_tables(conn, cur)
# ensure_domain_columns_and_backfill(conn, cur, corpus_table="corpus_main")

def _get_domain_for_proj(cur, project_id: int):
    try:
        row = execute_write(
            "SELECT domain FROM items WHERE id=?",
            (int(project_id),)
        ).fetchone()
        if not row:
            return None
        val = (row[0] or "").strip()
        return val or None
    except Exception:
        return None
try:
    execute_write("CREATE INDEX IF NOT EXISTS idx_term_ext_project ON term_ext(project_id)")
except Exception as e:
    print("索引创建跳过:", e)

def _has_col(table: str, col: str) -> bool:
    execute_write(f"PRAGMA table_info({table})")
    return any(r[1] == col for r in cur.fetchall())

def ensure_col(table: str, col: str, col_type: str):
    execute_write(f"PRAGMA table_info({table})")
    cols = {r[1] for r in cur.fetchall()}
    if col not in cols:
        execute_write(f"ALTER TABLE {table} ADD COLUMN {col} {col_type}")

# —— 建表
execute_write("""
CREATE TABLE IF NOT EXISTS items (
    id INTEGER PRIMARY KEY,
    title TEXT NOT NULL,
    body TEXT,
    tags TEXT,
    domain TEXT,
    type TEXT,
    created_at TEXT DEFAULT (datetime('now'))
);
""")

execute_write("""
CREATE TABLE IF NOT EXISTS item_ext (
    id INTEGER PRIMARY KEY,
    item_id INTEGER,
    src_path TEXT,
    FOREIGN KEY(item_id) REFERENCES items(id)
);
""")

execute_write("""
CREATE TABLE IF NOT EXISTS project_files (
    id INTEGER PRIMARY KEY,
    project_id INTEGER NOT NULL,
    file_path TEXT NOT NULL,
    file_name TEXT,
    uploaded_at TEXT DEFAULT (datetime('now')),
    note TEXT,
    FOREIGN KEY(project_id) REFERENCES items(id)
);
""")

execute_write("""
CREATE TABLE IF NOT EXISTS term_ext (
    id INTEGER PRIMARY KEY,
    source_term TEXT NOT NULL,
    target_term TEXT,
    domain TEXT,
    project_id INTEGER,
    strategy TEXT,
    example TEXT,
    category TEXT,
    FOREIGN KEY(project_id) REFERENCES items(id)
);
""")

execute_write("""
CREATE TABLE IF NOT EXISTS trans_ext (
    id INTEGER PRIMARY KEY,
    project_id INTEGER,
    src_text TEXT,
    tgt_text TEXT,
    mode TEXT,
    segment_count INTEGER,
    created_at TEXT DEFAULT (datetime('now')),
    FOREIGN KEY(project_id) REFERENCES items(id)
);
""")

execute_write("""
CREATE TABLE IF NOT EXISTS corpus (
    id INTEGER PRIMARY KEY,
    project_id INTEGER,
    text TEXT,
    lang TEXT,
    source TEXT,
    created_at TEXT DEFAULT (datetime('now')),
    FOREIGN KEY(project_id) REFERENCES items(id)
);
""")

# —— 兜底补列
for t, cols in {
    "items": [("type","TEXT"),("tags","TEXT"),("scene","TEXT"),("prompt","TEXT"),
              ("mode","TEXT"),("body","TEXT"),("created_at","TEXT"),("updated_at","TEXT"),("trans_type","TEXT")],
    "item_ext": [("src_path","TEXT")],
    "term_ext": [("domain","TEXT"),("project_id","INTEGER"),("strategy","TEXT"),
                 ("example","TEXT"),("note","TEXT"), ("category","TEXT")],
    "trans_ext": [("stats_json","TEXT"),("segments","INTEGER"),("term_hit_total","INTEGER")],
    "corpus": [
        ("title","TEXT"),("project_id","INTEGER"),("lang_pair","TEXT"),
        ("src_text","TEXT"),("tgt_text","TEXT"),("note","TEXT"),
        ("created_at","TEXT"),("domain","TEXT"),("scene","TEXT"),
        ("unit","TEXT"),("source_type","TEXT"),
    ],
}.items():
    for c, tp in cols:
        ensure_col(t, c, tp)
execute_write("UPDATE items SET type='project' WHERE IFNULL(type,'')=''")
execute_write("UPDATE items SET created_at = COALESCE(created_at, strftime('%Y-%m-%d %H:%M:%S','now'))")
ensure_col("term_ext", "example_vector_id", "INTEGER")

# --- 术语表字段兼容:缺少 project_id 时补建 ---
try:
    execute_write("PRAGMA table_info(term_ext)")
    cols = [c[1] for c in cur.fetchall()]
    if "project_id" not in cols:
        execute_write("ALTER TABLE term_ext ADD COLUMN project_id INTEGER")
except Exception as e:
    st.warning(f"术语表结构检查:{e}")

# ========== DeepSeek 参数/调用 ==========
def get_deepseek():
    try:
        ak = st.secrets["deepseek"]["api_key"]
        model = st.secrets["deepseek"].get("model", "deepseek-chat")
        return ak, model
    except Exception:
        return None, None

        # === 新增:术语提示 + 参考例句 ===
def _build_ref_context(project_id: int,
                       query_text: str,
                       topk: int = 20,
                       min_sim: float = 0.25,
                       prefer_side: str = "both",   # 当前暂未使用，保留以兼容旧参数
                       scope: str = "project",
                       top_n: int = 5) -> str:
    try:
        hits = semantic_retrieve(
            project_id,
            query_text,
            topk=topk,
            scope=scope,
        )
    except Exception as e:
        # 召回失败不影响主流程，只在 UI 环境下做个轻提示
        try:
            st.warning(f"参考检索失败: {e}")
        except Exception as e2:
            log_exception("Failed to render warning for semantic retrieval error", e2, level="WARNING")
        return ""

    if not hits:
        return ""

    # 1) 过滤低相似度 + 去重
    seen = set()
    selected = []

    for sc, meta, src_sent, tgt_sent in hits:
        try:
            score = float(sc or 0.0)
        except Exception:
            score = 0.0

        if score < min_sim:
            continue

        ch = (src_sent or "").strip()
        en = (tgt_sent or "").strip()

        # 中英文都空，就跳过
        if not ch and not en:
            continue

        key = (ch, en)
        if key in seen:
            continue
        seen.add(key)

        selected.append((score, meta, ch, en))
        if len(selected) >= top_n:
            break

    if not selected:
        best = hits[0]
        sc, meta, ch, en = best
        ch = (ch or "").strip()
        en = (en or "").strip()
        if not ch and not en:
            return ""
        try:
            sc = float(sc or 0.0)
        except Exception:
            sc = 0.0
        selected = [(sc, meta, ch, en)]

    # 2) 拼成多行文本
    ctx_lines = ["参考例句(用于保持术语与风格一致):"]
    for idx, (sc, meta, ch, en) in enumerate(selected, 1):
        dom = (meta.get("domain") or "").strip() if isinstance(meta, dict) else ""
        title = (meta.get("title") or "").strip() if isinstance(meta, dict) else ""
        tag_info = " · ".join(x for x in [dom, title] if x)

        ch_show = ch.replace("\n", " ").strip()
        en_show = en.replace("\n", " ").strip()

        if en_show:
            line = (
                f"例句{idx} 原文:{ch_show}\n"
                f"       译文:{en_show}"
                f"(sim={sc:.2f}{'，'+tag_info if tag_info else ''})"
            )
        else:
            line = f"例句{idx}:{ch_show}(sim={sc:.2f}{'，'+tag_info if tag_info else ''})"

        ctx_lines.append(line)

        # 控制总长度，避免 prompt 过长
        if sum(len(x) for x in ctx_lines) > 1800:
            break

    return "\n".join(ctx_lines) if len(ctx_lines) > 1 else ""
# -------- Glossary & Instruction helpers (放在 ds_translate 上方) --------
def build_term_hint(term_dict: dict, lang_pair: str, max_terms: int = 80) -> str:
    lines = []
    seen = set()
    items = list(term_dict.items())[: max_terms * 2]  # 稍多取一些.过滤空后再截断

    for src, val in items:
        if src is None: 
            continue
        src = str(src).strip()
        if not src or src in seen:
            continue

        tgt, pos, note = None, None, None
        if isinstance(val, dict):
            tgt  = (val.get("target") or val.get("tgt") or "").strip()
            pos  = (val.get("pos") or "").strip() or None
            note = (val.get("usage_note") or val.get("note") or "").strip() or None
        elif isinstance(val, (list, tuple)) and len(val) >= 1:
            tgt = str(val[0]).strip()
            if len(val) >= 2:
                pos = (str(val[1]).strip() or None)
        else:
            tgt = str(val or "").strip()

        if not tgt:
            continue

        seen.add(src)
        if pos:
            line = f"- When '{src}' is a {pos}, translate it as '{tgt}'."
        else:
            line = f"- Translate '{src}' as '{tgt}'."
        if note:
            line += f" ({note})"
        lines.append(line)

    if not lines:
        return ""  # 没有可用术语就返回空串.不要干扰提示

    header = "GLOSSARY (STRICT):\n"
    return header + "\n".join(lines[:max_terms]) + "\n"

def build_instruction(lang_pair: str) -> str:
    lp = (lang_pair or "").replace(" ", "")
    if "中→英" in lp or "中->英" in lp or "zh" in lp.lower() and "en" in lp.lower():
        return (
            "Translate the source text from Chinese to English. "
            "Use a professional, natural style; follow the GLOSSARY (STRICT) exactly; "
            "preserve proper nouns and numbers; keep paragraph structure. "
            "Do not add explanations."
        )
    if "英→中" in lp or "英->中" in lp or "en" in lp.lower() and "zh" in lp.lower():
        return (
            "Translate the source text from English to Chinese. "
            "用专业、通顺、符合领域文体的中文表达;严格遵守上方 GLOSSARY (STRICT);"
            "专有名词、数字与计量单位保持准确;段落结构保持一致。不得添加解释。"
        )
    # 兜底
    return (
        "Translate the source text. Follow the GLOSSARY (STRICT) exactly. "
        "Keep the original structure and do not add explanations."
    )

def ds_translate(
    block: str,
    term_dict: dict,
    lang_pair: str,
    ak: str,
    model: str,
    ref_context: str = "",
    fewshot_examples=None,
    extra_prompt: str = "",
) -> str:
    import requests

    if not block.strip():
        return ""

    # 使用统一的函数构建术语提示和指令，避免重复代码
    term_hint = build_term_hint(term_dict, lang_pair)
    instr = build_instruction(lang_pair)

    system_msg = (
        "You are a senior professional translator. Prioritize accuracy, faithfulness, and consistent terminology. "
        "No hallucinations. If a term mapping is provided, follow it strictly. "
        "A locked glossary may be included—always use the exact target given, no synonyms or rewording."
    )

    extra_block = ""
    if extra_prompt and str(extra_prompt).strip():
        extra_block = "ADDITIONAL REQUIREMENTS:\n" + str(extra_prompt).strip() + "\n"

    user_msg = (
        f"{term_hint}"
        + (f"REFERENCE CONTEXT (use if relevant):\n{ref_context}\n\n" if ref_context else "")
        + (extra_block + "\n" if extra_block else "")
        + "INSTRUCTION:\n" + instr + "\n"
        + "RESPONSE FORMAT:\n- Return ONLY the final translation text, no explanations, no backticks.\n\n"
        + "SOURCE:\n" + block
    )

    messages = [{"role": "system", "content": system_msg}]
    if fewshot_examples:
        for ex in fewshot_examples:
            src_demo = (ex.get("src") or "").strip()
            tgt_demo = (ex.get("tgt") or "").strip()
            if not (src_demo and tgt_demo):
                continue
            title = ex.get("title") or ""
            demo_user = f"【参考示例:{title}】\n源文:\n{src_demo}"
            messages.append({"role": "user", "content": demo_user})
            messages.append({"role": "assistant", "content": tgt_demo})
    messages.append({"role": "user", "content": user_msg})

    url = "https://api.deepseek.com/v1/chat/completions"
    headers = {"Authorization": f"Bearer {ak}", "Content-Type": "application/json"}
    payload = {
        "model": model or "deepseek-chat",
        "messages": messages,
        "temperature": 0.2,
    }

    for attempt in range(3):
        try:
            resp = requests.post(url, headers=headers, json=payload, timeout=60)
            if resp.status_code == 200:
                data = resp.json()
                return data["choices"][0]["message"]["content"].strip()
            else:
                txt = f"[DeepSeek {resp.status_code}] {resp.text}"
                if resp.status_code in (429, 500, 502, 503, 504) and attempt < 2:
                    time.sleep(1.5 * (attempt + 1))
                    continue
                # 最终失败: 记一条错误日志
                log_event(
                    "ERROR",
                    "DeepSeek HTTP error",
                    status_code=resp.status_code,
                    body=resp.text[:500],
                )
                return txt
        except Exception as e:
            if attempt < 2:
                time.sleep(1.5 * (attempt + 1))
                continue
            # 重试用尽仍然异常: 记一条错误日志
            log_event(
                "ERROR",
                "DeepSeek request exception",
                error=str(e),
            )
            return f"[DeepSeek Request Error] {e}"

    log_event("ERROR", "DeepSeek unknown failure")
    return "[DeepSeek Error] Unknown failure."

def translate_block_with_kb(
    cur,
    project_id: int,
    block_text: str,
    lang_pair: str,
    ak: str,
    model: str,
    use_semantic: bool = True,
    scope: str = "project",
    fewshot_examples=None,
    locked_term_map: dict | None = None,
    extra_prompt: str = "",
):
    blk = (block_text or "").strip()
    if not blk:
        return {
            "src": "",
            "tgt": "",
            "project_id": project_id,
            "lang_pair": lang_pair,
            "term_map_all": {},
            "terms_in_block": {},
            "terms_corpus_dyn": {},
            "terms_final": {},
            "term_meta": [],
            "ref_context": "",
            "poly_hint": "",
            "violated_terms": [],
        }

    # 1) 静态术语（项目+全局），统一接口
    term_map_all_db, term_meta = get_terms_for_project(cur, project_id, use_dynamic=True)
    # 若传入锁定术语表，优先用锁定表，确保“先定术语表，再翻译”
    term_map_all = dict(locked_term_map) if locked_term_map else term_map_all_db

    # 2) 命中检测工具：给“本段”和“语料参考”共用
    def _detect_hits(text: str, term_map: dict[str, str]) -> dict[str, str]:
        txt_low = (text or "").lower()
        out = {}
        for k, v in (term_map or {}).items():
            if not k:
                continue
            key_low = k.lower()
            if key_low in txt_low or k in text:
                out[k] = v
        return out

    # 2.1 当前段落文本中命中的术语
    terms_in_block = _detect_hits(blk, term_map_all)

    # 3) 参考例句(来自语料库语义召回)
    if use_semantic:
        try:
            ref_context = _build_ref_context(
                project_id,
                blk,
                topk=20,
                min_sim=0.25,
                prefer_side="both",
                scope=scope,
            )
        except Exception:
            # 召回失败不影响主流程
            ref_context = ""
    else:
        ref_context = ""

    # 3.1 语料驱动术语：在参考例句文本中命中的静态术语
    if ref_context:
        terms_corpus_dyn = _detect_hits(ref_context, term_map_all)
    else:
        terms_corpus_dyn = {}

    # 3.2 最终注入 Prompt 的术语:
    #    - 不再全量注入；只注入“本段命中 + 参考命中”的并集
    terms_final = dict(terms_in_block)
    for k, v in terms_corpus_dyn.items():
        if k not in terms_final:
            terms_final[k] = v
    # 3.3 一词多义提示（软约束）
    try:
        domain_val = _get_domain_for_proj(cur, int(project_id))
    except Exception:
        domain_val = None
    try:
        scene_val = _get_scene_for_proj(cur, int(project_id))
    except Exception:
        scene_val = None
    poly_hint = ""
    try:
        poly_hint = build_poly_hint(cur, int(project_id), blk, domain_val, scene_val, lang_pair, topk=8)
    except Exception:
        poly_hint = ""
    if poly_hint:
        ref_context = (poly_hint + "\n\n" + (ref_context or "")).strip()

    # 4) 调用 DeepSeek 翻译（只喂最终术语）
    tgt = ds_translate(
        block=blk,
        term_dict=terms_final,
        lang_pair=lang_pair,
        ak=ak,
        model=model,
        ref_context=ref_context,
        fewshot_examples=fewshot_examples,
        extra_prompt=extra_prompt,
    )

    # 5) 粗略术语一致性检查(以“最终注入”的术语为准)
    violated = check_term_consistency(tgt, terms_final, blk)

    return {
        "src": blk,
        "tgt": tgt,
        "project_id": project_id,
        "lang_pair": lang_pair,
        "term_map_all": term_map_all,
        "terms_in_block": terms_in_block,
        "terms_corpus_dyn": terms_corpus_dyn,
        "terms_final": terms_final,
        "term_meta": term_meta,
        "ref_context": ref_context,
        "poly_hint": poly_hint,
        "violated_terms": violated,
    }

def ds_extract_terms(text: str, ak: str, model: str, src_lang: str = "zh", tgt_lang: str = "en"):

    if not text.strip():
        return []

    system_msg = (
        "You are a terminology mining assistant. Extract high-value bilingual term pairs "
        "suitable for a project glossary. Return JSON array only. No extra text."
    )
    user_msg = f"""
Source language: {src_lang}
Target language: {tgt_lang}
任务:从给定文本中抽取双语术语条目.输出 JSON 数组。字段名与取值必须是中文。
字段定义:
- source_term: 源语(中文术语或专名)
- target_term: 译文(英文)
- target_term_alt: 译文备选(英文, 与 target_term 不同)
- domain: 领域.取值集合之一:["政治","经济","文化","文物","金融","法律","其他"]
- strategy: 翻译策略.取值集合之一:["直译","意译","转译","音译","省略","增译","规范化","其他"]
- example: 例句(原文中包含该术语的一句.尽量保留标点)

要求:
1) 仅输出 JSON.不要多余说明。
2) 同一术语重复时合并.选择最典型的例句。
3) 若无法判断 domain/strategy.填“其他”。
4) 不要限制数量，尽量多抽取高价值术语。

Text:
{text}
"""

    url = "https://api.deepseek.com/v1/chat/completions"
    headers = {"Authorization": f"Bearer {ak}", "Content-Type": "application/json"}
    payload = {
        "model": model or "deepseek-chat",
        "messages": [
            {"role": "system", "content": system_msg},
            {"role": "user", "content": user_msg},
        ],
        "temperature": 0.1,
    }
    try:
        r = requests.post(url, headers=headers, json=payload, timeout=60)
        r.raise_for_status()
        data = r.json()
    except Exception as e:
        log_event("ERROR", "DeepSeek term extraction request failed", error=str(e))
        raise RuntimeError(f"DeepSeek 请求失败: {e}") from e
    
    # 提取响应文本
    try:
        txt = data["choices"][0]["message"]["content"].strip()
    except (KeyError, IndexError, TypeError) as e:
        log_event("ERROR", "DeepSeek term extraction missing content", data_preview=str(data)[:500])
        raise RuntimeError("DeepSeek 返回内容缺少 message") from e

    # 1. 去除 Markdown 标记
    if "```json" in txt:
        txt = txt.split("```json")[1].split("```")[0]
    elif "```" in txt:
        txt = txt.split("```")[1].split("```")[0]
    
    # 2. 寻找 JSON 数组边界
    start = txt.find("[")
    end = txt.rfind("]")
    if start == -1 or end == -1:
        log_event("ERROR", "DeepSeek term extraction missing JSON array", response_preview=txt[:500])
        raise RuntimeError("DeepSeek 返回未包含 JSON 数组")

    # 3. 解析 JSON
    try:
        arr = json.loads(txt[start:end+1])
    except Exception as e:
        log_event("ERROR", "DeepSeek term extraction JSON parse failed", response_preview=txt[start:end+1][:500])
        raise RuntimeError("DeepSeek JSON 解析失败") from e

    out = []
    for o in arr:
        src = (o.get("source_term") or o.get("source") or "").strip()
        tgt = (o.get("target_term") or o.get("target") or "").strip()
        tgt_alt = (o.get("target_term_alt") or o.get("alt") or "").strip()
        dom = (o.get("domain") or "").strip() or None
        strat = (o.get("strategy") or "").strip() or None
        ex = (o.get("example") or "").strip() or None
        if src:
            out.append({
                "source_term": src,
                "target_term": tgt,
                "target_term_alt": tgt_alt,
                "domain": dom,
                "strategy": strat,
                "example": ex,
            })
    return out

# ========== 文件读写与导出 ==========

def _read_docx_xml(path: str) -> str:
    import zipfile
    import xml.etree.ElementTree as ET
    try:
        with zipfile.ZipFile(path) as zf:
            xml_names = [n for n in zf.namelist() if n.startswith("word/") and n.endswith(".xml")]
            if not xml_names:
                return ""
            ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
            paras = []
            for name in xml_names:
                try:
                    xml = zf.read(name)
                except Exception:
                    continue
                try:
                    root = ET.fromstring(xml)
                except Exception:
                    continue
                for pnode in root.findall('.//w:p', ns):
                    texts = [t.text for t in pnode.findall('.//w:t', ns) if t.text]
                    if texts:
                        paras.append(''.join(texts))
            return "\n".join(paras)
    except Exception:
        return ""

def read_source_file(path: str) -> str:
    if not path or not os.path.exists(path):
        return ""
    ext = os.path.splitext(path)[1].lower()
    if ext == ".txt":
        for enc in ["utf-8", "utf-8-sig", "gb18030", "gbk"]:
            try:
                with open(path, "r", encoding=enc) as f:
                    return f.read()
            except Exception:
                continue
        with open(path, "r", errors="ignore") as f:
            return f.read()
    elif ext == ".docx":
        text_out = ""
        try:
            from docx import Document
            doc = Document(path)
            paras = [p.text for p in doc.paragraphs if p.text]
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for p in cell.paragraphs:
                            if p.text:
                                paras.append(p.text)
            text_out = "\n".join(paras).strip()
        except Exception as e:
            log_exception(f"read_source_file: failed to read DOCX file", e, level="WARNING")

        if not text_out:
            try:
                text_out = _read_docx_xml(path).strip()
            except Exception as e:
                log_exception(f"read_source_file: failed to read DOCX xml", e, level="WARNING")

        return text_out
    elif ext == ".pdf":
        try:
            from pdfplumber import open as pdf_open
            with pdf_open(path) as pdf:
                pages = [page.extract_text() or "" for page in pdf.pages]
            return "\n\n".join(pages)
        except Exception as e:
            log_exception(f"read_source_file: failed to read PDF file", e, level="WARNING")
            return ""
    elif ext == ".xlsx":
        try:
            xls = pd.ExcelFile(path)
            parts = []
            for sheet in xls.sheet_names:
                df = xls.parse(sheet)
                parts.append(df.astype(str).to_csv(sep=" ", index=False, header=False))
            return "\n".join(parts)
        except Exception as e:
            log_exception(f"read_source_file: failed to read Excel file", e, level="WARNING")
            return ""
    else:
        # 兜底尝试文本读取
        try:
            with open(path, "r", encoding="utf-8") as f:
                return f.read()
        except Exception as e:
            log_exception(f"read_source_file: failed to read file (fallback)", e, level="WARNING")
            return ""

def build_bilingual_lines(src_text: str, tgt_text: str):
    return pair_paragraphs(src_text, tgt_text)

def export_csv_bilingual(src_text: str, tgt_text: str) -> bytes:
    s, t = build_bilingual_lines(src_text, tgt_text)
    df = pd.DataFrame({"Source": s, "Target": t})
    buf = io.StringIO()
    df.to_csv(buf, index=False)
    return buf.getvalue().encode("utf-8-sig")

def export_docx_bilingual(src_text: str, tgt_text: str) -> bytes:
    try:
        from docx import Document
        from docx.oxml.ns import qn
    except Exception as e:
        log_exception("export_docx_bilingual: missing python-docx library", e, level="ERROR")
        st.error("缺少 python-docx.请先安装:pip install python-docx")
        return b""
    doc = Document()
    # 基础字体
    try:
        doc.styles['Normal'].font.name = 'Calibri'
        doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    except Exception as e:
        log_exception("export_docx_bilingual: failed to set default fonts", e, level="WARNING")
    s, t = build_bilingual_lines(src_text, tgt_text)
    table = doc.add_table(rows=1, cols=2)
    hdr = table.rows[0].cells
    hdr[0].text = "Source"
    hdr[1].text = "Target"
    for a, b in zip(s, t):
        row = table.add_row().cells
        row[0].text = a
        row[1].text = b
    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()

# ==== 工具:文件读取 / 分句 / 向量 / 对齐 / 索引 ====

def _lazy_import_doc_pdf():
    docx = pdfplumber = None
    try:
        import docx as _docx
        docx = _docx
    except Exception as e:
        log_exception("_lazy_import_doc_pdf: docx import failed", e, level="WARNING")
    try:
        import pdfplumber as _pdfplumber
        pdfplumber = _pdfplumber
    except Exception as e:
        log_exception("_lazy_import_doc_pdf: pdfplumber import failed", e, level="WARNING")
    return docx, pdfplumber
# 段落切分
def split_paragraphs(text: str) -> list[str]:
    text = (text or "").replace("\r\n", "\n").replace("\r", "\n")

    if "\n\n" not in text and "\n \n" not in text:
        lines = [ln.strip() for ln in text.split("\n")]
        return [ln for ln in lines if ln]

    # 正常:有空行分段
    parts = re.split(r"\n\s*\n+", text)
    paras = []
    for p in parts:
        p = p.strip()
        if not p:
            continue
        # 段内如果还有软换行，压成空格，避免导出时被拆成多行
        p = re.sub(r"\s*\n\s*", " ", p)
        paras.append(p)
    return paras

def pair_paragraphs(src_full: str, tgt_full: str) -> tuple[list[str], list[str]]:
    src_paras = split_paragraphs(src_full or "")
    tgt_paras = split_paragraphs(tgt_full or "")

    n = max(len(src_paras), len(tgt_paras))
    src_paras += [""] * (n - len(src_paras))
    tgt_paras += [""] * (n - len(tgt_paras))
    return src_paras, tgt_paras

# 预编译(可放全局)
_RE_WS = re.compile(r"[ \t\u00A0\u200B\u200C\u200D]+")
_RE_ZH_SENT = re.compile(r"(?<=[。！？；])\s*")           # 中文句末
_RE_EN_SENT = re.compile(r"(?<=[\.\?\!;:])\s+")          # 英文句末(放宽，不强制大写)
_RE_BLANK_PARA = re.compile(r"\n{2,}")                   # 空行分段

def _norm_text(text: str) -> str:
    t = (text or "").replace("\r\n", "\n").replace("\r", "\n").replace("\x0b", "\n")
    t = _RE_WS.sub(" ", t)
    t = re.sub(r"\n{3,}", "\n\n", t)  
    return t.strip()

def _is_zh(text: str) -> bool:
    # 简单判定:含有较多中文字符
    zh_hits = len(re.findall(r"[\u4e00-\u9fff]", text))
    en_hits = len(re.findall(r"[A-Za-z]", text))
    return zh_hits >= en_hits

def split_sents(
    text: str,
    lang_hint: str = "auto",
    min_char: int = 1,
    prefer_newline: bool = True,
    **kwargs,
):
    # 兼容旧参数名 lang=
    lang = kwargs.get("lang", lang_hint)

    t = _norm_text(text)
    if not t:
        return []

    pieces = []

    # A) 若文本中有换行 & prefer_newline=True:先按行切，再在行内按句末细分
    if prefer_newline and ("\n" in t):
        lines = [ln.strip() for ln in t.split("\n") if ln.strip()]
        for ln in lines:
            if lang == "auto":
                cur_lang = "zh" if _is_zh(ln) else "en"
            else:
                cur_lang = lang

            if cur_lang.startswith(("zh", "cn")):
                sents = [s.strip() for s in _RE_ZH_SENT.split(ln) if s and s.strip()]
            else:
                sents = [s.strip() for s in _RE_EN_SENT.split(ln) if s and s.strip()]

            pieces.extend(sents if sents else [ln])
    else:
        # B) 没有换行或不偏好换行:整块按句末标点切
        if lang == "auto":
            cur_lang = "zh" if _is_zh(t) else "en"
        else:
            cur_lang = lang

        if cur_lang.startswith(("zh", "cn")):
            pieces = [s.strip() for s in _RE_ZH_SENT.split(t) if s and s.strip()]
        else:
            pieces = [s.strip() for s in _RE_EN_SENT.split(t) if s and s.strip()]

        if not pieces:
            pieces = [t]

    # 过滤过短片段
    return [x for x in pieces if len(x) >= min_char]

# 兼容旧函数名
split_sentences = split_sents

def split_blocks(text: str, max_len: int = 1200):
    blocks, curbuf = [], ""
    for line in text.splitlines(True):
        if len(curbuf) + len(line) > max_len:
            if curbuf:
                blocks.append(curbuf)
            curbuf = ""
        curbuf += line
    if curbuf:
        blocks.append(curbuf)
    return blocks
# =========================
# 术语提示 & 译后一致性检查
# =========================
def check_term_consistency(out_text: str, term_dict: dict, source_text: str = "") -> list:
    if not out_text or not term_dict:
        return []
    warns = []
    s = (source_text or "")[:2000]  # 限长度.防极端长文本
    out_low = out_text.lower()
    for k, v in term_dict.items():
        if not k or not v:
            continue
        # 如果源文包含该术语键(大小写忽略英文;中文直接包含)
        hit_src = (k.lower() in s.lower()) if any(ord(ch) < 128 for ch in k) else (k in s)
        if hit_src:
            # 译文是否出现目标术语(同理大小写容忍英文)
            ok = (v.lower() in out_low) if any(ord(ch) < 128 for ch in v) else (v in out_text)
            if not ok:
                warns.append(f"{k}→{v}")
    return warns

def _lazy_embedder():
    # 优先 sentence-transformers;失败退化到 TF-IDF
    try:
        emb = _make_sbert_encoder(
            "sentence-transformers/paraphrase-multilingual-MiniLM-L12-v2",
            normalize_embeddings=True,
            extra_normalize=False,
        )
        return emb, "sbert"
    except Exception:
        from sklearn.feature_extraction.text import TfidfVectorizer
        def _emb(texts):
            vec = TfidfVectorizer(min_df=1).fit_transform(texts)
            # 归一化
            norms = _np.sqrt((vec.multiply(vec)).sum(axis=1)).A.ravel() + 1e-8
            vec = vec.multiply(1/norms[:,None])
            return vec
        return _emb, "tfidf"
def import_corpus_from_upload(
    st,
    cur,
    conn,
    *,
    pid: int | None,
    title: str | None,
    lp: str,
    pairs,
    src_text: str | None,
    tgt_text: str | None,
    default_title: str = "",
    build_after_import: bool = False,
):
    # 统一标题
    base_title = (title or default_title or "").strip() or "未命名语料"

    # 小工具：把 [(s,t,score)] / [(s,t)] 统一成 [(s,t)]
    def normalize_pairs_to2(pairs_in):
        if not pairs_in:
            return []
        if len(pairs_in[0]) == 3:
            return [(s, t) for (s, t, _) in pairs_in]
        return pairs_in

    # 1) 双语 pairs 情况
    if pairs:
        pairs2 = normalize_pairs_to2(pairs)
        ins = 0
        for s, t in pairs2:
            s = (s or "").strip()
            t = (t or "").strip()
            if not (s or t):
                continue
            execute_write(
                """
                INSERT INTO corpus(title, project_id, lang_pair, src_text, tgt_text, note, created_at)
                VALUES (?, ?, ?, ?, ?, ?, datetime('now'))
                """,
                (
                    base_title,
                    pid,
                    lp,
                    s or None,
                    t or None,
                    "auto-import",
                ),
            )
            ins += 1
        st.success(f" 已写入语料库 {ins} 条。")

        # 导入后重建当前项目语义索引
        if build_after_import and pid:
            res_idx = rebuild_project_semantic_index(pid)
            if res_idx.get("ok"):
                st.success(
                    f" 向量索引已更新: 新增 {res_idx['added']}，总量 {res_idx['total']}。"
                )
            else:
                st.warning(f"索引未更新: {res_idx.get('msg','未知错误')}")

        return

    # 2) 单语 src_text 情况（策略/单语语料）
    if src_text and not tgt_text:
        lang_hint = "zh" if (lp or "").startswith("中") else "en"
        sents = split_sents(src_text, lang_hint)
        ins = 0
        for s in sents:
            s = (s or "").strip()
            if not s:
                continue
            execute_write(
                """
                INSERT INTO corpus(title, project_id, lang_pair, src_text, tgt_text, note, created_at)
                VALUES (?, ?, ?, ?, NULL, ?, datetime('now'))
                """,
                (
                    base_title,
                    pid,
                    lp,
                    s,
                    "mono",  # 标记：单语策略/语料
                ),
            )
            ins += 1
        st.success(f" 已写入语料库 {ins} 条。")

        if build_after_import and pid:
            res_idx = rebuild_project_semantic_index(pid)
            if res_idx.get("ok"):
                st.success(
                    f" 向量索引已更新: 新增 {res_idx['added']}，总量 {res_idx['total']}。"
                )
            else:
                st.warning(f"索引未更新: {res_idx.get('msg','未知错误')}")
        return

    # 3) 其他情况
    st.warning("原文和译文都为空，无法写入语料库。")


def _extract_terms_for_upload(text_chunks, *, lp: str, pid: int | None):
    """结合上传文本调用 DeepSeek 提取术语，写入 term_ext。"""
    ak, model = get_deepseek()
    if not ak:
        st.warning("未检测到 DeepSeek Key(请在“设置”页配置)，已跳过术语提取。")
        return

    big = "\n".join([t for t in text_chunks if t]).strip()
    if not big:
        st.info("没有可用文本，跳过术语提取。")
        return

    if lp.startswith("中"):
        src_lang, tgt_lang = "zh", "en"
    elif lp.startswith("英"):
        src_lang, tgt_lang = "en", "zh"
    else:
        src_lang, tgt_lang = "zh", "en"

    res = safe_api_call(
        ds_extract_terms,
        big,
        ak,
        model,
        context="新建语料术语抽取",
        fallback=[],
        src_lang=src_lang,
        tgt_lang=tgt_lang,
    )
    if not res:
        st.info("未提取到术语或解析失败。")
        return

    ins = 0
    for o in res:
        execute_write(
            """
            INSERT INTO term_ext (
                source_term, target_term, domain, project_id,
                strategy, example
            )
            VALUES (?, ?, ?, ?, ?, ?)
            """,
            (
                o.get("source_term") or "",
                (o.get("target_term") or None),
                (o.get("domain") or None),
                pid,
                (o.get("strategy") or "corpus-import"),
                (o.get("example") or None),
            ),
        )
        ins += 1
    st.success(f" 已写入术语库 {ins} 条。")

from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity
import numpy as np

def align_semantic(src_sents, tgt_sents, max_jump=3):
    if not src_sents or not tgt_sents:
        return []

    emb, kind = _lazy_embedder()

    # === 优先使用 SBERT ===
    if kind == "sbert":
        E1 = emb(src_sents)
        E2 = emb(tgt_sents)
        sims = E1 @ E2.T  # (n, m)
    else:
        # === TF-IDF 回退:确保同一词表维度 ===
        vec = TfidfVectorizer(
            analyzer="char_wb",  # 字符 n-gram 对中英混合最稳
            ngram_range=(1, 2),
            min_df=1
        )
        combo = src_sents + tgt_sents
        X = vec.fit_transform(combo)
        n = len(src_sents)
        E1 = X[:n, :]
        E2 = X[n:, :]

        # 稀疏→相似度矩阵
        sims = cosine_similarity(E1, E2, dense_output=True)  # shape (n, m)

    # === 贪心对齐 ===
    i = j = 0
    n, m = len(src_sents), len(tgt_sents)
    pairs = []
    while i < n and j < m:
        j_min = max(0, j - max_jump)
        j_max = min(m, j + max_jump + 1)
        window = sims[i, j_min:j_max]
        if window.size == 0:
            break
        k = int(window.argmax())
        j_sel = j_min + k
        score = float(sims[i, j_sel])
        pairs.append((src_sents[i], tgt_sents[j_sel], score))
        i += 1
        j = j_sel + 1
    return pairs

# ========== 术语库管理 ==========
def render_term_management(st, cur, conn, base_dir, key_prefix="term"):
    sk = make_sk(key_prefix)

    st.subheader(" 术语库管理")
    sub_tabs = st.tabs(["查询与编辑", "批量导入 CSV", "统计与导出", "快速搜索", "批量挂接项目", "从历史提取术语", "分类管理"])

    # —— 查询与编辑
    with sub_tabs[0]:
        sk0 = lambda name: f"{key_prefix}_t0_{name}"

        c1, c2, c3, c4 = st.columns(4)
        with c1: kw = st.text_input("关键词(源/目标/例句)", "", key=sk("kw_example"))
        with c2: dom = st.text_input("领域", "", key=sk("dom"))
        with c3: strat = st.text_input("策略", "", key=sk("strat"))
        with c4: 
            pid = st.text_input("项目ID过滤", "", key=sk("pid"))
            # P0-2: 添加输入验证
            if pid and not pid.isdigit():
                st.error("❌ 项目ID必须是数字，当前输入无效")
                pid = ""
        cat = st.text_input("分类(支持子串)", "", key=sk("cat"))

        # —— 兼容老库:如无 category 列则补列(已有会忽略)
        try:
            execute_write("ALTER TABLE term_ext ADD COLUMN category TEXT;")
        except Exception as e:
            log_exception("render_terminology_mgmt: add category column failed (likely already exists)", e, level="INFO")

        # 1) 以数据库真实列为准检测是否存在 category
        cols_db = [c[1].lower() for c in cur.execute("PRAGMA table_info(term_ext);").fetchall()]
        has_category = ("category" in cols_db)

        # 2) 拼 SQL(只在 DB 真的有该列时才 SELECT category)
        base_cols = "id, source_term, target_term, domain, project_id, strategy, example"
        sel_cols = base_cols + (", category" if has_category else "")
        sql = f"SELECT {sel_cols} FROM term_ext WHERE 1=1"
        params = []

        if kw:
            like = f"%{kw}%"
            sql += " AND (IFNULL(source_term,'') LIKE ? OR IFNULL(target_term,'') LIKE ? OR IFNULL(example,'') LIKE ?)"
            params += [like, like, like]

        if dom:
            sql += " AND IFNULL(domain,'') LIKE ?"
            params += [f"%{dom}%"]

        if strat:
            sql += " AND IFNULL(strategy,'') LIKE ?"
            params += [f"%{strat}%"]

        if pid and str(pid).isdigit():
            sql += " AND IFNULL(project_id,0) = ?"
            params += [int(pid)]

        if has_category and cat:
            sql += " AND IFNULL(category,'') LIKE ?"
            params += [f"%{cat}%"]

        sql += " ORDER BY source_term COLLATE NOCASE LIMIT 1000"

        # 3) 查询并构造 DataFrame(表头与实际列对齐)
        rows = cur.execute(sql, params).fetchall()
        headers = ["ID","源术语","目标术语","领域","项目ID","策略","例句"]
        if has_category:
            headers.append("分类")

        df = pd.DataFrame(rows, columns=headers)
        st.caption(f"当前查询返回:{len(df)} 条")

        # 4) 空数据就不渲染编辑器
        if df.empty:
            st.info("没有匹配的术语。")
        else:
            # === 用 session_state 维护“当前表格”，含选择列 ===
            editor_df_key = sk0("editor_df")   # 专门存 DataFrame
            editor_key    = sk0("editor")      # data_editor 小部件本身

            # 初始化 / 尺寸变化时重置
            if editor_df_key not in st.session_state:
                work_df = df.copy()
                if "sel" not in work_df.columns:
                    work_df.insert(0, "sel", False)
                st.session_state[editor_df_key] = work_df
            else:
                work_df = st.session_state[editor_df_key]
                if len(work_df) != len(df):
                    work_df = df.copy()
                    if "sel" not in work_df.columns:
                        work_df.insert(0, "sel", False)
                    st.session_state[editor_df_key] = work_df

            # 动态构建列配置.只有当 DB 真有“分类”时才加入
            col_cfg = {
                "ID": st.column_config.NumberColumn("ID", disabled=True),
                "sel": st.column_config.CheckboxColumn("选择"),
                "源术语": "源术语",
                "目标术语": "目标术语",
                "领域": "领域",
                "项目ID": st.column_config.NumberColumn("项目ID", step=1, required=False),
                "策略": "策略",
                "例句": st.column_config.TextColumn("例句"),
            }
            if has_category:
                col_cfg["分类"] = st.column_config.TextColumn("分类")

            # 真正的编辑器:以 session_state 里的 DataFrame 为准
            edited = st.data_editor(
                st.session_state[editor_df_key],
                num_rows="dynamic",
                key=editor_key,
                column_config=col_cfg,
            )
            # 把编辑结果写回 session_state
            st.session_state[editor_df_key] = edited

            c1, c2, c3 = st.columns([1, 1, 2])

            # ---------------- 保存修改 ----------------
            with c1:
                if st.button(" 保存修改", type="primary", key=sk("save_terms")):
                    updated = inserted = 0
                    for _, row in edited.iterrows():
                        if pd.notna(row["ID"]):  # 更新
                            if has_category:
                                execute_write("""
                                    UPDATE term_ext
                                    SET source_term=?, target_term=?, domain=?,
                                        project_id=?, strategy=?, example=?, category=?
                                    WHERE id=?;
                                """, (
                                    (row["源术语"] or "").strip(),
                                    (row["目标术语"] or None),
                                    (row["领域"] or None),
                                    (int(row["项目ID"]) if pd.notna(row["项目ID"]) else None),
                                    (row["策略"] or None),
                                    (row["例句"] or None),
                                    (row.get("分类") or None),
                                    int(row["ID"])
                                ))
                            else:
                                execute_write("""
                                    UPDATE term_ext
                                    SET source_term=?, target_term=?, domain=?, project_id=?, strategy=?, example=?
                                    WHERE id=?;
                                """, (
                                    (row["源术语"] or "").strip(),
                                    (row["目标术语"] or None),
                                    (row["领域"] or None),
                                    (int(row["项目ID"]) if pd.notna(row["项目ID"]) else None),
                                    (row["策略"] or None),
                                    (row["例句"] or None),
                                    int(row["ID"])
                                ))
                            updated += cur.rowcount
                        else:  # 新增
                            if str(row["源术语"]).strip():
                                if has_category:
                                    execute_write("""
                                        INSERT INTO term_ext (
                                            source_term, target_term, domain, project_id,
                                            strategy, example, category
                                        )
                                        VALUES (?, ?, ?, ?, ?, ?, ?)
                                    """, (
                                        (row["源术语"] or "").strip(),
                                        (row["目标术语"] or None),
                                        (row["领域"] or None),
                                        (int(row["项目ID"]) if pd.notna(row["项目ID"]) else None),
                                        (row["策略"] or None),
                                        (row["例句"] or None),
                                        (row.get("分类") or None)
                                    ))
                                else:
                                    execute_write("""
                                        INSERT INTO term_ext (
                                            source_term, target_term, domain, project_id,
                                            strategy, example
                                        )
                                        VALUES (?, ?, ?, ?, ?, ?)
                                    """, (
                                        (row["源术语"] or "").strip(),
                                        (row["目标术语"] or None),
                                        (row["领域"] or None),
                                        (int(row["项目ID"]) if pd.notna(row["项目ID"]) else None),
                                        (row["策略"] or None),
                                        (row["例句"] or None),
                                    ))
                                inserted += 1
                    st.success(f" 已保存:更新 {updated} 条.新增 {inserted} 条。")
                    st.rerun()

            # ---------------- 全选 / 清空 / 删除 ----------------
            with c2:
                cc2a, cc2b, cc2c = st.columns([1, 1, 2])
                # 这里统一操作 session_state 里的 DataFrame
                cur_df = st.session_state[editor_df_key]

                if cc2a.button("全选", key=sk("sel_all")):
                    cur_df.loc[:, "sel"] = True
                    st.session_state[editor_df_key] = cur_df
                    st.rerun()

                if cc2b.button("清空", key=sk("sel_clear")):
                    cur_df.loc[:, "sel"] = False
                    st.session_state[editor_df_key] = cur_df
                    st.rerun()

                if cc2c.button(" 删除已勾选", key=sk("del_sel")):
                    to_delete = cur_df[(cur_df["sel"] == True) & pd.notna(cur_df["ID"])]["ID"].astype(int).tolist()
                    if not to_delete:
                        st.warning("未勾选任何记录")
                    else:
                        execute_write("DELETE FROM term_ext WHERE id=?", [(i,) for i in to_delete])
                        st.success(f" 已删除 {len(to_delete)} 条")
                        st.rerun()

                with c3:
                    proj_opts = cur.execute("SELECT id, title FROM items ORDER BY id DESC").fetchall()
                    proj_map = {"(不挂接/置空)": None, **{f"#{i} {t}": i for (i, t) in proj_opts}}

                    cc3a, cc3b = st.columns([2, 1])
                    target_proj_label = cc3a.selectbox("批量挂接到项目", list(proj_map.keys()), key=sk("bind_proj_sel"))
                    if cc3b.button("执行挂接", type="primary", key=sk("bind_apply")):
                        if "ID" in edited.columns:
                            to_update = (
                                edited[(edited["sel"] == True) & pd.notna(edited["ID"])]["ID"]
                                .astype(int)
                                .tolist()
                            )
                        else:
                            to_update = []
                        if not to_update:
                            st.warning("未勾选任何记录")
                        else:
                            pid_val = proj_map.get(target_proj_label)
                            q_marks = ",".join("?" for _ in to_update)
                            execute_write(
                                f"UPDATE term_ext SET project_id=? WHERE id IN ({q_marks})",
                                (pid_val, *to_update),
                            )
                            st.success(f" 已挂接 {len(to_update)} 条到项目:{target_proj_label or '(空)'}")
                            st.rerun()

        with st.form(sk0("term_edit")):
            col1, col2, col3 = st.columns(3)
            with col1:
                rid_edit = st.text_input("要编辑的记录 ID(留空则新增)", "", key=sk("rid_edit"))
                source_term = st.text_input("源语言术语(必填)*", key=sk("source_term"))
                target_term = st.text_input("目标语言术语", key=sk("target_term"))
            with col2:
                domain = st.text_input("领域", key=sk("domain"))
                project_id = st.text_input("项目ID(可空)", key=sk("project_id"))
                strategy = st.text_input("策略(直译/意译/转译/音译/省略/增译/规范化…)", key=sk("strategy"))
            with col3:
                example = st.text_area("例句", height=80, key=sk("example"))
                category = st.text_input("分类(可选)", key=sk("category"))

            b1, b2 = st.columns(2)
            add = b1.form_submit_button("保存(新增或更新)")
            delbtn = b2.form_submit_button("删除(按 ID)")

        if add:
            if not source_term.strip():
                st.error("源术语必填")
            else:
                if rid_edit and rid_edit.isdigit():
                    if _has_col("term_ext", "category"):
                        execute_write("""
                        UPDATE term_ext
                        SET source_term=?, target_term=?, domain=?, project_id=?, strategy=?, example=?, category=?
                        WHERE id=?;
                        """, (source_term.strip(), target_term.strip() or None, domain or None,
                            int(project_id) if project_id.isdigit() else None, strategy or None, example or None,
                            category or None, int(rid_edit)))
                    else:
                        execute_write("""
                        UPDATE term_ext
                        SET source_term=?, target_term=?, domain=?, project_id=?, strategy=?, example=?
                        WHERE id=?;
                        """, (source_term.strip(), target_term.strip() or None, domain or None,
                            int(project_id) if project_id.isdigit() else None, strategy or None, example or None,
                            int(rid_edit)))
                    st.success(" 已更新")
                    st.rerun()
                else:
                    if _has_col("term_ext", "category"):
                        execute_write("""
                        INSERT INTO term_ext (
                            source_term, target_term, domain, project_id,
                            strategy, example, category
                        )
                        VALUES (?, ?, ?, ?, ?, ?, ?)
                        """, (
                            source_term.strip(),
                            (target_term.strip() or None) if target_term else None,
                            (domain or None),
                            (int(project_id) if project_id.isdigit() else None) if project_id else None,
                            (strategy or None),
                            (example or None),
                            (category or None)
                        ))
                    else:
                        execute_write("""
                        INSERT INTO term_ext (
                            source_term, target_term, domain, project_id,
                            strategy, example
                        )
                        VALUES (?, ?, ?, ?, ?, ?)
                        """, (
                            source_term.strip(),
                            (target_term.strip() or None) if target_term else None,
                            (domain or None),
                            (int(project_id) if project_id.isdigit() else None) if project_id else None,
                            (strategy or None),
                            (example or None),
                        ))
                    st.success(" 已新增")
                    st.rerun()

        if delbtn:
            if rid_edit and rid_edit.isdigit():
                execute_write("DELETE FROM term_ext WHERE id=?", (int(rid_edit),))
                st.success(" 已删除")
                st.rerun()
            else:
                st.error("请填写要删除的 ID")

    # —— 批量导入 CSV(增强版:列名规范化 / 动态带或不带 category / 去重或Upsert / 逐行容错)
    with sub_tabs[1]:
        sk1 = lambda n: f"{key_prefix}_t1_{n}"
        st.caption("CSV 推荐列:source_term, target_term, domain, project_id, strategy, example(可选:category / 分类)")
        up = st.file_uploader("上传 CSV 文件", type=["csv"], key=sk1("csv"))

        def _ensure_col(table, col, type_):
            try:
                execute_write(f"ALTER TABLE {table} ADD COLUMN {col} {type_};")
            except Exception as e:
                log_exception(f"term_ext CSV import: add column {col} on {table} failed", e, level="INFO")

        def _ensure_unique_index():
            try:
                execute_write("""
                CREATE UNIQUE INDEX IF NOT EXISTS ux_term_proj
                ON term_ext(LOWER(TRIM(source_term)), IFNULL(project_id,-1));
                """)
            except Exception as e:
                log_exception("term_ext CSV import: ensure unique index failed", e, level="INFO")

        def _norm_cols(df):
            # 列名:去BOM/两端空格 -> 小写 -> 替换空格为下划线 -> 中英列名映射
            df.columns = [str(c).replace("\ufeff","").strip() for c in df.columns]
            df.columns = [c.lower().replace(" ", "_") for c in df.columns]
            mapping = {
                "源术语": "source_term", "source": "source_term",
                "目标术语": "target_term", "target": "target_term",
                "领域": "domain",
                "项目id": "project_id", "项目_id": "project_id",
                "策略": "strategy",
                "例句": "example",
                "分类": "category",
            }
            df = df.rename(columns={k.lower(): v for k, v in mapping.items()})
            return df

        def _norm(v):
            if v is None: return None
            v = str(v).strip().replace("\u3000", " ")
            return v if v else None

        def _to_int(v):
            try:
                return int(v) if v is not None and str(v).strip() != "" else None
            except Exception as e:
                log_exception("_to_int: failed to convert value to int", e, level="DEBUG")
                return None

        if up is not None:
            try:
                df_up = pd.read_csv(up, encoding="utf-8-sig")
            except Exception as e:
                log_exception("term_ext CSV import: utf-8-sig decode failed, trying utf-8", e, level="DEBUG")
                df_up = pd.read_csv(up, encoding="utf-8", errors="ignore")

            df_up = _norm_cols(df_up)
            st.write("检测到列:", list(df_up.columns))
            render_table(df_up.head(10), key=sk("csv_preview"))

            # DB 侧确保有 category 列(兼容老库;若已有会忽略)
            _ensure_col("term_ext", "category", "TEXT")

            # 侧边选项
            c1, c2, c3 = st.columns(3)
            with c1:
                dedup = st.checkbox("去重(源术语+项目ID)", value=True, key=sk1("dedup"))
            with c2:
                use_upsert = st.checkbox("已存在则更新(Upsert)", value=False, key=sk1("upsert"))
            with c3:
                skip_empty = st.checkbox("跳过空译文", value=False, key=sk1("skip_empty"))

            # Upsert 需要唯一索引
            if use_upsert:
                _ensure_unique_index()

            if st.button("导入术语库", key=sk1("import_btn")):
                # 是否包含 category 列(以CSV为准;DB已有不强制CSV必须有)
                has_category_col = ("category" in df_up.columns)

                # 去重缓存(仅在非Upsert模式下使用)
                existing = set()
                if dedup and not use_upsert:
                    rows_exist = execute_write("""
                        SELECT LOWER(TRIM(source_term)), IFNULL(project_id,-1)
                        FROM term_ext
                    """).fetchall()
                    existing = set(rows_exist)

                ins = skp = upd = 0
                errors = []

                for idx, row in df_up.iterrows():
                    src = _norm(row.get("source_term"))
                    if not src:
                        skp += 1
                        continue

                    tgt = _norm(row.get("target_term"))
                    if skip_empty and not tgt:
                        skp += 1
                        continue

                    dom = _norm(row.get("domain"))
                    pid = _to_int(row.get("project_id"))
                    stg = _norm(row.get("strategy"))
                    exa = _norm(row.get("example"))
                    cat = _norm(row.get("category")) if has_category_col else None

                    try:
                        if use_upsert:
                            # Upsert 分支(需要唯一索引)
                            if has_category_col:
                                execute_write("""
                                INSERT INTO term_ext (
                                    source_term, target_term, domain, project_id,
                                    strategy, example, category
                                )
                                VALUES (?,?,?,?,?,?,?)
                                ON CONFLICT(LOWER(TRIM(source_term)), IFNULL(project_id,-1))
                                DO UPDATE SET
                                    target_term=COALESCE(excluded.target_term, term_ext.target_term),
                                    domain     =COALESCE(excluded.domain,      term_ext.domain),
                                    strategy   =COALESCE(excluded.strategy,    term_ext.strategy),
                                    example    =COALESCE(excluded.example,     term_ext.example),
                                    category   =COALESCE(excluded.category,    term_ext.category);
                                """, (src, tgt, dom, pid, stg, exa, cat))
                            else:
                                execute_write("""
                                INSERT INTO term_ext (
                                    source_term, target_term, domain, project_id,
                                    strategy, example
                                )
                                VALUES (?,?,?,?,?,?)
                                ON CONFLICT(LOWER(TRIM(source_term)), IFNULL(project_id,-1))
                                DO UPDATE SET
                                    target_term=COALESCE(excluded.target_term, term_ext.target_term),
                                    domain     =COALESCE(excluded.domain,      term_ext.domain),
                                    strategy   =COALESCE(excluded.strategy,    term_ext.strategy),
                                    example    =COALESCE(excluded.example,     term_ext.example);
                                """, (src, tgt, dom, pid, stg, exa))
                            upd += 1  # 计为“处理成功”.不区分新旧
                        else:
                            # 非 Upsert:去重(源术语+项目ID)
                            key = (src.lower(), pid if pid is not None else -1)
                            if dedup and key in existing:
                                skp += 1
                                continue

                            if has_category_col:
                                execute_write("""
                                    INSERT INTO term_ext (
                                        source_term, target_term, domain, project_id,
                                        strategy, example, category
                                    )
                                    VALUES (?,?,?,?,?,?,?)
                                """, (src, tgt, dom, pid, stg, exa, cat))
                            else:
                                execute_write("""
                                    INSERT INTO term_ext (
                                        source_term, target_term, domain, project_id,
                                        strategy, example
                                    )
                                    VALUES (?,?,?,?,?,?)
                                """, (src, tgt, dom, pid, stg, exa))
                            ins += 1
                            if dedup:
                                existing.add(key)

                    except Exception as e:
                        errors.append((idx+1, src, str(e)))
                        skp += 1
                        continue

                # 结果提示
                if use_upsert:
                    st.success(f" 已处理 {ins+upd} 条(其中可能含新增+更新).跳过 {skp} 条。")
                else:
                    st.success(f" 新增 {ins} 条.跳过 {skp} 条。")

                if errors:
                    with st.expander(" 行级错误明细(不影响其他行写入)", expanded=False):
                        for i, s, e in errors:
                            st.write(f"第 {i} 行({s}):{e}")

    # —— 统计与导出
    with sub_tabs[2]:
        sk2 = lambda n: f"{key_prefix}_t2_{n}"
        st.markdown("### 翻译流程")
        df_stats = pd.read_sql_query(
            "SELECT strategy, domain, category FROM term_ext WHERE source_term IS NOT NULL",
            RAW_CONN,
        )
        if df_stats.empty:
            st.info("术语库为空.请先添加或导入")
        else:
            # 统一预处理：空值 → 未标注
            df_stats["strategy"] = df_stats["strategy"].fillna("未标注").replace("", "未标注")
            df_stats["domain"]   = df_stats["domain"].fillna("未标注").replace("", "未标注")

            # 选择统计维度
            dim_label = st.selectbox(
                "统计维度",
                ["按领域 (domain)", "按翻译策略 (strategy)"],
                index=0,
                key=sk2("dim_sel"),
            )

            if "领域" in dim_label:
                dim_col = "domain"
                dim_title = "领域"
            else:
                dim_col = "strategy"
                dim_title = "翻译策略"

            # 选择展示方式
            chart_type = st.radio(
                "展示方式",
                ["柱状图", "饼图", "数据表"],
                index=0,
                horizontal=True,
                key=sk2("chart_type"),
            )

            # 做计数
            count_df = (
                df_stats.groupby(dim_col)[dim_col]
                .count()
                .reset_index(name="term_count")
                .sort_values("term_count", ascending=False)
            )

            # ===== 不同展示方式 =====
            if chart_type == "柱状图":
                st.markdown(f"**{dim_title} 分布（柱状图）**")

                chart = (
                    alt.Chart(count_df)
                    .mark_bar()
                    .encode(
                        x=alt.X("term_count:Q", title="术语数量"),
                        y=alt.Y(f"{dim_col}:N", sort="-x", title=dim_title),
                        tooltip=[dim_col, "term_count"],
                    )
                    .properties(height=320)
                )
                st.altair_chart(chart, width='stretch')

            elif chart_type == "饼图":
                st.markdown(f"**{dim_title} 分布（饼图）**")

                chart = (
                    alt.Chart(count_df)
                    .mark_arc()
                    .encode(
                        theta=alt.Theta("term_count:Q", title="术语数量"),
                        color=alt.Color(f"{dim_col}:N", title=dim_title),
                        tooltip=[dim_col, "term_count"],
                    )
                    .properties(height=320)
                )
                st.altair_chart(chart, width='stretch')

            else:  # 数据表
                st.markdown(f"**{dim_title} 分布（数据表）**")
                tbl = count_df.rename(
                    columns={
                        dim_col: dim_title,
                        "term_count": "术语数量",
                    }
                )
                render_table(tbl, hide_index=True, key=sk2("tbl"))

        st.markdown("---")

        target_pid = st.session_state.get(CORPUS_TARGET_PROJECT)
        st.markdown("### 翻译流程")
        df_exp = pd.read_sql_query("""
            SELECT source_term AS '源术语',
                   target_term AS '目标术语',
                   domain AS '领域',
                   strategy AS '翻译策略',
                   example AS '示例句',
                   category AS '分类'
            FROM term_ext ORDER BY source_term COLLATE NOCASE
        """, RAW_CONN)

        buff = io.BytesIO()
        with pd.ExcelWriter(buff, engine="xlsxwriter") as writer:
            df_exp.to_excel(writer, index=False, sheet_name="术语库")
        st.download_button(" 下载 Excel",
                           buff.getvalue(),
                           file_name="terms.xlsx",
                           mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                           key=sk2("dl_terms"))

    # —— 快速搜索
    with sub_tabs[3]:
        sk3 = lambda n: f"{key_prefix}_t3_{n}"
        q = st.text_input("快速搜索(前缀/子串)", "", key=sk3("q"))
        limit = st.number_input("返回上限", 1, 5000, 1000, 100, key=sk3("limit"))
        if st.button("搜索", key=sk3("q_btn")):
            if q:
                like = f"%{q}%"
                rows = execute_write("""
                    SELECT id, source_term, target_term, domain, project_id
                    FROM term_ext
                    WHERE source_term LIKE ? OR target_term LIKE ?
                    ORDER BY id DESC
                    LIMIT ?
                """, (like, like, int(limit))).fetchall()
                render_table(pd.DataFrame(rows, columns=["ID","源术语","目标术语","领域","项目"]),
                             key=sk3("q_grid"),editable=True)
            else:
                st.warning("请输入关键词")

    # —— 批量挂接项目
    with sub_tabs[4]:
        sk4 = lambda n: f"{key_prefix}_t4_{n}"
        st.caption("将一批术语统一设置 project_id.便于项目内优先匹配")
        ids_txt = st.text_area("术语ID列表(逗号/空格/换行分隔)", key=sk4("ids"))
        pid_to = st.text_input("目标项目ID", key=sk4("pid_to"))
        if st.button("批量挂接", key=sk4("batch_btn")):
            import re
            if not pid_to.isdigit():
                st.error("项目ID需为数字")
            else:
                raw = re.split(r"[,\s]+", ids_txt.strip())
                ids = [int(x) for x in raw if x.isdigit()]
                if not ids:
                    st.warning("未识别到有效ID")
                else:
                    qmarks = ",".join(["?"]*len(ids))
                    execute_write(f"UPDATE term_ext SET project_id=? WHERE id IN ({qmarks})", (int(pid_to), *ids))
                    st.success(f" 已挂接 {len(ids)} 条到项目 {pid_to}")

    # —— 从历史提取术语
    with sub_tabs[5]:
        sk5 = lambda n: f"{key_prefix}_t5_{n}"
        st.markdown("### 翻译流程")
        ak, model = get_deepseek()
        if not ak:
            st.warning("未检测到 DeepSeek Key.请先在“设置”中配置。")
        else:
            mode_pick = st.radio(
                "选择来源",
                ["按项目抽取(合并多条)", "按单条记录抽取"],
                horizontal=True,
                key=sk5("ext_mode"),
            )
            if mode_pick == "按项目抽取(合并多条)":
                pid_ext = st.text_input("项目ID", key=sk5("ext_pid"))
                max_chars = st.number_input("采样最大字符数", 1000, 20000, 5000, 500, key=sk5("ext_max"))
                if st.button("开始抽取", key=sk5("ext_go_proj")):
                    if pid_ext.isdigit():
                        rows = execute_write(
                            "SELECT src_path, output_text FROM trans_ext WHERE project_id=? ORDER BY id DESC LIMIT 10",
                            (int(pid_ext),),
                        ).fetchall()
                        parts = []
                        total = 0
                        for sp, ot in rows:
                            src = read_source_file(sp) if sp else ""
                            txt = (src + "\n" + (ot or "")).strip()
                            if not txt:
                                continue
                            if total + len(txt) > int(max_chars):
                                remain = max(0, int(max_chars) - total)
                                parts.append(txt[:remain])
                                break
                            else:
                                parts.append(txt)
                                total += len(txt)
                        big = "\n\n".join(parts)

                        if not big.strip():
                            st.warning("该项目下没有可用的翻译历史文本，无法抽取术语。")
                            return

                        st.write({
                            "history_rows": len(rows),
                            "sample_chars": len(big),
                            "sample_preview": big[:300]
                        })

                        # P1-1: 使用 safe_api_call 消除重复的异常处理
                        res = safe_api_call(
                            ds_extract_terms,
                            big, ak, model,
                            context="术语抽取(项目模式)",
                            fallback=[],
                            src_lang="zh", tgt_lang="en"
                        )

                        st.write({"extract_result_preview": str(res)[:500]})
                        if not res:
                            st.info("未抽取到术语或解析失败")
                        else:
                            st.success(f"抽取到 {len(res)} 条.准备批量写入……")
                            ins = 0
                            for o in res:
                                execute_write(
                                    "INSERT INTO term_ext (source_term, target_term, domain, project_id, "
                                    "strategy, example) "
                                    "VALUES (?, ?, ?, ?, ?, ?)",
                                    (
                                        o["source_term"],
                                        o.get("target_term") or None,
                                        o.get("domain"),
                                        int(pid_ext),
                                        o.get("strategy"),
                                        o.get("example"),
                                    ),
                                )
                                ins += 1
                            st.success(f" 已写入术语库 {ins} 条")
                    else:
                        st.warning("请输入数字型项目ID")
            else:
                rid_ext = st.text_input("历史记录ID", key=sk5("ext_rid"))
                if st.button("开始抽取", key=sk5("ext_go_rec")):
                    if rid_ext.isdigit():
                        row = execute_write(
                            "SELECT src_path, output_text, project_id FROM trans_ext WHERE id=?",
                            (int(rid_ext),),
                        ).fetchone()
                        if not row:
                            st.warning("未找到该记录")
                        else:
                            sp, ot, pid0 = row
                            src = read_source_file(sp) if sp else ""
                            big = (src + "\n" + (ot or "")).strip()
                            
                            # P1-1: 使用 safe_api_call 消除重复的异常处理
                            res = safe_api_call(
                                ds_extract_terms,
                                big, ak, model,
                                context="术语抽取(单条记录)",
                                fallback=[],
                                src_lang="zh", tgt_lang="en"
                            )
                            if not res:
                                st.info("未抽取到术语或解析失败")
                            else:
                                st.success(f"抽取到 {len(res)} 条.准备批量写入……")
                                ins = 0
                                for o in res:
                                    execute_write(
                                        "INSERT INTO term_ext (source_term, target_term, domain, project_id, "
                                        "strategy, example) "
                                        "VALUES (?, ?, ?, ?, ?, ?)",
                                        (
                                            o["source_term"],
                                            o.get("target_term") or None,
                                            o.get("domain"),
                                            pid0,
                                            o.get("strategy"),
                                            o.get("example"),
                                        ),
                                    )
                                    ins += 1
                                st.success(f" 已写入术语库 {ins} 条(project_id={pid0})")

    # —— 分类管理
    with sub_tabs[6]:
        sk6 = lambda n: f"{key_prefix}_t6_{n}"
        st.markdown("### 翻译流程")
        c1, c2 = st.columns(2)
        with c1:
            ids_txt = st.text_area("按 ID 批量设置分类(逗号/空格/换行分隔)", key=sk6("cat_ids"))
            cat_to = st.text_input("要设置的分类名", key=sk6("cat_name"))
            if st.button("批量设置分类", key=sk6("cat_set_ids")):
                raw = re.split(r"[,\s]+", (ids_txt or "").strip())
                ids = [int(x) for x in raw if x.isdigit()]
                if not ids or not cat_to.strip():
                    st.warning("请填入ID列表与分类名称")
                else:
                    qmarks = ",".join(["?"] * len(ids))
                    execute_write(f"UPDATE term_ext SET category=? WHERE id IN ({qmarks})", (cat_to.strip(), *ids))
                    st.success(f" 已设置 {len(ids)} 条为分类:{cat_to.strip()}")

        with c2:
            pid_cat = st.text_input("将某项目ID全部术语设置为分类", key=sk6("cat_pid"))
            cat2_to = st.text_input("分类名", key=sk6("cat2_name"))
            if st.button("按项目ID统一分类", key=sk6("cat_set_pid")):
                if pid_cat.isdigit() and cat2_to.strip():
                    execute_write("UPDATE term_ext SET category=? WHERE project_id=?", (cat2_to.strip(), int(pid_cat)))
                    st.success(f" 已将项目 {pid_cat} 的术语分类设为:{cat2_to.strip()}")
                else:
                    st.warning("请填写项目ID与分类名")

# ========== Session 初始化 ==========
if KB_EMBEDDER not in st.session_state and KBEmbedder:
    st.session_state[KB_EMBEDDER] = KBEmbedder(lazy=True)

def render_index_manager_by_domain(st, conn, cur):
    st.subheader(" 领域级索引管理 (双语对照 + 翻译策略)")

    # 1) 收集所有可能的领域值
    domains = set()

    # from items
    try:
        rows = cur.execute("SELECT DISTINCT IFNULL(domain,'未分类') FROM items").fetchall()
        for (d,) in rows:
            if d and d.strip():
                domains.add(d.strip())
            else:
                domains.add("未分类")
    except Exception as e:
        log_exception("render_index_manager_by_domain: fetch domains from items failed", e)

    # from corpus.domain (如果有该字段)
    try:
        cols = [r[1] for r in cur.execute("PRAGMA table_info(corpus)").fetchall()]
        if "domain" in cols:
            rows = cur.execute("SELECT DISTINCT IFNULL(domain,'未分类') FROM corpus").fetchall()
            for (d,) in rows:
                if d and d.strip():
                    domains.add(d.strip())
                else:
                    domains.add("未分类")
    except Exception as e:
        log_exception("render_index_manager_by_domain: fetch domains from corpus failed", e, level="WARNING")

    # from strategy_texts(如果已经存在)
    try:
        execute_write(
            "CREATE TABLE IF NOT EXISTS strategy_texts ("
            "id INTEGER PRIMARY KEY,"
            "domain TEXT,"
            "title TEXT,"
            "content TEXT NOT NULL,"
            "collection TEXT,"
            "source TEXT,"
            "created_at TEXT DEFAULT (datetime('now'))"
            ");"
        )
        rows = cur.execute("SELECT DISTINCT IFNULL(domain,'未分类') FROM strategy_texts").fetchall()
        for (d,) in rows:
            if d and d.strip():
                domains.add(d.strip())
            else:
                domains.add("未分类")
    except Exception as e:
        log_exception("render_index_manager_by_domain: fetch domains from strategy_texts failed", e, level="WARNING")

    if not domains:
        st.info("当前尚未设置任何领域(domain)。请先在项目或语料中设置领域。")
        return

    domains_list = sorted(domains)
    dom_sel = st.selectbox("选择要管理的领域", domains_list)
    dom_key = (dom_sel or "").strip() or "未分类"

    st.markdown(f"### 当前领域: `{dom_key}`")

    # 2) 统计该领域下的项目 & 语料 & 索引情况
    # 2.1 项目列表
    proj_rows = execute_write(
        "SELECT id, title FROM items WHERE IFNULL(domain,'未分类') = ? ORDER BY id ASC",
        (dom_key,)
    ).fetchall()
    proj_ids = [pid for (pid, _) in proj_rows]

    # 2.2 语料条数(如果 corpus 有 domain 字段)
    corpus_cnt = None
    try:
        cols = [r[1] for r in cur.execute("PRAGMA table_info(corpus)").fetchall()]
        if "domain" in cols:
            corpus_cnt = execute_write(
                "SELECT COUNT(*) FROM corpus WHERE IFNULL(domain,'未分类') = ?",
                (dom_key,)
            ).fetchone()[0]
    except Exception:
        corpus_cnt = None

    # 2.3 双语索引条数 = 该领域所有项目索引的 mapping 长度之和
    idx_bilingual_total = 0
    for pid in proj_ids:
        try:
            mode, index, mapping, vecs = _load_index(int(pid))
        except Exception:
            continue
        if isinstance(mapping, list):
            idx_bilingual_total += len(mapping)

    # 2.4 策略文本数量 & 策略索引条数
    try:
        strategy_cnt = execute_write(
            "SELECT COUNT(*) FROM strategy_texts WHERE IFNULL(domain,'未分类') = ?",
            (dom_key,)
        ).fetchone()[0]
    except Exception:
        strategy_cnt = 0

    mode_s, index_s, mapping_s, vecs_s = _load_index_domain(dom_key, "strategy")
    if isinstance(mapping_s, list):
        idx_strategy_total = len(mapping_s)
    else:
        idx_strategy_total = 0

    c1, c2 = st.columns(2)
    with c1:
        st.markdown("### 翻译流程")
        st.write(f"- 该领域下项目数: **{len(proj_ids)}**")
        if corpus_cnt is not None:
            st.write(f"- 语料库中双语条目(按 domain 计): **{corpus_cnt}**")
        st.write(f"- 已建立索引的句对条数(合计): **{idx_bilingual_total}**")

        if proj_ids:
            if st.button(" 重建该领域所有项目的【双语对照】索引", key=f"rebuild_bi_{dom_key}"):
                added_sum = 0
                total_sum = 0
                for pid in proj_ids:
                    try:
                        res = build_project_vector_index(int(pid))
                        added_sum += res.get("added", 0)
                        total_sum = res.get("total", total_sum)
                    except Exception as e:
                        st.warning(f"项目 {pid} 重建索引时出错: {e}")
                st.success(
                    f"已重建该领域所有项目索引。"
                    f"新增句对: {added_sum}，最后一个项目返回的索引总量: {total_sum}"
                )

    with c2:
        st.markdown("### 翻译流程")
        st.write(f"- 该领域下策略文本条数: **{strategy_cnt}**")
        st.write(f"- 已建立索引的策略向量条数: **{idx_strategy_total}**")

        if st.button(" 重建该领域的【翻译策略】索引", key=f"rebuild_strategy_{dom_key}"):
            try:
                res = build_strategy_index_for_domain(dom_key)
                st.success(
                    f"已重建策略索引。新增策略段落: {res.get('added', 0)}，"
                    f"索引总量: {res.get('total', 0)}"
                )
            except Exception as e:
                st.error(f"重建策略索引时出错: {e}")

    st.markdown("---")
    with st.expander(" 查看该领域下的项目列表", expanded=False):
        if proj_rows:
            for pid, title in proj_rows:
                st.write(f"- 项目 {pid}: {title}")
        else:
            st.write("暂无项目。")

    with st.expander(" 查看该领域下的策略文本(前几条)", expanded=False):
        try:
            rows = execute_write(
                "SELECT id, title, substr(content,1,200) FROM strategy_texts "
                "WHERE IFNULL(domain,'未分类') = ? ORDER BY id DESC LIMIT 20",
                (dom_key,)
            ).fetchall()
            if not rows:
                st.write("暂无策略文本。")
            else:
                for sid, ttl, preview in rows:
                    st.write(f"**[{sid}] {ttl or '(无标题)'}**")
                    st.write(preview + ("..." if len(preview) >= 200 else ""))
                    st.markdown("---")
        except Exception as e:
            st.write(f"读取策略文本出错: {e}")

# ========== 页面结构 ==========
st.sidebar.title("导航")
nav_options = [
    " 翻译项目管理",
    " 术语库管理",
    " 翻译历史",
    " 语料库管理",
    " 索引管理",
]
choice = st.sidebar.radio(
    "功能选择",
    nav_options,
    key="MAIN_NAV_CHOICE",
)
choice_idx = nav_options.index(choice)

st.title("齐译知识增强型翻译支持与管理系统1.0")

# ========== Tab1:翻译项目管理 ==========
if choice_idx == 0:
    st.subheader("翻译项目管理")
    if SCOPE_NEWPROJ not in st.session_state:
        st.session_state[SCOPE_NEWPROJ] = "project"
    if USE_SEMANTIC_GLOBAL not in st.session_state:
        st.session_state[USE_SEMANTIC_GLOBAL] = True
    pending_tab = st.session_state.pop(TRANS_MGMT_TAB_PENDING, None)
    if pending_tab:
        st.session_state[TRANS_MGMT_TAB] = pending_tab
    trans_tab = st.radio(
        '项目管理导航',
        ['创建项目', '项目列表', '翻译流程'],
        horizontal=True,
        key=TRANS_MGMT_TAB,
        label_visibility='collapsed',
    )
    if trans_tab == '创建项目':
        with st.form('new_project'):
            TAG_OPTIONS = ['政治', '经济', '文化', '文物', '金融', '法律']
            SCENE_OPTIONS = ['学术', '配音稿', '正式会议']

            c1, c2 = st.columns([3, 2])
            with c1:
                title = st.text_input('项目名称')
                tags_sel = st.multiselect('项目标签(可多选)', TAG_OPTIONS)
                scene_sel = st.selectbox('场合', SCENE_OPTIONS, index=0)
            with c2:
                translation_type = st.selectbox('翻译方式', ['使用术语库', '纯机器翻译'])
                translation_mode = st.radio('模式', ['标准模式', '术语约束模式'], horizontal=True)
                prompt_text = st.text_area(
                    '翻译提示(注入模型 System Prompt)',
                    placeholder='写下对 DeepSeek 的硬性/优先级指令.如:时态统一为过去式.专有名词保持原文……',
                    height=120,
                    key='new_proj_prompt'
                )

            domain_val = tags_sel[0] if tags_sel else None

            submitted = st.form_submit_button('创建项目')
            if submitted:
                if not title:
                    st.error('请填写项目名称')
                else:
                    try:
                        execute_write('PRAGMA table_info(items);')
                        cols = [r[1] for r in cur.fetchall()]
                        if 'domain' not in cols:
                            execute_write('ALTER TABLE items ADD COLUMN domain TEXT;')

                        execute_write("INSERT INTO items(title, body, tags, domain, type) VALUES (?, ?, ?, ?, 'project')", (
                            title,
                            prompt_text or '',
                            ','.join(tags_sel or []),
                            domain_val
                        ))

                        st.success(f"项目 '{title}' 已创建(领域:{domain_val or '未指定'})")
                    except Exception as e:
                        st.error(f'创建项目失败: {e}')
    if trans_tab == "翻译流程":
        st.markdown("### 翻译流程")
        use_semantic_global = st.checkbox(
            "在翻译时启用语义召回参考",
            value=st.session_state.get(USE_SEMANTIC_GLOBAL, True),
            key=USE_SEMANTIC_GLOBAL
        )
        scope_options = {
            "仅当前项目": "project",
            "同领域 + 当前项目": "domain",
            "全库": "all",
        }
        scope_default = st.session_state.get(SCOPE_NEWPROJ, "project")
        scope_label_map = {v: k for k, v in scope_options.items()}
        scope_label_default = scope_label_map.get(scope_default, "仅当前项目")
        scope_label = st.selectbox(
            "语义召回范围",
            list(scope_options.keys()),
            index=list(scope_options.keys()).index(scope_label_default),
            key="scope_sel_translate"
        )
        st.session_state[SCOPE_NEWPROJ] = scope_options[scope_label]

        proj_rows = execute_write(
            "SELECT i.id, IFNULL(i.title,''), IFNULL(e.src_path,'') "
            "FROM items i LEFT JOIN item_ext e ON e.item_id=i.id "
            "WHERE COALESCE(i.type,'')='project' ORDER BY i.id DESC"
        ).fetchall()
        if not proj_rows:
            st.info("暂无项目")
        else:
            proj_map = {f"[{pid}] {ttl or '(未命名)'}": pid for pid, ttl, _ in proj_rows}
            labels = list(proj_map.keys())
            default_pid = st.session_state.get(ACTIVE_TRANSLATION_PID)
            default_label = labels[0] if labels else None
            if default_pid is not None:
                for lbl, pid in proj_map.items():
                    if pid == default_pid:
                        default_label = lbl
                        break
            sel_label = st.selectbox(
                "选择项目",
                labels,
                index=labels.index(default_label) if default_label in labels else 0,
                key="trans_proj_sel"
            )
            pid = proj_map.get(sel_label)

            selected_src_path = None
            if pid:
                row = cur.execute("SELECT IFNULL(src_path,'') FROM item_ext WHERE item_id=?", (pid,)).fetchone()
                legacy_path = (row[0] if row else "") or None
                ensure_legacy_file_record(cur, conn, pid, legacy_path)
                file_records = fetch_project_files(cur, pid)
                if file_records:
                    option_labels = []
                    option_map = {}
                    for rec in file_records:
                        label = f"[#{rec['id']}] {rec['name']}"
                        if rec["uploaded_at"]:
                            label += f"｜{rec['uploaded_at']}"
                        option_labels.append(label)
                        option_map[label] = rec
                    sel_key = SessionManager._build_key(SessionManager.TRANS_FILE_SEL, pid=pid)
                    default_path = st.session_state.get(ACTIVE_TRANSLATION_SRC)
                    default_label_file = option_labels[0]
                    if default_path:
                        for lbl in option_labels:
                            if option_map[lbl]["path"] == default_path:
                                default_label_file = lbl
                                break
                    chosen_label = st.selectbox(
                                "选择要翻译的源文件",
                        option_labels,
                        index=option_labels.index(default_label_file),
                        key=sel_key
                    )
                    selected_src_path = option_map[chosen_label]["path"]
                else:
                    selected_src_path = legacy_path
                    st.info("该项目暂无上传文件，请先在项目列表上传。")

            if pid and selected_src_path:
                st.session_state[ACTIVE_TRANSLATION_PID] = pid
                st.session_state[ACTIVE_TRANSLATION_SRC] = selected_src_path
                run_project_translation_ui(
                    pid=pid,
                    project_title=sel_label.split("] ", 1)[-1],
                    src_path=selected_src_path,
                    conn=conn,
                    cur=cur,
                )
            elif pid:
                st.warning("请先选择或上传源文件后再开始翻译。")
    if trans_tab == "项目列表":
        rows = execute_write("""
            SELECT
                i.id,
                i.title,
                COALESCE(i.tags,'')         AS tags,
                COALESCE(e.src_path,'')     AS src_path,
                COALESCE(i.created_at,'')   AS created_at,
                COALESCE(i.scene,'')        AS scene,
                COALESCE(i.prompt,'')       AS prompt,
                COALESCE(i.mode,'')         AS mode,
                COALESCE(i.trans_type,'')   AS trans_type
            FROM items i
            LEFT JOIN item_ext e ON e.item_id = i.id
            WHERE COALESCE(i.type,'')='project'
            ORDER BY i.id DESC
        """).fetchall()
    
        if not rows:
            st.info("暂无项目")
        else:
            # 用于收集本轮勾选的项目(ID + 文件路径)
            batch_to_delete = []
    
            for pid, title, tags_str, path, ct, scene, prompt_ro, mode, trans_type in rows:
                ensure_legacy_file_record(cur, conn, pid, path or None)
                file_records = fetch_project_files(cur, pid)
                tag_display = tags_str or "无"
                file_display = f"{len(file_records)} 个文件" if file_records else "无"
                selected_src_path = None
    
                with st.expander(
                    f"{title}｜方式:{mode or '未设'}｜标签:{tag_display}｜场合:{scene or '未填'}"
                    f"｜文件:{file_display}｜创建:{ct}"
                ):
                    #  批量操作用的勾选框
                    sel = st.checkbox("选择此项目(用于批量删除)", key=f"sel_proj_{pid}")
                    if sel:
                        batch_to_delete.append(pid)
    
                    c1, c2, c3 = st.columns([2, 2, 1])
                    with c1:
                        st.selectbox(
                            "翻译方向",
                            ["中译英", "英译中"],
                            key=SessionManager._build_key(SessionManager.LANG, pid=pid),
                        )
                    with c2:
                        max_len = st.number_input("分块长度", 600, 2000, 1200, 100, key=f"len_{pid}")
                    with c3:
                        use_terms = st.checkbox("使用术语库", value=(mode == "术语约束模式"), key=f"ut_{pid}")
    
                    st.caption(f"标签:{tag_display}")
                    st.caption(f"场合:{scene or '未填写'}")
                 
                    # === 领域(domain)设置:跟随第一个标签 或 手动选择 ===
                    # 读取当前项目的 domain / tags
                    # 保底:items 表若没有 domain 列.动态补列(兼容旧库)
                    cols_items = [r[1] for r in cur.execute("PRAGMA table_info(items)").fetchall()]
                    if "domain" not in cols_items:
                        try:
                            execute_write("ALTER TABLE items ADD COLUMN domain TEXT;")
                        except Exception as e:
                            # 并发或已有列时忽略
                            log_exception("add domain column to items failed (likely already exists)", e, level="INFO")
     
                    row = execute_write(
                        "SELECT IFNULL(domain,''), IFNULL(tags,'') FROM items WHERE id=?",
                        (pid,)
                    ).fetchone()
                    domain0, tags0 = (row or ["", ""])
                    tags_list = [t for t in (tags0.split(",") if tags0 else []) if t]
    
                    DOMAIN_OPTIONS = ["政治", "经济", "文化", "文物", "金融", "法律"]
    
                    dom_mode = st.radio(
                        "领域设置方式",
                        ["跟随第一个标签", "手动选择"],
                        horizontal=True,
                        key=f"dom_mode_{pid}"
                    )
    
                    if dom_mode == "跟随第一个标签":
                        domain_val = (tags_list[0] if tags_list else (domain0 or None))
                        st.caption(f"当前领域(自动):{domain_val or '未指定'}(由第一个标签决定)")
                    else:
                        idx = DOMAIN_OPTIONS.index(domain0) if domain0 in DOMAIN_OPTIONS else 0
                        domain_val = st.selectbox(
                            "领域(手动选择)",
                            DOMAIN_OPTIONS,
                            index=idx,
                            key=f"dom_sel_{pid}"
                        )
    
                    sync_corpus = st.checkbox(
                        "同时回填该项目下语料的领域(仅补空或原领域相同时覆盖)",
                        value=True,
                        key=f"sync_corpus_{pid}"
                    )
    
                    if st.button(" 保存领域设置", key=f"save_dom_{pid}", type="secondary"):
                        try:
                            # 确保 items.domain 存在
                            cols_items = [r[1] for r in cur.execute("PRAGMA table_info(items)").fetchall()]
                            if "domain" not in cols_items:
                                execute_write("ALTER TABLE items ADD COLUMN domain TEXT;")
    
                            # 更新 items.domain
                            execute_write("UPDATE items SET domain=? WHERE id=?", (domain_val, pid))
    
                            # 同步语料库的 domain(优先 corpus_main.退回 corpus)
                            def _table_exists(tb: str) -> bool:
                                return bool(execute_write(
                                    "SELECT name FROM sqlite_master WHERE type='table' AND name=?;", (tb,)
                                ).fetchone())
    
                            corpus_tb = (
                                "corpus_main"
                                if _table_exists("corpus_main")
                                else ("corpus" if _table_exists("corpus") else None)
                            )
                            if sync_corpus and corpus_tb and domain_val:
                                # 确保列存在
                                cols_corpus = [r[1] for r in cur.execute(f"PRAGMA table_info({corpus_tb})").fetchall()]
                                if "domain" not in cols_corpus:
                                    execute_write(f"ALTER TABLE {corpus_tb} ADD COLUMN domain TEXT;")
    
                                # 仅补空或原 domain 与 domain0 相同时覆盖.避免误伤跨领域数据
                                execute_write(f"""
                                    UPDATE {corpus_tb}
                                    SET domain = ?
                                    WHERE project_id = ?
                                    AND (domain IS NULL OR TRIM(domain)='' OR domain = ?)
                                """, (domain_val, pid, domain0 or ""))
    
                            st.success(" 已保存领域设置")
                            st.rerun()
    
                        except Exception as e:
                            st.error(f" 保存失败:{e}")
    
                    if prompt_ro:
                        try:
                            execute_write("SELECT IFNULL(prompt, '') FROM items WHERE id=?", (pid,))
                            prompt_ro = (cur.fetchone() or [""])[0]
                        except Exception:
                            prompt_ro = ""
    
                        st.text_area("翻译提示(只读)", prompt_ro or "", height=120, key=f"proj_prompt_ro_{pid}")
    
                    file_col, action_col = st.columns([3, 1])
    
                    with file_col:
                        if file_records:
                            option_labels = []
                            option_map = {}
                            for rec in file_records:
                                label = f"[#{rec['id']}] {rec['name']}"
                                if rec["uploaded_at"]:
                                    label += f"｜{rec['uploaded_at']}"
                                option_labels.append(label)
                                option_map[label] = rec
                            sel_key = SessionManager._build_key(SessionManager.FILE_SEL, pid=pid)
                            default_label = SessionManager.get(SessionManager.FILE_SEL, pid=pid)
                            if default_label not in option_labels:
                                default_label = option_labels[0]
                                SessionManager.set(SessionManager.FILE_SEL, default_label, pid=pid)
                            chosen_label = st.selectbox(
                                "选择要翻译的源文件",
                                option_labels,
                                index=option_labels.index(default_label),
                                key=sel_key,
                            )
                            selected_src_path = option_map[chosen_label]["path"]
                            st.caption(f"已上传 {len(file_records)} 个附件，当前选中:{option_map[chosen_label]['name']}")
                        else:
                            selected_src_path = path or None
                            st.info("尚未上传源文件，可在下方上传一个或多个文件。")                      
    
                        upload_key = f"up_multi_{pid}"
                        processed_key = f"up_multi_processed_{pid}"
                        if upload_key not in st.session_state:
                            st.session_state[processed_key] = set()
                        uploads = st.file_uploader(
                            "新增/补传文件(可多选)",
                            type=["txt", "docx", "xlsx", "pdf"],
                            accept_multiple_files=True,
                            key=upload_key
                        )
                        if uploads:
                            processed_names = st.session_state.setdefault(processed_key, set())
                            saved = 0
                            failed_files = []
                            for uf in uploads:
                                if not uf or uf.name in processed_names:
                                    continue
                                data = uf.read()
                                if not data:
                                    failed_files.append(f"{uf.name}（文件为空）")
                                    continue
                                
                                new_path = register_project_file(cur, conn, pid, uf.name, data)
                                if new_path:
                                    execute_write("SELECT id FROM item_ext WHERE item_id=?", (pid,))
                                    r = cur.fetchone()
                                    if r:
                                        execute_write("UPDATE item_ext SET src_path=? WHERE id=?", (new_path, r[0]))
                                    else:
                                        execute_write(
                                            "INSERT INTO item_ext (item_id, src_path) VALUES (?, ?)",
                                            (pid, new_path),
                                        )
                                    saved += 1
                                    processed_names.add(uf.name)
                                else:
                                    failed_files.append(f"{uf.name}（保存失败）")
                            
                            if saved:
                                st.success(f"✓ 已上传 {saved} 个文件")
                            if failed_files:
                                st.error(f"✗ 以下文件上传失败，请检查日志:\n" + "\n".join(failed_files))
                        else:
                            st.session_state.pop(processed_key, None)
    
                        if file_records:
                            st.markdown("附件列表:")
                            for rec in file_records:
                                info_cols = st.columns([5, 1])
                                info = f"[#{rec['id']}] {rec['name']}｜{os.path.basename(rec['path'])}"
                                if rec["uploaded_at"]:
                                    info += f"｜{rec['uploaded_at']}"
                                info_cols[0].write(info)
                                if info_cols[1].button("删除", key=f"del_file_{rec['id']}"):
                                    result = remove_project_file(cur, conn, rec["id"])
                                    if result:
                                        st.success("文件已删除")
                                        st.rerun()
                                    else:
                                        st.error("文件删除失败，请检查文件是否被占用或重试")
                                        st.rerun()
    
                    with action_col:
                        if st.button("删除项目", key=f"del_proj_{pid}"):
                            cleanup_project_files(cur, conn, pid)
                            execute_write("DELETE FROM items WHERE id=?", (pid,))
                            execute_write("DELETE FROM item_ext WHERE item_id=?", (pid,))
                            st.success("项目已删除")
                            st.rerun()
    
                        if st.button("开始翻译", key=f"run_{pid}", type="primary"):
                            st.session_state[ACTIVE_TRANSLATION_PID] = pid
                            st.session_state[ACTIVE_TRANSLATION_SRC] = selected_src_path
                            st.session_state[TRANS_MGMT_TAB_PENDING] = "翻译流程"
                            st.rerun()
    
                    # —— 新增：进入翻译工作台
                    if st.button("进入翻译工作台", key=f"workspace_{pid}", type="secondary"):
                        # 1) 环境检查：有源文件吗？
                        if not selected_src_path or not os.path.exists(selected_src_path):
                            st.error("缺少源文件，请先在上面选择或上传源文件。")
                            st.stop()
                        SessionManager.set(SessionManager.WORKSPACE_ACTIVATED, True, pid=pid)
    
                        # 2) 读取源文本（这里定义 src_text）
                        src_text = read_source_file(selected_src_path)

                        # 3) 分段
                        blocks = split_paragraphs(src_text)
                        if not blocks:
                            # 若未切出段落但文本非空，退化为整篇作为单段，避免误判
                            if src_text and src_text.strip():
                                blocks = [src_text.strip()]
                            else:
                                st.error("源文件内容为空，或未识别到有效段落")
                                st.stop()
    
                        # 4) 术语：用统一接口 + 转成 term_pairs 供高亮用
                        term_map_all, term_meta = get_terms_for_project(cur, pid, use_dynamic=True)
                        term_pairs = list(term_map_all.items())
    
                        # 5) 用统一管线 translate_block_with_kb 做初译
                        ak, model = get_deepseek()
                        if not ak:
                            st.error("未检测到 DeepSeek Key，请配置 deepseek")
                            st.stop()
    
                        # 翻译方向（沿用项目里已有的变量）
                        lang_pair_val = SessionManager.get(SessionManager.LANG, default="中译英", pid=pid)
    
                        # 是否启用语义召回 / 召回范围，直接沿用上面 Tab1 的设置
                        use_semantic_val = bool(st.session_state.get(USE_SEMANTIC_GLOBAL, True))
                        scope_val_local = st.session_state.get(SCOPE_NEWPROJ, "project")
                        locked_terms_ws = SessionManager.get(SessionManager.TERM_FINAL_MAP, default={}, pid=pid)
    
                        draft = []
                        for blk in blocks:
                            blk = (blk or "").strip()
                            if not blk:
                                draft.append("")
                                continue
    
                            res = translate_block_with_kb(
                                cur=cur,
                                project_id=pid,
                                block_text=blk,
                                lang_pair=lang_pair_val,
                                ak=ak,
                                model=model,
                                use_semantic=use_semantic_val,
                                scope=scope_val_local,
                                fewshot_examples=None,  # 工作台模式先不注入 few-shot
                                locked_term_map=locked_terms_ws,
                            )
                            draft.append(res["tgt"])
    
                        # 6) 保存到 session_state，供下面编辑界面使用
                        SessionManager.set(SessionManager.WORKSPACE_SRC, blocks, pid=pid)
                        SessionManager.set(SessionManager.WORKSPACE_DRAFT, draft, pid=pid)
                        SessionManager.set(SessionManager.WORKSPACE_TERMS, term_pairs, pid=pid)
    
                        st.success("草稿已生成，请下方开始编辑 ↓")
    
                    # ③ 翻译工作台 UI：只有当 session 里有草稿时才显示
                    if (
                        SessionManager.get(SessionManager.WORKSPACE_DRAFT, pid=pid)
                        and SessionManager.get(SessionManager.WORKSPACE_ACTIVATED, pid=pid, default=False)
                    ):
    
                        st.markdown("##  翻译工作台（可编辑）")
    
                        # 从 session 中取回草稿和术语
                        blocks = SessionManager.get(SessionManager.WORKSPACE_SRC, default=[], pid=pid)
                        draft  = SessionManager.get(SessionManager.WORKSPACE_DRAFT, default=[], pid=pid)
                        terms  = SessionManager.get(SessionManager.WORKSPACE_TERMS, default=[], pid=pid)
    
                        if not blocks or not draft:
                            st.info("当前暂无草稿，请先点击“进入翻译工作台（可编辑）”生成初稿。")
                        else:
                            edited_blocks = []
    
                            for i, (src, trg) in enumerate(zip(blocks, draft), 1):
                                st.markdown(f"### 段落 {i}")
    
                                col1, col2 = st.columns(2)
    
                                with col1:
                                    st.markdown("**原文**")
                                    st.markdown(
                                        f"<div style='padding:8px;border:1px solid #ccc;"
                                        f"background:#f8f8f8'>{src}</div>",
                                        unsafe_allow_html=True
                                    )
    
                                with col2:
                                    st.markdown("**译文（可编辑）**")
                                    new_trg = st.text_area(
                                        label="编辑后的译文",
                                        value=trg,
                                        key=f"edit_{pid}_{i}",
                                        height=120
                                    )
                                    edited_blocks.append(new_trg)
    
                                    if "highlight_terms" in globals():
                                        highlighted = highlight_terms(new_trg, terms)
                                        st.markdown("术语高亮：")
                                        st.markdown(
                                            f"<div style='padding:8px;border:1px solid #ccc;"
                                            f"background:#f0fff0'>{highlighted}</div>",
                                            unsafe_allow_html=True
                                        )
    
                            # —— 确认生成最终译文 —— 
                            if st.button(" 确认生成最终译文", key=f"confirm_{pid}", type="primary"):
                                final_text = "\n\n".join(edited_blocks)
    
                                lang_pair_val = SessionManager.get(SessionManager.LANG, default="中译英", pid=pid)
    
                                execute_write("""
                                    INSERT INTO trans_ext (
                                        project_id, src_path, lang_pair, mode, output_text, created_at
                                    )
                                    VALUES (?, ?, ?, ?, ?, datetime('now'))
                                """, (pid, selected_src_path, lang_pair_val, "工作台模式", final_text))
    
                                st.success("最终译文已生成并写入历史！")
    
                                # 清空工作台草稿
                                SessionManager.pop(SessionManager.WORKSPACE_SRC, None, pid=pid)
                                SessionManager.pop(SessionManager.WORKSPACE_DRAFT, None, pid=pid)
                                SessionManager.pop(SessionManager.WORKSPACE_TERMS, None, pid=pid)
     
            # —— 批量删除按钮(在项目列表底部)
            if batch_to_delete:
                st.warning(f"已勾选 {len(batch_to_delete)} 个项目，操作不可撤销。")
                if st.button(" 批量删除选中项目", key="batch_del_projects"):
                    deleted = 0
                    for pid_del in batch_to_delete:
                        cleanup_project_files(cur, conn, pid_del)
                        execute_write("DELETE FROM items WHERE id=?", (pid_del,))
                        execute_write("DELETE FROM item_ext WHERE item_id=?", (pid_del,))
                        deleted += 1
                    st.success(f"已批量删除 {deleted} 个项目")
                    st.rerun()
            else:
                st.caption("提示:如需批量删除，可在上方勾选多个项目。")
    
    # ========== Tab2:术语库管理 ==========
elif choice_idx == 1:
    render_term_management(st, cur, conn, BASE_DIR, key_prefix="term")

# ========== Tab3:翻译历史(增强版) ==========
elif choice_idx == 2:
    st.subheader(" 翻译历史记录(可写入语料 / 抽取术语 / 下载对照 / 删除)")

    try:
        rows = cur.execute(
            """
            SELECT id, project_id, lang_pair,
                   substr(IFNULL(output_text,''),1,120) AS prev, created_at
            FROM trans_ext
            ORDER BY datetime(created_at) DESC
            LIMIT 200
            """
        ).fetchall()
    except Exception as e:
        log_event("WARNING", "read trans_ext failed", err=str(e))
        rows = []

    if not rows:
        st.info("暂无历史记录。")
    else:
        for rid, pid, lp, prev, ts in rows:
            # 项目标题(做语料标题/展示)
            ttl_row = cur.execute("SELECT IFNULL(title,'') FROM items WHERE id=?", (pid,)).fetchone()
            proj_title = (ttl_row or [""])[0] or f"project#{pid}"

            with st.expander(f"#{rid}｜项目 {pid}｜{proj_title}｜{lp}｜{ts}", expanded=False):
                # 译文全文 & 源文件路径
                det = cur.execute("SELECT output_text, src_path FROM trans_ext WHERE id=?", (rid,)).fetchone()
                tgt_full, src_path = det or ("", "")
                st.code(prev or "", language="text")
                st.text_area("译文全文", tgt_full or "", height=220, key=f"hist_full_{rid}")

                # 尝试读取原文(如果当时保存了源文件路径)
                try:
                    src_full = read_source_file(src_path) if src_path else ""
                except Exception:
                    src_full = ""

                with st.expander("原文预览(若上传了源文件)", expanded=False):
                    st.text_area("原文全文", src_full or "(未保存/未上传源文件)", height=160, key=f"hist_src_{rid}")

                c1, c2, c3, c4, c5 = st.columns(5)

                # 1) 添加进语料库 / 添加+重建索引
                with c1:
                    # 1-1 只写入语料库
                    if st.button(" 添加进语料库", key=f"hist_add_corpus_{rid}"):
                        if not src_full and not tgt_full:
                            st.warning("原文和译文都为空，无法写入语料库。")
                        else:
                            execute_write("""
                                INSERT INTO corpus (title, project_id, lang_pair, src_text, tgt_text, note, created_at)
                                VALUES (?, ?, ?, ?, ?, ?, datetime('now'))
                            """, (
                                f"{proj_title} · history#{rid}",
                                pid,
                                lp or "",
                                src_full or None,
                                tgt_full or "",
                                f"from trans_ext#{rid}",
                            ))
                            st.success(" 已写入语料库")

                    # 1-2 写入语料库并重建索引
                    if st.button(" 添加并重建索引", key=f"hist_add_corpus_rebuild_{rid}"):
                        if not src_full and not tgt_full:
                            st.warning("原文和译文都为空，无法写入语料库。")
                        else:
                            # 先写入语料库
                            execute_write("""
                                INSERT INTO corpus (title, project_id, lang_pair, src_text, tgt_text, note, created_at)
                                VALUES (?, ?, ?, ?, ?, ?, datetime('now'))
                            """, (
                                f"{proj_title} · history#{rid}",
                                pid,
                                lp or "",
                                src_full or None,
                                tgt_full or "",
                                f"from trans_ext#{rid}",
                            ))

                            # 再重建该项目的语义索引
                            res_idx = rebuild_project_semantic_index(pid)
                            if res_idx.get("ok"):
                                st.success(
                                    f" 已写入语料库并重建索引: 新增 {res_idx['added']} 条, 总量 {res_idx['total']} 条"
                                )
                            else:
                                st.warning(
                                    f"已写入语料库，但重建索引失败: {res_idx.get('msg','未知错误')}"
                                )

                # 2) 提取术语
                with c2:
                    if st.button(" 提取术语", key=f"hist_extract_terms_{rid}"):
                        ak, model = get_deepseek()
                        if not ak:
                            st.warning("未检测到 DeepSeek Key(请到“设置”页配置)")
                        else:
                            # 合并原文+译文.提高候选质量
                            big = ((src_full or "") + "\n" + (tgt_full or "")).strip()
                            res = safe_api_call(
                                ds_extract_terms,
                                big, ak, model,
                                context="历史记录术语抽取",
                                fallback=[],
                                src_lang="zh", tgt_lang="en"
                            )
                            if not res:
                                st.info("未抽取到术语或解析失败")
                            else:
                                ins = 0
                                for o in res:
                                    execute_write("""
                                        INSERT INTO term_ext (
                                            source_term, target_term, domain, project_id,
                                            strategy, example
                                        )
                                        VALUES (?, ?, ?, ?, ?, ?)
                                    """, (
                                        o.get("source_term") or "",
                                        (o.get("target_term") or None),
                                        (o.get("domain") or None),
                                        pid,
                                        (o.get("strategy") or "history-extract"),
                                        (o.get("example") or None),
                                    ))
                                    ins += 1
                                st.success(f" 已写入术语库 {ins} 条")

                # 3) 下载双语对照(CSV / DOCX)
                with c3:
                    if st.button("⬇ CSV 对照", key=f"hist_dl_bicsv_btn_{rid}"):
                        if not src_full:
                            st.warning("找不到原文(未上传源文件).无法生成 CSV 对照")
                        else:
                            try:
                                csv_name = f"bilingual_history_{rid}.csv"
                                csv_bytes = export_csv_bilingual((src_full, tgt_full),
                                    filename=f"bilingual_history_{rid}.csv"
                                )
                            except TypeError:
                                # 如果导出函数是 text→bytes 版本
                                csv_name = f"bilingual_history_{rid}.csv"
                                csv_bytes = export_csv_bilingual(src_full, tgt_full)
                            st.download_button("下载 CSV", data=csv_bytes,
                                               file_name=csv_name, mime="text/csv",
                                               key=f"hist_dl_bicsv_{rid}")

                with c4:
                    if st.button("⬇ DOCX 对照", key=f"hist_dl_bidocx_btn_{rid}"):
                        if not src_full:
                            st.warning("找不到原文(未上传源文件).无法生成 DOCX 对照")
                        else:
                            try:
                                docx_path = export_docx_bilingual(
                                    filename=f"bilingual_history_{rid}.docx"
                                )
                                with open(docx_path, "rb") as f:
                                    data_docx = f.read()
                            except TypeError:
                                data_docx = export_docx_bilingual(src_full, tgt_full)
                            st.download_button("下载 DOCX", data=data_docx,
                                               file_name=f"bilingual_history_{rid}.docx",
                                               mime=(
                                                   "application/vnd.openxmlformats-"
                                                   "officedocument.wordprocessingml.document"
                                               ),
                                               key=f"hist_dl_bidocx_{rid}")

                # 4)  删除本条历史(安全确认)
                with c5:
                    with st.expander(" 删除本条历史(不可恢复)", expanded=False):
                        st.warning("此操作将永久删除该条 trans_ext 记录(不影响已写入语料库/术语表的数据)。")
                        ok = st.checkbox(f"我确认删除 #{rid}", key=f"hist_del_ck_{rid}")
                        if st.button("确认删除", key=f"hist_del_btn_{rid}") and ok:
                            execute_write("DELETE FROM trans_ext WHERE id=?", (rid,))
                            st.success("已删除.请刷新页面查看结果。")
                            st.stop()  # 终止本次渲染.避免在已删除数据上继续操作

                # 原有的“下载译文 TXT”
                st.download_button("下载译文 (TXT)", tgt_full or "",
                                   file_name=f"history_{rid}.txt",
                                   mime="text/plain",
                                   key=f"hist_dl_txt_{rid}")

# ========== Tab4:语料库管理 ==========
elif choice_idx == 3:
    def render_corpus_manager(st, cur, conn, pid_prefix="corpus"):
        st.header(" 语料库管理")
        sk = make_sk(pid_prefix)

    render_corpus_manager(st, cur, conn)

    _ensure_project_ref_map()
    _ensure_project_switch_map()
    st.session_state.setdefault(CORPUS_TARGET_PROJECT, None)
    st.session_state.setdefault(CORPUS_TARGET_LABEL, "(请选择 Few-shot 目标项目)")

    sub = st.tabs(["新建语料", "浏览/检索", "使用与导出"])

    # -------- 新建语料 --------
    with sub[0]:
        st.subheader(" 上传 / 对齐 / 入库")

        colA, colB = st.columns(2)
        with colA:
            one_file = st.file_uploader("① 单个文件(DOCX 表格对照 / 单语 DOCX/TXT/PDF)",
                                        type=["docx", "txt", "pdf"], key="up_one")
        with colB:
            two_src = st.file_uploader("② 原文文件(可选:与 ③ 搭配做对齐)",
                                       type=["docx", "txt", "csv", "pdf"], key="up_src")
            two_tgt = st.file_uploader("③ 译文文件(可选:与 ② 搭配做对齐)",
                                       type=["docx", "txt", "csv", "pdf"], key="up_tgt")

        st.divider()
        meta1, meta2, meta3 = st.columns([2, 1, 1])
        with meta1:
            title = st.text_input("语料标题", value="未命名语料")
        with meta2:
            lp = st.selectbox("方向", ["自动", "中译英", "英译中"])
        with meta3:
            pid_val = st.text_input("项目ID(可留空)")
        pid = int(pid_val) if pid_val.strip().isdigit() else None

        ins = 0
        pairs_key = sk("pairs")
        pairs = st.session_state.get(pairs_key, [])
        src_text = tgt_text = ""
        preview_df = None
        expected_first = _expected_first_from_lp(lp)

        # ========== 路径 A:单个文件 ==========
        if one_file is not None and (two_src is None and two_tgt is None):
            ext = (one_file.name.split(".")[-1] or "").lower()
            bio = io.BytesIO(one_file.getvalue())

            if ext == "docx":
                # try table pairs first
                tables = read_docx_tables_info(io.BytesIO(bio.getvalue()))
                if tables:
                    st.caption("检测到 DOCX 表格，优先作为双语对照导入。")
                    # default to first table col0/col1
                    pairs = extract_pairs_from_docx_table(
                        io.BytesIO(bio.getvalue()),
                        table_index=0,
                        src_col=0,
                        tgt_col=1,
                    )
                    st.session_state[pairs_key] = pairs
                if not pairs:
                    pairs = extract_pairs_from_docx_paragraphs(
                        io.BytesIO(bio.getvalue()),
                        expected_first=expected_first,
                    )
                    if pairs:
                        st.caption("检测到中英交替段落，已自动配对为双语语料。")
                        st.session_state[pairs_key] = pairs
                    else:
                        src_text = read_docx_text(io.BytesIO(bio.getvalue()))
            elif ext == "txt":
                src_text = read_txt(bio)

            elif ext == "pdf":
                src_text = read_pdf_text(io.BytesIO(bio.getvalue()))

        # ========== 路径 B:两个文件(原文 + 译文)==========
        elif two_src is not None and two_tgt is not None:
            def read_any(f):
                e = (f.name.split(".")[-1] or "").lower()
                b = io.BytesIO(f.getvalue())
                if e == "docx":
                    return read_docx_text(b)
                if e == "txt":
                    return read_txt(b)
                if e == "csv":
                    try:
                        df = pd.read_csv(b)
                        return "\n".join(df.iloc[:, 0].astype(str).fillna(""))
                    except Exception:
                        return ""
                if e == "pdf":
                    return read_pdf_text(b)
                return ""
            src_text = read_any(two_src)
            tgt_text = read_any(two_tgt)

        # ========== 预览与决定入库方式 ==========
        # 情况 1:有 pairs(来自 DOCX 表格或自动配对)
        if pairs:
            st.success(f"解析到 {len(pairs)} 对")
            preview_df = pd.DataFrame(pairs[:200], columns=["源句", "目标句"])

        # 情况 2:没有 pairs，但拿到 src/tgt 文本 -> 切句/对齐
        elif src_text and tgt_text:
            sents_src = split_sents(src_text, "zh" if lp.startswith("中") else "auto")
            sents_tgt = split_sents(tgt_text, "en" if lp.startswith("英") else "auto")
            st.caption(f"将对齐: src={len(sents_src)}  tgt={len(sents_tgt)}")
            if st.button(" 执行语义对齐", key="do_align"):
                pairs_aligned = align_semantic(sents_src, sents_tgt, max_jump=5)
                st.info(f"对齐得到 {len(pairs_aligned)} 对")
                pairs = [(s, t) for (s, t, score) in pairs_aligned]
                if pairs:
                    preview_df = pd.DataFrame(pairs[:200], columns=["源句", "目标句"])
                    st.session_state[pairs_key] = pairs

        # 情况 3:只有单语文本(PDF/DOCX/TXT)
        elif src_text and not tgt_text:
            sents_src = split_sents(src_text, "zh" if lp.startswith("中") else "auto")
            st.info(f"检测到单语文本，共 {len(sents_src)} 句。将以单语语料写入(译文为空)。")
            preview_df = pd.DataFrame(
                [{"源句": s, "目标句": ""} for s in sents_src[:200]]
            )

        if preview_df is not None:
            st.dataframe(preview_df, width='stretch')

        # \u2014\u2014 \u6309\u94ae + \u9009\u9879:\u5bfc\u5165\u8bed\u6599\u5e93 | \u540c\u65f6\u91cd\u5efa\u7d22\u5f15
        c_imp, c_opt, c_build = st.columns([1, 1, 1])
        do_import = c_imp.button(" 写入语料库并提取术语", type="primary", key="write_pairs_btn")
        do_build_opt = c_opt.checkbox(
            "\u5bfc\u5165\u540e\u7acb\u5373\u91cd\u5efa\u7d22\u5f15",
            value=True,
            key="build_vec_opt",
        )
        only_build_now = c_build.button(" \u4ec5\u91cd\u5efa\u7d22\u5f15(\u4e0d\u5bfc\u5165)", key="only_build")

        st.caption(
            "\u63d0\u793a: \u7d22\u5f15\u4e5f\u53ef\u4ee5\u7a0d\u540e\u5728\u201c\u4f7f\u7528\u4e0e\u7d22\u5f15 / "
            "\u5bfc\u51fa\u201d\u9875\u7684 C \u533a\u7edf\u4e00\u91cd\u5efa\u3002"
            "\u9875\u7684 C \u533a\u7edf\u4e00\u91cd\u5efa\u3002"
        )

        default_title = ""
        if one_file is not None:
            default_title = one_file.name
        elif two_src is not None:
            default_title = two_src.name

        if do_import:
            import_corpus_from_upload(
                st,
                cur,
                conn,
                pid=pid,              # meta 区的项目 ID
                title=title,          # 输入的语料标题
                lp=lp,                # 方向选择
                pairs=pairs,
                src_text=src_text,
                tgt_text=tgt_text,
                default_title=default_title,
                build_after_import=do_build_opt,
            )
            # 术语提取：对齐对或原文/译文组合文本
            text_chunks = []
            if pairs:
                text_chunks.extend([f"{s}\n{t}" for s, t in pairs])
            text_chunks.extend([src_text, tgt_text])
            _extract_terms_for_upload(text_chunks, lp=lp, pid=pid)
            st.session_state.pop(pairs_key, None)

        if only_build_now:
            # 仅重建当前项目的语义索引(不导入新语料)
            if pid:
                res_idx = rebuild_project_semantic_index(pid)
                if res_idx.get("ok"):
                    st.success(
                        f" 索引已重建: 新增 {res_idx['added']}，总量 {res_idx['total']}。"
                    )
                else:
                    st.error(f"重建失败: {res_idx.get('msg','未知错误')}")
            else:
                st.warning("请先在上方填写有效的项目ID,再重建索引。")

    # -------- 浏览/检索 --------
    with sub[1]:
        st.subheader(" 浏览/检索")
        k1, k2, k3 = st.columns([2, 1, 1])
        with k1:
            kw = st.text_input("关键词(标题/源文/译文/备注)", "", key=sk("kw"))
        with k2:
            lp_filter = st.selectbox("方向过滤", ["全部", "中译英", "英译中", "自动"], key=sk("lp_filter"))
        with k3:
            limit = st.number_input("条数", min_value=10, max_value=1000, value=200, step=10, key=sk("limit"))

        st.markdown("---")
        sql = """
        SELECT id, title, IFNULL(project_id,''), IFNULL(lang_pair,''), 
               substr(IFNULL(tgt_text,''),1,80), created_at
        FROM corpus
        WHERE 1=1
        """
        params = []
        if kw.strip():
            like = f"%{kw.strip()}%"
            sql += (
                " AND (title LIKE ? OR IFNULL(note,'') LIKE ? OR "
                "IFNULL(src_text,'') LIKE ? OR IFNULL(tgt_text,'') LIKE ?)"
            )
            params.extend([like, like, like, like])
        if lp_filter != "全部":
            sql += " AND lang_pair=?"
            params.append(lp_filter)
        sql += " ORDER BY id DESC LIMIT ?"
        params.append(int(limit))

        rows = cur.execute(sql, params).fetchall()
        if not rows:
            st.info("暂无匹配语料。")

        else:
            sel_key = sk("corpus_sel_set")
            sel_set = st.session_state.setdefault(sel_key, set())
            select_all = st.checkbox("本页全选", value=False, key=sk("corpus_sel_all"))
            if select_all:
                sel_set.update([r[0] for r in rows])
            else:
                # 不自动清空，用户可单独勾选；若想清空可点下方“清空选择”
                pass

            col_btns = st.columns(3)
            with col_btns[0]:
                if st.button("删除已选语料", type="primary", key=sk("del_corpus_selected")):
                    if sel_set:
                        qmarks = ",".join("?" * len(sel_set))
                        execute_write(f"DELETE FROM corpus WHERE id IN ({qmarks})", tuple(sel_set))
                        st.success(f"已删除 {len(sel_set)} 条语料。")
                        sel_set.clear()
                    else:
                        st.warning("未选择任何语料。")
            with col_btns[1]:
                if st.button("清空选择", key=sk("clear_corpus_sel")):
                    sel_set.clear()
            st.divider()

            for rid, ttl, pj, lpv, prev, ctime in rows:
                with st.expander(f"[{rid}] {ttl} | 项目:{pj} | 方向:{lpv} | {ctime}"):
                    checked = st.checkbox("选中删除", value=(rid in sel_set), key=sk(f"sel_{rid}"))
                    if checked:
                        sel_set.add(rid)
                    else:
                        sel_set.discard(rid)
                    st.write(f"**ID**: {rid}  **项目ID**: {pj}  **方向**: {lpv}  **时间**: {ctime}")
                    st.write(f"**预览(译文前80字)**: {prev}")
                    det = cur.execute(
                        "SELECT IFNULL(src_text,''), IFNULL(tgt_text,''), IFNULL(domain,'') "
                        "FROM corpus WHERE id=?",
                        (rid,)
                    ).fetchone()
                    src_all, tgt_all, dom_val = det or ("", "", "")
                    st.text_area("源文", src_all, height=160, key=sk(f"src_{rid}"))
                    st.text_area("译文", tgt_all, height=160, key=sk(f"tgt_{rid}"))

                    with st.expander("一词多义(快速添加)"):
                        kw_term = (kw or "").strip()
                        if not kw_term:
                            st.info("请先在上方输入中文关键词。")
                        else:
                            en_pick = st.text_input(
                                "译入语(手工输入)",
                                key=sk(f"poly_en_pick_{rid}"),
                            )
                            if st.button("一键加入一词多义列表", key=sk(f"poly_add_{rid}")):
                                if not en_pick.strip():
                                    st.warning("请先输入译入语。")
                                else:
                                    ex_src = _pick_sentence(src_all, kw_term, "zh")
                                    ex_tgt = _pick_sentence(tgt_all, en_pick.strip(), "en")
                                    added = add_polysemy_manual(
                                        cur,
                                        conn,
                                        kw_term,
                                        [en_pick.strip()],
                                        dom_val,
                                        ex_src or src_all,
                                        ex_tgt or tgt_all,
                                    )
                                    if added:
                                        st.success("已加入一词多义列表。")
                                    else:
                                        st.info("未添加新译入语（可能已存在）。")

                    c1, c2, c3, c4 = st.columns(4)
                    with c1:
                        if st.button("加入参考集合", key=sk(f"add_ref_{rid}")):
                            target_pid = st.session_state.get(CORPUS_TARGET_PROJECT)
                            if not target_pid:
                                st.warning("请先在上方选择 Few-shot 目标项目，再添加参考语料。")
                            else:
                                refs = get_project_ref_ids(target_pid)
                                refs.add(int(rid))
                                st.success(f" 已加入项目 #{target_pid} 的参考集合(“使用与导出”查看)")
                    with c2:
                        if st.button("导出TXT", key=sk(f"cor_txt_{rid}")):
                            st.download_button(
                                "下载译文TXT",
                                tgt_all or "",
                                file_name=f"corpus_{rid}.txt",
                                mime="text/plain",
                                key=sk(f"cor_txt_dl_{rid}")
                            )
                    with c3:
                        if st.button("导出CSV(中英对照)", key=sk(f"cor_csv_{rid}")):
                            df_out = pd.DataFrame(
                                [{"source": src_all, "target": tgt_all}]
                            )
                            csv_data = df_out.to_csv(index=False)
                            st.download_button(
                                "下载CSV",
                                csv_data,
                                file_name=f"corpus_{rid}.csv",
                                mime="text/csv",
                                key=sk(f"cor_csv_dl_{rid}")
                            )
                    with c4:
                        if st.button("删除", key=sk(f"del_{rid}")):
                            execute_write("DELETE FROM corpus WHERE id=?", (rid,))
                            st.warning(" 已删除，刷新后生效")
                            st.rerun()

    # -------- 使用与导出 --------
    with sub[2]:
        st.subheader(" 一词多义库管理 ")
        
        with st.expander("➕ 手工新增一词多义规则", expanded=False):
            with st.form("manual_poly_add"):
                c1, c2 = st.columns([1, 2])
                with c1:
                    new_term = st.text_input("中文一词多义 (如: 精神)")
                with c2:
                    new_senses = st.text_area("英文译项 (每行一个，可加备注)", placeholder="spirit (用于会议精神)\nmind (用于精神健康)")
                
                submitted = st.form_submit_button("添加规则")
                if submitted:
                    if not new_term.strip() or not new_senses.strip():
                        st.error("请填写完整信息")
                    else:
                        # 解析输入
                        lines = [line.strip() for line in new_senses.split('\n') if line.strip()]
                        added_count = 0
                        
                        # 1. 确保主词存在
                        execute_write(
                            "INSERT INTO poly_term(source_term, created_at) VALUES(?, datetime('now')) "
                            "ON CONFLICT(source_term) DO NOTHING", (new_term.strip(),)
                        )
                        row = execute_write(
                            "SELECT id FROM poly_term WHERE source_term=?",
                            (new_term.strip(),),
                        ).fetchone()
                        term_id = row[0]

                        # 2. 插入义项
                        for line in lines:
                            parts = line.split(None, 1)
                            kw = parts[0]
                            desc = parts[1] if len(parts) > 1 else ""
                            
                            execute_write(
                                "INSERT INTO poly_sense(term_id, target_term, domain, example_src) "
                                "VALUES(?, ?, ?, ?)",
                                (term_id, kw, "人工录入", desc) # 将备注存入 example_src 或 domain 字段供提示显示
                            )
                            added_count += 1
                        
                        st.success(f"已添加 '{new_term}' 的 {added_count} 个译项。")
                        st.rerun()

        st.markdown("---")

        # === 现有列表展示与管理 ===
        st.write("📚 **一词多义列表**")
        poly_kw = st.text_input("查询列表", "", key=sk("poly_search"))
        
        params = []
        sql = """
        SELECT t.id, t.source_term, COUNT(s.id) AS sense_cnt
        FROM poly_term t
        LEFT JOIN poly_sense s ON t.id = s.term_id
        WHERE 1=1
        """
        if poly_kw.strip():
            like = f"%{poly_kw.strip()}%"
            sql += " AND (t.source_term LIKE ? OR s.target_term LIKE ?)"
            params = [like, like]
        
        sql += " GROUP BY t.id, t.source_term ORDER BY t.id DESC LIMIT 100"
        
        terms = execute_write(sql, params).fetchall()
        
        if not terms:
            st.info("暂无数据。")
        else:
            for term_id, term, sense_cnt in terms:
                with st.expander(f"{term} (包含 {sense_cnt} 个译项)", expanded=False):
                    # 获取详情
                    senses = execute_write(
                        "SELECT id, target_term, IFNULL(example_src,'') FROM poly_sense WHERE term_id=?", 
                        (term_id,)
                    ).fetchall()
                    
                    # 显示列表
                    if senses:
                        sense_data = [{"ID": s[0], "英文译法": s[1], "备注/例句": s[2]} for s in senses]
                        st.table(pd.DataFrame(sense_data))
                    
                    # 删除按钮
                    c_del_term, c_del_sense = st.columns(2)
                    with c_del_term:
                        if st.button(f"删除整个词条 '{term}'", key=sk(f"del_t_{term_id}")):
                            execute_write("DELETE FROM poly_sense WHERE term_id=?", (term_id,))
                            execute_write("DELETE FROM poly_term WHERE id=?", (term_id,))
                            st.success("已删除")
                            st.rerun()

# ========== Tab5:索引管理 ==========

elif choice_idx == 4:
    st.subheader('索引管理')
    render_index_manager(st, conn, cur)
    st.markdown('---')
    render_index_manager_by_domain(st, conn, cur)



# ===== Override embedder with fallbacks (PyTorch unavailable on py3.13) =====
try:
    import numpy as np  # ensure available
except Exception:
    pass

def _make_fastembed_encoder(model_name: str = 'BAAI/bge-small-en-v1.5'):
    try:
        from fastembed import TextEmbedding
    except Exception as e:
        raise RuntimeError(f'fastembed not available: {e}')
    model = TextEmbedding(model_name)
    def encode(texts: list[str]):
        if not texts:
            return np.zeros((0, model.dim), dtype='float32')
        vecs = list(model.embed(texts))
        arr = np.asarray(vecs, dtype='float32')
        arr = arr / (np.linalg.norm(arr, axis=1, keepdims=True) + 1e-12)
        return arr
    return encode

def _make_tfidf_encoder():
    try:
        from sklearn.feature_extraction.text import TfidfVectorizer
    except Exception as e:
        raise RuntimeError(f'tfidf backend unavailable: {e}')
    vec = TfidfVectorizer(min_df=1)
    def encode(texts: list[str]):
        if not texts:
            return np.zeros((0,1), dtype='float32')
        X = vec.fit_transform(texts)
        arr = X.toarray().astype('float32')
        arr = arr / (np.linalg.norm(arr, axis=1, keepdims=True) + 1e-12)
        return arr
    return encode

@st.cache_resource(show_spinner=False)
def get_embedder():
    # 1) sentence-transformers (needs torch)
    try:
        import torch  # noqa: F401
        encode_st = _make_sbert_encoder(
            'distiluse-base-multilingual-cased-v1',
            normalize_embeddings=True,
            extra_normalize=True,
            batch_size=32,
            handle_empty=True,
        )
        return 'st', encode_st
    except Exception as e:
        log_event('WARNING', 'sentence-transformers unavailable, fallback to fastembed/tfidf', err=str(e))

    # 2) fastembed
    try:
        encode_fe = _make_fastembed_encoder()
        return 'fastembed', encode_fe
    except Exception as e:
        log_event('WARNING', 'fastembed unavailable, fallback to tfidf', err=str(e))

    # 3) TF-IDF fallback
    encode_tf = _make_tfidf_encoder()
    return 'tfidf', encode_tf
