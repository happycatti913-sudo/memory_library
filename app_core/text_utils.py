# -*- coding: utf-8 -*-
"""文本读取与分句等通用工具。"""

import io
import re
from typing import Tuple


def _lazy_docx():
    try:
        import docx  # python-docx
        return docx
    except Exception:
        return None


def _normalize(t: str) -> str:
    if not t:
        return ""
    t = t.replace("\xa0", " ").replace("\u200b", "")
    t = re.sub(r"[ \t]+", " ", t)
    t = re.sub(r"\n{2,}", "\n", t)
    return t.strip()


def read_docx_tables_info(file_like):
    docx = _lazy_docx()
    if not docx:
        return {}
    try:
        doc = docx.Document(file_like)
    except Exception:
        return {}
    info = {}
    for ti, tbl in enumerate(doc.tables):
        rows = len(tbl.rows)
        cols = len(tbl.columns) if rows else 0
        prev = []
        for r in tbl.rows[: min(6, rows)]:
            prev.append(tuple(_normalize(c.text) for c in r.cells))
        info[ti] = {"rows": rows, "cols": cols, "preview": prev}
    return info


def extract_pairs_from_docx_table(
    file_like,
    table_index: int = 0,
    src_col: int = 0,
    tgt_col: int = 1,
    ffill: bool = True,
    drop_empty_both: bool = True,
    dedup: bool = True,
):
    docx = _lazy_docx()
    if not docx:
        return []
    try:
        doc = docx.Document(file_like)
    except Exception:
        return []
    if table_index >= len(doc.tables):
        return []
    tbl = doc.tables[table_index]
    rows = []
    for r in tbl.rows:
        rows.append([_normalize(c.text) for c in r.cells])
    if not rows:
        return []
    max_cols = max(len(r) for r in rows)
    if src_col >= max_cols or tgt_col >= max_cols:
        return []

    if ffill:
        for col in (src_col, tgt_col):
            last = ""
            for i in range(len(rows)):
                val = rows[i][col] if col < len(rows[i]) else ""
                if val:
                    last = val
                else:
                    rows[i][col] = last

    pairs = []
    for r in rows:
        s = r[src_col] if src_col < len(r) else ""
        t = r[tgt_col] if tgt_col < len(r) else ""
        s, t = s.strip(), t.strip()
        if drop_empty_both and (not s and not t):
            continue
        pairs.append((s, t))

    if dedup:
        seen, out = set(), []
        for p in pairs:
            if p in seen:
                continue
            seen.add(p)
            out.append(p)
        pairs = out
    return pairs


def read_docx_text(file_like) -> str:
    docx = _lazy_docx()
    if not docx:
        return ""
    try:
        doc = docx.Document(file_like)
    except Exception:
        return ""
    blocks = []
    for p in doc.paragraphs:
        t = _normalize(p.text)
        if t:
            blocks.append(t)
    for tbl in doc.tables:
        for r in tbl.rows:
            line = " ".join(_normalize(c.text) for c in r.cells if _normalize(c.text))
            if line:
                blocks.append(line)
    return "\n".join(blocks)


def read_txt(file_like_or_bytes) -> str:
    try:
        if hasattr(file_like_or_bytes, "read"):
            data = file_like_or_bytes.read()
        else:
            data = file_like_or_bytes
        if isinstance(data, bytes):
            try:
                return data.decode("utf-8")
            except Exception:
                return data.decode("utf-8", errors="ignore")
        return str(data)
    except Exception:
        return ""


def read_pdf_text(file_like) -> str:
    try:
        from PyPDF2 import PdfReader

        reader = PdfReader(file_like)
        txts = []
        for page in reader.pages:
            t = page.extract_text() or ""
            t = _normalize(t)
            if t:
                txts.append(t)
        return "\n".join(txts)
    except Exception:
        try:
            from pdfminer.high_level import extract_text

            if hasattr(file_like, "read"):
                data = file_like.read()
            else:
                data = file_like
            return _normalize(extract_text(io.BytesIO(data)))
        except Exception:
            return ""


def _lazy_import_doc_pdf():
    docx = pdfplumber = None
    try:
        import docx as _docx

        docx = _docx
    except Exception:
        pass
    try:
        import pdfplumber as _pdfplumber

        pdfplumber = _pdfplumber
    except Exception:
        pass
    return docx, pdfplumber


def split_paragraphs(text: str) -> list[str]:
    text = (text or "").replace("\r\n", "\n").replace("\r", "\n")
    if "\n\n" not in text and "\n \n" not in text:
        lines = [ln.strip() for ln in text.split("\n")]
        return [ln for ln in lines if ln]

    parts = re.split(r"\n\s*\n+", text)
    paras = []
    for p in parts:
        p = p.strip()
        if not p:
            continue
        p = re.sub(r"\s*\n\s*", " ", p)
        paras.append(p)
    return paras


def pair_paragraphs(src_full: str, tgt_full: str) -> Tuple[list[str], list[str]]:
    src_paras = split_paragraphs(src_full or "")
    tgt_paras = split_paragraphs(tgt_full or "")

    n = max(len(src_paras), len(tgt_paras))
    src_paras += [""] * (n - len(src_paras))
    tgt_paras += [""] * (n - len(tgt_paras))
    return src_paras, tgt_paras


_RE_WS = re.compile(r"[ \t\u00A0\u200B\u200C\u200D]+")
_RE_ZH_SENT = re.compile(r"(?<=[。！？；])\s*")
_RE_EN_SENT = re.compile(r"(?<=[\.\?\!;:])\s+")
_RE_BLANK_PARA = re.compile(r"\n{2,}")


def _norm_text(text: str) -> str:
    t = (text or "").replace("\r\n", "\n").replace("\r", "\n").replace("\x0b", "\n")
    t = _RE_WS.sub(" ", t)
    t = re.sub(r"\n{3,}", "\n\n", t)
    t = _RE_BLANK_PARA.sub("\n\n", t)
    return t.strip()


def _is_zh(text: str) -> bool:
    if not text:
        return False
    zh_chars = sum(1 for ch in text if "\u4e00" <= ch <= "\u9fff")
    return zh_chars >= max(1, len(text) // 3)


def split_sents(
    text: str,
    lang_hint: str = "zh",
    min_char: int = 2,
    prefer_newline: bool = False,
    **kwargs,
):
    lang = kwargs.get("lang", lang_hint)

    t = _norm_text(text)
    if not t:
        return []

    pieces = []

    if prefer_newline and ("\n" in t):
        lines = [ln.strip() for ln in t.split("\n") if ln.strip()]
        for ln in lines:
            if lang == "auto":
                cur_lang = "zh" if _is_zh(ln) else "en"
            else:
                cur_lang = lang

            if cur_lang.startswith(("zh", "cn")):
                sents = [s.strip() for s in _RE_ZH_SENT.split(ln) if s and s.strip()]
            else:
                sents = [s.strip() for s in _RE_EN_SENT.split(ln) if s and s.strip()]

            pieces.extend(sents if sents else [ln])
    else:
        if lang == "auto":
            cur_lang = "zh" if _is_zh(t) else "en"
        else:
            cur_lang = lang

        if cur_lang.startswith(("zh", "cn")):
            pieces = [s.strip() for s in _RE_ZH_SENT.split(t) if s and s.strip()]
        else:
            pieces = [s.strip() for s in _RE_EN_SENT.split(t) if s and s.strip()]

        if not pieces:
            pieces = [t]

    return [x for x in pieces if len(x) >= min_char]


def _split_pair_for_index(src_text: str, tgt_text: str) -> Tuple[list[str], list[str]]:
    try:
        return split_sents(src_text, lang_hint="zh"), split_sents(tgt_text, lang_hint="en")
    except Exception:
        src_sents = src_text.split("。") if src_text else []
        tgt_sents = tgt_text.split(".") if tgt_text else []
        return src_sents, tgt_sents


split_sentences = split_sents


def split_blocks(text: str, max_len: int = 1200):
    blocks, curbuf = [], ""
    for line in text.splitlines(True):
        if len(curbuf) + len(line) > max_len:
            if curbuf:
                blocks.append(curbuf)
            curbuf = ""
        curbuf += line
    if curbuf:
        blocks.append(curbuf)
    return blocks


def build_bilingual_lines(src_text: str, tgt_text: str):
    return pair_paragraphs(src_text, tgt_text)

