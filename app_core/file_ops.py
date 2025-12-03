# -*- coding: utf-8 -*-
"""通用文件读取与双语导出工具。"""

from __future__ import annotations

import io
import os

import pandas as pd

from app_core.text_utils import build_bilingual_lines


def read_source_file(path: str) -> str:
    """读取 txt/docx/xlsx 等源文件文本，兜底兼容多种编码。"""
    if not path or not os.path.exists(path):
        return ""
    ext = os.path.splitext(path)[1].lower()
    if ext == ".txt":
        for enc in ["utf-8", "utf-8-sig", "gb18030", "gbk"]:
            try:
                with open(path, "r", encoding=enc) as f:
                    return f.read()
            except Exception:
                continue
        with open(path, "r", errors="ignore") as f:
            return f.read()
    elif ext == ".docx":
        try:
            from docx import Document

            doc = Document(path)
            return "\n".join(p.text for p in doc.paragraphs)
        except Exception:
            return ""
    elif ext == ".xlsx":
        try:
            xls = pd.ExcelFile(path)
            parts = []
            for sheet in xls.sheet_names:
                df = xls.parse(sheet)
                parts.append(df.astype(str).to_csv(sep=" ", index=False, header=False))
            return "\n".join(parts)
        except Exception:
            return ""
    else:
        try:
            with open(path, "r", encoding="utf-8") as f:
                return f.read()
        except Exception:
            try:
                with open(path, "r", errors="ignore") as f:
                    return f.read()
            except Exception:
                return ""


def export_csv_bilingual(src_text: str, tgt_text: str) -> bytes:
    s, t = build_bilingual_lines(src_text, tgt_text)
    df = pd.DataFrame({"Source": s, "Target": t})
    buf = io.StringIO()
    df.to_csv(buf, index=False)
    return buf.getvalue().encode("utf-8-sig")


def export_docx_bilingual(src_text: str, tgt_text: str) -> bytes:
    try:
        from docx import Document
        from docx.oxml.ns import qn
    except Exception:
        return b""

    doc = Document()
    try:
        doc.styles["Normal"].font.name = "Calibri"
        doc.styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    except Exception:
        pass

    s, t = build_bilingual_lines(src_text, tgt_text)
    table = doc.add_table(rows=1, cols=2)
    hdr = table.rows[0].cells
    hdr[0].text = "Source"
    hdr[1].text = "Target"
    for a, b in zip(s, t):
        row = table.add_row().cells
        row[0].text = a
        row[1].text = b
    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()
