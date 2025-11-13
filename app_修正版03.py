# -*- coding: utf-8 -*-
"""
个人翻译知识库管理系统(修正版03)
- Tab1 📂 翻译项目管理:新建项目、文件上传、执行翻译(DeepSeek API).导出对照/原格式.写入历史
- Tab2 📘 术语库管理:查询/编辑/删除、CSV批量导入、统计/导出、快速搜索、批量挂接项目、历史抽取术语、分类管理
- Tab3 📊 翻译历史:查看、下载译文(简版占位.按需扩展)
- Tab4 📚 语料库管理:新增/检索/合并/Few-shot 注入
- Tab5 ⚙ 设置: DeepSeek Key 提示
"""

# ==== stdlib ====
import os
import re
import io
import sys
import json
import uuid
import time
import sqlite3
import streamlit as st
import pandas as pd
from pathlib import Path
from datetime import datetime

# 让同目录下的 kb_dynamic.py 可被导入(如果存在)
sys.path.append(os.path.dirname(__file__))

# ======== 基本路径设置 ========

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
DB_PATH = os.path.join(BASE_DIR, "kb.db")

PROJECT_DIR = os.path.join(BASE_DIR, "uploads")
os.makedirs(PROJECT_DIR, exist_ok=True)

INDEX_DIR = Path(BASE_DIR) / ".cache_index"
INDEX_DIR.mkdir(exist_ok=True)

# ---------- 优化:给术语表 project_id 建索引 ----------
def _index_paths(project_id: int):
    base_dir = os.path.join(BASE_DIR, ".cache_index")
    os.makedirs(base_dir, exist_ok=True)
    return (
        os.path.join(base_dir, f"faiss_{project_id}.bin"),
        os.path.join(base_dir, f"vecmap_{project_id}.json"),
        os.path.join(base_dir, f"vectors_{project_id}.npy"),
    )

# ==== third-party ====
try:
    from docx import Document  # 在需要处仍会 try/except
except Exception:
    Document = None

# ========== 页面设置 ==========
st.set_page_config(page_title="个人翻译知识库管理系统 · 修正版03", layout="wide")

# ========== kb_dynamic (可选) ==========
KBEmbedder = None
recommend_for_segment = None
build_prompt_strict = None
build_prompt_soft = None
try:
    from kb_dynamic import (
        KBEmbedder as _KBEmbedder,
        recommend_for_segment as _recommend_for_segment,
        build_prompt_strict as _build_prompt_strict,
        build_prompt_soft as _build_prompt_soft,
    )
    KBEmbedder = _KBEmbedder
    recommend_for_segment = _recommend_for_segment
    build_prompt_strict = _build_prompt_strict
    build_prompt_soft = _build_prompt_soft
except Exception:
    pass  # 允许缺失;动态术语推荐功能将自动降级

# ========== 工具函数 ==========
def make_sk(prefix: str):
    """返回一个带有前缀的 key 生成器"""
    return lambda name, id=None: f"{prefix}_{name}_{id}" if id else f"{prefix}_{name}"
# 全局默认 key 生成器(替代被删除的计数器版 sk)
sk = make_sk("global")

def render_table(df, *, key=None, hide_index=True, editable=False):
    """
    统一渲染表格(对旧/新 Streamlit 都安全):
    - editable=False: 只读(用 data_editor disabled=True 以保留 key)
    - editable=True : 可编辑
    - 不再传 width 参数.避免 'str' as int 的报错
    """
    try:
        if editable:
            return st.data_editor(
                df,
                hide_index=hide_index,
                key=key,
            )
        if key is not None:
            return st.data_editor(
                df,
                hide_index=hide_index,
                disabled=True,
                key=key,
            )
        return st.dataframe(df, hide_index=hide_index)
    except TypeError:
        return st.data_editor(
            df,
            hide_index=hide_index,
            disabled=not editable,
            key=key,
        )

# ======= 获取某条历史记录对应的原文(优先 items.body.兜底 src_path 仅作为标题提示)=======
def _get_source_text_for_history(cur, project_id):
    row = cur.execute("SELECT body FROM items WHERE id=?", (project_id,)).fetchone()
    return (row[0] if row and row[0] else "") or ""

# ======= 轻量术语候选(中英都可;你后续可换成 DeepSeek 抽取)=======
def _simple_term_candidates(text, topn=50):
    import re, collections
    # very light: 英文按词.中文用简单正则;停用词可按需扩展
    STOP_EN = set("a an the and or of to in for on with by from as at is are was were be been being this that these those it its".split())
    is_en = len(re.sub(r'[^A-Za-z]', '', text)) >= len(re.sub(r'[^\u4e00-\u9fff]', '', text))
    if is_en:
        toks = [t.lower() for t in re.findall(r"[A-Za-z][A-Za-z\-']+", text)]
        toks = [t for t in toks if t not in STOP_EN and len(t) > 2]
    else:
        import jieba
        toks = [t.strip() for t in jieba.cut(text) if re.search(r"[\u4e00-\u9fff]", t)]
        toks = [t for t in toks if len(t) >= 2 and not re.match(r"^[0-9]+$", t)]
    cnt = collections.Counter(toks)
    return [{"term": k, "freq": v} for k, v in cnt.most_common(topn)]

# ======= 对齐并导出(依赖你已有的 split_blocks / align_export)=======
def _split_bilingual_pairs(split_blocks, src_text, tgt_text):
    src_blocks = split_blocks(src_text, max_len=1200)
    tgt_blocks = split_blocks(tgt_text, max_len=1200)
    # 简单对齐:按 zip 对齐;长度不等时以较短为准
    n = min(len(src_blocks), len(tgt_blocks))
    return list(zip(src_blocks[:n], tgt_blocks[:n]))

def quick_diagnose_vectors(pid: int):
    """
    打印/提示项目向量索引状态.帮助排查“检索为空/维度不匹配/未建索引”等问题。
    """
    try:
        mode, index, mapping, vecs = _load_index(pid)
        if mode == "none":
            st.warning(f"项目 {pid} 尚未建立向量索引（.cache_index 无对应文件）。")
            return
        msg = f"索引模式: {mode}; 映射条数: {len(mapping)}"
        if mode == "faiss" and index is not None:
            msg += f"; FAISS ntotal: {index.ntotal}"
        elif vecs is not None:
            msg += f"; NPY shape: {getattr(vecs, 'shape', None)}"
        st.info(msg)

        # 抽样验证映射的 corpus_id 是否都能回查到文本
        bad = 0
        for m in mapping[:10]:
            cid = int(m.get("corpus_id") or -1)
            row = cur.execute("SELECT id FROM corpus WHERE id=?", (cid,)).fetchone()
            if not row:
                bad += 1
        if bad:
            st.warning(f"映射中有 {bad} 条 corpus_id 无法回查.请考虑重建索引。")
    except Exception as e:
        st.error(f"向量诊断异常:{e}")

# ---------- 文件读取工具 ----------
def _lazy_docx():
    try:
        import docx  # python-docx
        return docx
    except Exception:
        return None

def _normalize(t: str) -> str:
    if not t: return ""
    t = t.replace("\xa0", " ").replace("\u200b", "")
    t = re.sub(r"[ \t]+", " ", t)
    t = re.sub(r"\n{2,}", "\n", t)
    return t.strip()

def read_docx_tables_info(file_like):
    docx = _lazy_docx()
    if not docx: return {}
    try:
        doc = docx.Document(file_like)
    except Exception:
        return {}
    info = {}
    for ti, tbl in enumerate(doc.tables):
        rows = len(tbl.rows)
        cols = len(tbl.columns) if rows else 0
        prev = []
        for r in tbl.rows[: min(6, rows)]:
            prev.append(tuple(_normalize(c.text) for c in r.cells))
        info[ti] = {"rows": rows, "cols": cols, "preview": prev}
    return info

def extract_pairs_from_docx_table(file_like, table_index=0, src_col=0, tgt_col=1,
                                  ffill=True, drop_empty_both=True, dedup=True):
    docx = _lazy_docx()
    if not docx: return []
    try:
        doc = docx.Document(file_like)
    except Exception:
        return []
    if table_index >= len(doc.tables): return []
    tbl = doc.tables[table_index]
    rows = []
    for r in tbl.rows:
        rows.append([_normalize(c.text) for c in r.cells])
    if not rows: return []
    max_cols = max(len(r) for r in rows)
    if src_col >= max_cols or tgt_col >= max_cols: return []

    if ffill:
        for col in (src_col, tgt_col):
            last = ""
            for i in range(len(rows)):
                val = rows[i][col] if col < len(rows[i]) else ""
                if val: last = val
                else: rows[i][col] = last

    pairs = []
    for r in rows:
        s = r[src_col] if src_col < len(r) else ""
        t = r[tgt_col] if tgt_col < len(r) else ""
        s, t = s.strip(), t.strip()
        if drop_empty_both and (not s and not t):
            continue
        pairs.append((s, t))

    if dedup:
        seen, out = set(), []
        for p in pairs:
            if p in seen: continue
            seen.add(p); out.append(p)
        pairs = out
    return pairs

def read_docx_text(file_like) -> str:
    docx = _lazy_docx()
    if not docx: return ""
    try:
        doc = docx.Document(file_like)
    except Exception:
        return ""
    blocks = []
    for p in doc.paragraphs:
        t = _normalize(p.text)
        if t: blocks.append(t)
    # 把表格单元也拼成行.避免漏掉内容
    for tbl in doc.tables:
        for r in tbl.rows:
            line = " ".join(_normalize(c.text) for c in r.cells if _normalize(c.text))
            if line: blocks.append(line)
    return "\n".join(blocks)

def read_txt(file_like_or_bytes) -> str:
    try:
        if hasattr(file_like_or_bytes, "read"):
            data = file_like_or_bytes.read()
        else:
            data = file_like_or_bytes
        if isinstance(data, bytes):
            try: return data.decode("utf-8")
            except: return data.decode("utf-8", errors="ignore")
        return str(data)
    except Exception:
        return ""

def read_csv_two_cols(file_like, col_a=0, col_b=1):
    try:
        df = pd.read_csv(file_like)
        a = df.iloc[:, col_a].astype(str).fillna("").tolist()
        b = df.iloc[:, col_b].astype(str).fillna("").tolist()
        return list(zip(a, b))
    except Exception:
        return []

def read_pdf_text(file_like) -> str:
    # 优先 pypdf;失败再试 pdfminer.six
    try:
        from PyPDF2 import PdfReader
        reader = PdfReader(file_like)
        txts = []
        for page in reader.pages:
            t = page.extract_text() or ""
            t = _normalize(t)
            if t: txts.append(t)
        return "\n".join(txts)
    except Exception:
        try:
            from pdfminer.high_level import extract_text
            if hasattr(file_like, "read"):
                data = file_like.read()
            else:
                data = file_like
            return _normalize(extract_text(io.BytesIO(data)))
        except Exception:
            return ""
        
# ========== 向量召回(多后端:Sentence-Transformers → Fastembed → TF-IDF)==========
def _lazy_import_vec():
    import importlib
    np = importlib.import_module("numpy")

    # 尝试 faiss / faiss_cpu(可选)
    try:
        faiss = importlib.import_module("faiss")
    except Exception:
        try:
            faiss = importlib.import_module("faiss_cpu")
        except Exception:
            faiss = None

    # 1) 优先:sentence_transformers(需要 torch)
    SentenceTransformer = None
    try:
        SentenceTransformer = importlib.import_module("sentence_transformers").SentenceTransformer
    except Exception:
        SentenceTransformer = None

    # 2) 备选:fastembed(不需要 torch.后端是 onnxruntime)
    FastEmbedModel = None
    try:
        FastEmbedModel = importlib.import_module("fastembed").TextEmbedding
    except Exception:
        FastEmbedModel = None

    # 3) 兜底:scikit-learn TF-IDF
    TfidfVectorizer = None
    NearestNeighbors = None
    try:
        from sklearn.feature_extraction.text import TfidfVectorizer  # type: ignore
        from sklearn.neighbors import NearestNeighbors  # type: ignore
    except Exception:
        pass

    return np, faiss, SentenceTransformer, FastEmbedModel, TfidfVectorizer, NearestNeighbors

@st.cache_resource(show_spinner=False)
def get_embedder(model_name: str = "thenlper/gte-multilingual-base"):
    """
    返回 (backend, encoder):
      - backend: "st" | "fastembed" | "tfidf"
      - encoder: 可调用对象.具备 encode(texts)->np.ndarray 接口(tfidf 返回 (vectorizer, matrix) 的打包器)
    """
    np, faiss, SentenceTransformer, FastEmbedModel, TfidfVectorizer, _ = _lazy_import_vec()

    # 1) Sentence-Transformers(成功则优先)
    if SentenceTransformer is not None:
        try:
            model = SentenceTransformer(model_name)
            def _encode_st(texts):
                return model.encode(texts, normalize_embeddings=True, batch_size=64, convert_to_numpy=True).astype("float32")
            return "st", _encode_st
        except Exception as e:
            # torch/transformers/DLL 等失败则继续回退
            pass

    # 2) Fastembed(轻量.无 torch)
    if FastEmbedModel is not None:
        try:
            fe = FastEmbedModel(model_name="sentence-transformers/all-MiniLM-L6-v2")  # 多语表现也不错
            def _encode_fe(texts):
                # 返回生成器.需要堆叠; 向量已归一化
                vecs = [v for v in fe.embed(texts)]
                import numpy as np
                arr = np.asarray(vecs, dtype="float32")
                # fastembed 通常已是 L2 归一;稳妥起见再归一遍
                norms = np.linalg.norm(arr, axis=1, keepdims=True) + 1e-12
                return (arr / norms).astype("float32")
            return "fastembed", _encode_fe
        except Exception:
            pass

    # 3) 兜底:TF-IDF 近似(非真正“语义向量”.但能用)
    if TfidfVectorizer is not None:
        def _encode_tfidf(texts, _cache={"vec": None, "mat": None}):
            if _cache["vec"] is None:
                vec = TfidfVectorizer(max_features=50000, ngram_range=(1,2))
                mat = vec.fit_transform(texts)
                _cache["vec"], _cache["mat"] = vec, mat
            else:
                vec = _cache["vec"]
                mat = vec.transform(texts)
            # 为了接口一致.这里返回 dense(注意内存);小规模语料可接受
            import numpy as np
            arr = mat.astype("float32").toarray()
            norms = np.linalg.norm(arr, axis=1, keepdims=True) + 1e-12
            return (arr / norms).astype("float32")
        return "tfidf", _encode_tfidf

    raise RuntimeError("缺少向量后端:请安装 sentence-transformers 或 fastembed 或 scikit-learn。")

def _load_index(project_id: int):
    np, faiss, *_ = _lazy_import_vec()
    idx_path, map_path, vec_path = _index_paths(project_id)
    mapping = []
    if os.path.exists(map_path):
        with open(map_path, "r", encoding="utf-8") as f:
            mapping = json.load(f)
    # FAISS
    if faiss is not None and os.path.exists(idx_path):
        index = faiss.read_index(idx_path)
        return ("faiss", index, mapping, None)
    # 回退:.npy
    if os.path.exists(vec_path):
        vecs = np.load(vec_path).astype("float32")
        return ("fallback", None, mapping, vecs)
    return ("none", None, mapping, None)

def _save_index(project_id: int, mode: str, index, mapping, vecs=None):
    np, faiss, *_ = _lazy_import_vec()
    idx_path, map_path, vec_path = _index_paths(project_id)
    if mode == "faiss" and index is not None:
        faiss.write_index(index, idx_path)
    elif mode == "fallback" and vecs is not None:
        np.save(vec_path, vecs.astype("float32"))
    with open(map_path, "w", encoding="utf-8") as f:
        json.dump(mapping, f, ensure_ascii=False, indent=2)

def build_project_vector_index(project_id: int, use_src: bool = True, use_tgt: bool = True):
    """
    为指定项目构建/更新向量索引（.cache_index）:
    - 默认同时写入 src/tgt.但推荐检索时优先用 tgt
    - 写入映射时补充 domain/title/lang_pair/side
    - 统一 L2 归一化.FAISS 用内积检索≈余弦
    - 去重:同一 (corpus_id, side) 不重复
    返回: {"added": 新增条数, "total": 索引总条数}
    """
    import numpy as _np
    np, faiss, *_ = _lazy_import_vec()
    backend, encode = get_embedder()

    # 仅索引当前项目；需要 domain 用于后续“同领域”过滤
    rows = cur.execute("""
        SELECT c.id, IFNULL(c.src_text,''), IFNULL(c.tgt_text,''), 
               IFNULL(c.title,''), IFNULL(c.lang_pair,''), 
               IFNULL(c.project_id,0), IFNULL(i.domain,'')
        FROM corpus c
        LEFT JOIN items i ON i.id = c.project_id
        WHERE c.project_id = ?
        ORDER BY c.id ASC
    """, (int(project_id),)).fetchall()

    # 组装文本与元数据（优先译文侧；可同时写入 src/tgt）
    texts, metas = [], []
    for cid, s, t, ttl, lp, pj, dom in rows:
        s = (s or "").strip()
        t = (t or "").strip()
        if use_tgt and t:
            texts.append(t)
            metas.append({
                "corpus_id": cid,
                "project_id": pj,
                "domain": dom or "",
                "title": ttl,
                "lang_pair": lp,
                "side": "tgt"
            })
        if use_src and s:
            texts.append(s)
            metas.append({
                "corpus_id": cid,
                "project_id": pj,
                "domain": dom or "",
                "title": ttl,
                "lang_pair": lp,
                "side": "src"
            })

    if not texts:
        return {"added": 0, "total": 0}

    # 编码并 L2 归一化（IP≈cos）
    new_vecs = encode(texts)
    if hasattr(new_vecs, "toarray"):  # 兼容稀疏
        new_vecs = new_vecs.toarray()
    new_vecs = _np.asarray(new_vecs, dtype="float32")
    new_vecs = new_vecs / (_np.linalg.norm(new_vecs, axis=1, keepdims=True) + 1e-12)

    # 载入已有索引
    mode, index, mapping, vecs = _load_index(project_id)
    mapping = list(mapping or [])

    # —— 去重:已有映射的 (corpus_id, side) 不再重复加入
    seen = {(m.get("corpus_id"), m.get("side")) for m in mapping}
    keep_idx = []
    for i, m in enumerate(metas):
        key = (int(m["corpus_id"]), m["side"])
        if key not in seen:
            keep_idx.append(i)
            seen.add(key)

    if not keep_idx:
        # 没有新增
        total = (index.ntotal if (faiss is not None and mode == "faiss" and index is not None)
                 else (vecs.shape[0] if isinstance(vecs, _np.ndarray) else len(mapping)))
        return {"added": 0, "total": int(total)}

    metas = [metas[i] for i in keep_idx]
    new_vecs = new_vecs[keep_idx, :]

    # —— 写入索引:优先 FAISS；否则 .npy 回退
    if faiss is not None and backend in ("st", "fastembed"):
        dim = int(new_vecs.shape[1])
        if mode != "faiss" or index is None:
            index = faiss.IndexFlatIP(dim)   # 余弦等价（向量已归一）
            mapping = []
        # 若已有 mapping/索引.但之前维度不一致.重建
        if index.d != dim:
            index = faiss.IndexFlatIP(dim)
            mapping = []
        index.add(new_vecs)
        mapping.extend(metas)
        _save_index(project_id, "faiss", index, mapping)
        total = int(index.ntotal)
    else:
        # 回退:拼接矩阵
        if vecs is None:
            vecs = new_vecs
            mapping = metas
        else:
            # 维度不一致则重建
            if vecs.shape[1] != new_vecs.shape[1]:
                vecs = new_vecs
                mapping = metas
            else:
                vecs = _np.concatenate([vecs, new_vecs], axis=0)
                mapping.extend(metas)
        _save_index(project_id, "fallback", None, mapping, vecs=vecs)
        total = int(vecs.shape[0])

    return {"added": len(keep_idx), "total": total}

# =========================
# 语义召回(支持范围:project/domain/all）
# =========================
def semantic_retrieve(project_id: int,
                      query_text: str,
                      topk: int = 20,
                      scope: str = "project",
                      min_char: int = 3):
    """
    语料库语义召回（自动切句版）
    返回 [(score, meta, txt)]
    兼容：沿用外部向量索引 (.cache_index/...) 与你现有的 mapping 结构。
    """

    q = (query_text or "").strip()
    if len(q) < min_char:
        return []

    # --- 工具：切句（优先用你已有的 split_sents）
    def _split(text: str) -> list[str]:
        try:
            if 'split_sents' in globals():
                segs = split_sents(text, lang_hint="auto")
                return [s for s in segs if s and len(s.strip()) >= min_char]
        except Exception:
            pass
        import re
        segs = re.split(r"(?<=[\.\!\?;。！？；])\s*", text)
        return [s.strip() for s in segs if s and len(s.strip()) >= min_char]

    # --- 取 embedder / 索引
    backend, encode = get_embedder()  # ("st"/"fastembed"/"tfidf", encoder)
    mode, index, mapping, vecs = _load_index(project_id)  # ("faiss"/"fallback"/"none", idx, list, np.array/None)
    if mode == "none" or not mapping:
        return []

    # --- 项目领域信息（scope=domain 时用）
    def _get_domain_for_proj(pid: int) -> str:
        row = cur.execute("SELECT IFNULL(domain,'') FROM items WHERE id=?", (pid,)).fetchone()
        return (row[0] if row else "") or ""
    proj_domain = _get_domain_for_proj(project_id) if scope == "domain" else ""

    # --- 单次搜索：给一个文本片段，返回候选 [(score, idx)]
    def _search_once(piece: str, per_segment_k: int):
        try:
            import numpy as np
            qv = encode([piece]).astype("float32")
            qv = qv[0] if qv.ndim == 2 else qv
            if mode == "faiss" and index is not None:
                D, I = index.search(qv.reshape(1, -1), min(per_segment_k, len(mapping)))
                return list(zip(D[0].tolist(), I[0].tolist()))
            else:
                if vecs is None:
                    return []
                sims = (vecs @ qv.reshape(-1, 1)).ravel()  # 内积（已归一化≈cos）
                order = sims.argsort()[::-1][:min(per_segment_k, sims.shape[0])]
                return [(float(sims[i]), int(i)) for i in order]
        except Exception:
            return []

    # --- 片段化检索：按片段平均分配候选预算
    parts = _split(q)
    if not parts:
        parts = [q]

    # 基于片段数动态调整每片召回预算；乘以 2 预留合并损耗
    import math
    per_k = max(5, math.ceil((topk * 2) / max(1, len(parts))))

    # 汇总：同一个 corpus_id 只保留最高得分
    best_by_cid = {}  # cid -> (score, meta, txt)

    # 简单的权重/惩罚
    PROJECT_BONUS = 0.05
    DOMAIN_BONUS = 0.02
    SHORT_LEN = 6
    SHORT_PENALTY = 0.5

    for piece in parts:
        piece = piece.strip()
        if len(piece) < min_char:
            continue

        hits = _search_once(piece, per_k)
        # 对短片段降权
        length_factor = SHORT_PENALTY if len(piece) < SHORT_LEN else 1.0

        for raw_sc, mi in hits:
            if mi < 0 or mi >= len(mapping):
                continue
            meta = dict(mapping[mi])  # 可能含: corpus_id/side/title/lang_pair
            cid = int(meta.get("corpus_id") or -1)
            if cid < 0:
                continue

            # 回表拿项目/领域/文本
            row = cur.execute(
                "SELECT IFNULL(project_id,0), IFNULL(domain,''), IFNULL(src_text,''), IFNULL(tgt_text,'') "
                "FROM corpus WHERE id=?", (cid,)
            ).fetchone()
            if not row:
                continue
            pid_c, dom_c, s_txt, t_txt = row
            meta["project_id"] = pid_c
            meta["domain"] = dom_c or ""

            # scope 过滤
            if scope == "project" and int(pid_c or 0) != int(project_id):
                continue
            if scope == "domain":
                # 允许：同项目；或不同项目但 domain 相同
                if (int(pid_c or 0) != int(project_id)) and ((dom_c or "") != proj_domain):
                    continue
            # scope == "all": 不过滤

            txt = (t_txt or s_txt or "").strip()
            if not txt:
                continue

            # 计算最终分数：原始相似度 × 片段长度因子 + 项目/领域奖励
            sc = float(raw_sc) * float(length_factor)
            if int(pid_c or 0) == int(project_id):
                sc += PROJECT_BONUS
            elif scope in ("domain", "all") and (dom_c or "") == proj_domain and proj_domain:
                sc += DOMAIN_BONUS

            # 合并：取同一 cid 的最高分
            prev = best_by_cid.get(cid)
            if (prev is None) or (sc > prev[0]):
                best_by_cid[cid] = (sc, meta, txt)

    if not best_by_cid:
        return []

    # 归一排序并截断
    merged = sorted(best_by_cid.values(), key=lambda x: x[0], reverse=True)[:topk]
    return merged

def semantic_consistency_report(project_id: int, blocks_src: list, blocks_tgt: list, term_map: dict, topk: int = 3, thr: float = 0.70):
    """
    返回 DataFrame:段号、最相似参考得分、是否低于阈值、未遵守术语条目等
    """
    emb = get_embedder()

    # 仅对“译文侧”做参考(更贴近人工审校)
    hits_all = []
    for i, (s, t) in enumerate(zip(blocks_src, blocks_tgt), 1):
        hits = semantic_retrieve(project_id, t, topk=topk)  # 用译文去检索参考译文
        top_score = 0.0
        # 只看 tgt 侧
        for sc, meta, txt in hits:
            if meta.get("side") == "tgt":
                top_score = sc
                break

        # 术语遵守:粗略检查“目标译名是否出现在译文中”
        violated = []
        for src_term, tgt_term in (term_map or {}).items():
            if src_term in (s or "") and tgt_term and (tgt_term not in (t or "")):
                violated.append(f"{src_term}->{tgt_term}")

        hits_all.append({
            "段号": i,
            "相似参考得分": round(top_score, 2),
            "低于阈值": (top_score < thr),
            "未遵守术语": ", ".join(violated) if violated else ""
        })

    return pd.DataFrame(hits_all)

def search_semantic(project_id, query_text, topk: int = 20, scope: str = "project"):
    """
    统一封装一层，内部直接调用 semantic_retrieve：
    - project_id: 当前项目ID（int 或 "global"）
    - query_text: 用户在界面输入的查询文本
    - topk: 返回多少条结果
    - scope: "project" / "domain" / "all"

    返回值：[(score, meta, txt), ...] —— 直接复用 semantic_retrieve 的格式
    """
    # 允许 UI 层传 "global" 或 None 进来，这里做个简单兼容
    if project_id in (None, "", "global"):
        # 你自己的语义检索实现是基于项目ID的，
        # 如果你想做“全局搜索”，可以约定一个特殊值，比如 0，
        # 再在 semantic_retrieve 里根据 scope="all" 走全库检索。
        pid = 0
    else:
        pid = int(project_id)

    return semantic_retrieve(
        project_id=pid,
        query_text=query_text,
        topk=topk,
        scope=scope
    )

# ========== 路径/DB ==========
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
DB_PATH = os.path.join(BASE_DIR, "kb.db")
PROJECT_DIR = os.path.join(BASE_DIR, "projects")
os.makedirs(PROJECT_DIR, exist_ok=True)

conn = sqlite3.connect(DB_PATH, check_same_thread=False)
cur = conn.cursor()

def ensure_domain_columns_and_backfill(conn, cur, corpus_table="corpus"):
    # items.domain
    cols = [r[1] for r in cur.execute("PRAGMA table_info(items)").fetchall()]
    if "domain" not in cols:
        cur.execute("ALTER TABLE items ADD COLUMN domain TEXT;")
        conn.commit()

    # 语料表 domain（如果你用 corpus_main 就把参数改成 corpus_main）
    cols = [r[1] for r in cur.execute(f"PRAGMA table_info({corpus_table})").fetchall()]
    if "domain" not in cols:
        cur.execute(f"ALTER TABLE {corpus_table} ADD COLUMN domain TEXT;")
        conn.commit()

    # 回填:用 items.domain 补 corpus.domain（有 project_id 的行）
    try:
        cur.execute(f"""
            UPDATE {corpus_table}
            SET domain = (
              SELECT i.domain FROM items i WHERE i.id = {corpus_table}.project_id
            )
            WHERE domain IS NULL AND project_id IS NOT NULL;
        """)
        conn.commit()
    except Exception:
        pass

# 调用（老库表名是 corpus）:
ensure_domain_columns_and_backfill(conn, cur, corpus_table="corpus")
# 若你已经切到 corpus_main / corpus_vec:
# ensure_domain_columns_and_backfill(conn, cur, corpus_table="corpus_main")

try:
    cur.execute("CREATE INDEX IF NOT EXISTS idx_term_ext_project ON term_ext(project_id)")
    conn.commit()
except Exception as e:
    print("索引创建跳过:", e)

def _has_col(table: str, col: str) -> bool:
    cur.execute(f"PRAGMA table_info({table})")
    return any(r[1] == col for r in cur.fetchall())

def ensure_col(table: str, col: str, col_type: str):
    cur.execute(f"PRAGMA table_info({table})")
    cols = {r[1] for r in cur.fetchall()}
    if col not in cols:
        cur.execute(f"ALTER TABLE {table} ADD COLUMN {col} {col_type}")

# ==== 术语加载:项目优先.缺省用全局 ====
def load_terms_for_project(cur, project_id: int | None) -> dict[str, str]:
    """
    返回 {source_term: target_term}。加载顺序:
      1) 全局术语(project_id IS NULL)
      2) 指定项目术语(覆盖同名源词)
    """
    out: dict[str, str] = {}
    # 全局
    for s, t in cur.execute("""
        SELECT source_term, target_term FROM term_ext WHERE project_id IS NULL
    """).fetchall():
        s = (s or "").strip(); t = (t or "").strip()
        if s and t:
            out[s] = t
    # 项目
    if project_id is not None:
        for s, t in cur.execute("""
            SELECT source_term, target_term FROM term_ext WHERE project_id=?
        """, (int(project_id),)).fetchall():
            s = (s or "").strip(); t = (t or "").strip()
            if s and t:
                out[s] = t  # 覆盖全局
    return out

# —— 建表
cur.execute("""
CREATE TABLE IF NOT EXISTS items (
    id INTEGER PRIMARY KEY,
    title TEXT NOT NULL,
    body TEXT,
    tags TEXT,
    domain TEXT,
    type TEXT,
    created_at TEXT DEFAULT (datetime('now'))
);
""")

cur.execute("""
CREATE TABLE IF NOT EXISTS item_ext (
    id INTEGER PRIMARY KEY,
    item_id INTEGER,
    src_path TEXT,
    FOREIGN KEY(item_id) REFERENCES items(id)
);
""")

cur.execute("""
CREATE TABLE IF NOT EXISTS term_ext (
    id INTEGER PRIMARY KEY,
    source_term TEXT NOT NULL,
    target_term TEXT,
    domain TEXT,
    project_id INTEGER,
    strategy TEXT,
    example TEXT,
    category TEXT,
    FOREIGN KEY(project_id) REFERENCES items(id)
);
""")

cur.execute("""
CREATE TABLE IF NOT EXISTS trans_ext (
    id INTEGER PRIMARY KEY,
    project_id INTEGER,
    src_text TEXT,
    tgt_text TEXT,
    mode TEXT,
    segment_count INTEGER,
    created_at TEXT DEFAULT (datetime('now')),
    FOREIGN KEY(project_id) REFERENCES items(id)
);
""")

cur.execute("""
CREATE TABLE IF NOT EXISTS corpus (
    id INTEGER PRIMARY KEY,
    project_id INTEGER,
    text TEXT,
    lang TEXT,
    source TEXT,
    created_at TEXT DEFAULT (datetime('now')),
    FOREIGN KEY(project_id) REFERENCES items(id)
);
""")
conn.commit()


# —— 兜底补列
for t, cols in {
    "items": [("type","TEXT"),("tags","TEXT"),("scene","TEXT"),("prompt","TEXT"),
              ("mode","TEXT"),("body","TEXT"),("created_at","TEXT"),("updated_at","TEXT"),("trans_type","TEXT")],
    "item_ext": [("src_path","TEXT")],
    "term_ext": [("domain","TEXT"),("project_id","INTEGER"),("strategy","TEXT"),
                 ("example","TEXT"),("note","TEXT"), ("category","TEXT")],
    "trans_ext": [("stats_json","TEXT"),("segments","INTEGER"),("term_hit_total","INTEGER")],
    "corpus": [("title","TEXT"),("project_id","INTEGER"),("lang_pair","TEXT"),("src_text","TEXT"),("tgt_text","TEXT"),("note","TEXT"),("created_at","TEXT")],
}.items():
    for c, tp in cols:
        ensure_col(t, c, tp)
cur.execute("UPDATE items SET type='project' WHERE IFNULL(type,'')=''")
cur.execute("UPDATE items SET created_at = COALESCE(created_at, strftime('%Y-%m-%d %H:%M:%S','now'))")
conn.commit()
ensure_col("term_ext", "example_vector_id", "INTEGER")

# --- 术语表字段兼容:缺少 project_id 时补建 ---
try:
    cur.execute("PRAGMA table_info(term_ext)")
    cols = [c[1] for c in cur.fetchall()]
    if "project_id" not in cols:
        cur.execute("ALTER TABLE term_ext ADD COLUMN project_id INTEGER")
        conn.commit()
except Exception as e:
    st.warning(f"术语表结构检查:{e}")

# ========== DeepSeek 参数/调用 ==========
def get_deepseek():
    """
    从 .streamlit/secrets.toml 读取:
    [deepseek]
    api_key="..."
    model="deepseek-chat"
    """
    try:
        ak = st.secrets["deepseek"]["api_key"]
        model = st.secrets["deepseek"].get("model", "deepseek-chat")
        return ak, model
    except Exception:
        return None, None

        # === 新增:术语提示 + 参考例句 ===
def _build_ref_context(project_id: int,
                       query_text: str,
                       topk: int = 20,
                       min_sim: float = 0.35,
                       prefer_side: str = "tgt",
                       scope: str = "project") -> str:
    """
    语义召回 → 组装参考块。
    scope: "project"|"domain"|"all"
    prefer_side: "tgt"|"src"|"both"
    返回一段可直接注入 Prompt 的参考文本。
    """
    try:
        hits = semantic_retrieve(project_id, query_text, topk=topk, scope=scope)  # 传入 scope
    except Exception as e:
        try:
            st.warning(f"参考检索失败:{e}")
        except Exception:
            pass
        return ""

    selected = [(sc, meta, txt) for (sc, meta, txt) in (hits or []) if (sc or 0) >= float(min_sim)]
    if not selected and hits:
        selected = [max(hits, key=lambda x: x[0])]

    ctx_parts, used = [], set()
    for idx, (sc, meta, txt) in enumerate(selected, 1):
        s_txt = (txt or "").strip()
        if not s_txt:
            continue
        key = s_txt[:120]
        if key in used:
            continue
        used.add(key)

        title = meta.get("domain", "") if isinstance(meta, dict) else ""
        side = meta.get("side", "tgt") if isinstance(meta, dict) else "tgt"

        if prefer_side == "both" and isinstance(meta, dict):
            row = cur.execute("SELECT src_text, tgt_text FROM corpus WHERE id=?", (meta.get("corpus_id"),)).fetchone()
            s0, t0 = (row or ["",""])
            snippet = f"例句 {idx} 原文:{(s0 or '').strip()}\n例句 {idx} 译文:{(t0 or '').strip()}"
            ctx_parts.append(f"[{title}] (sim={sc:.2f})\n{snippet}\n")
        else:
            tag = "参考译文" if side == "tgt" else "参考原文"
            snippet = s_txt.replace("\n", " ")[:400]
            ctx_parts.append(f"[{tag}·{title}] (sim={sc:.2f}) {snippet}")

        if sum(len(p) for p in ctx_parts) > 1800:
            break

    return "" if not ctx_parts else "(以下为相似语料.可参考术语与风格）\n" + "\n".join(ctx_parts)

def build_system_prompt(
    base_prompt: str,
    term_pairs: list,                 # 形如 [(src, tgt), ...]
    lang_pair: str,
    ref_context: str | None = None    # 外部已经准备好的语义召回/示例参考块
) -> str:
    """
    构造最终提示:翻译方向 + 术语约束 + 参考例句 + 通用指令 + 用户文本
    返回值:system_prompt + '\\n\\n' + user_prompt
    """

    # —— 翻译方向说明 ——
    lp_note = f"翻译方向:{lang_pair or '自动'}。"

    # —— 术语块(把 term_pairs → 文本提示) ——
    term_dict = {s: t for s, t in (term_pairs or []) if s and t}
    if term_dict:
        kb_lines = [f"- {s} -> {t}" for s, t in term_dict.items()]
        kb_block = "请严格遵循以下术语对照.不得改写或替换为近义表达:\n" + "\n".join(kb_lines)
        # 若你项目里已有 build_term_hint.就复用;没有也可删掉 term_hint 这行
        term_hint = build_term_hint(term_dict, lang_pair) or ""
    else:
        kb_block = ""
        term_hint = ""

    # —— 通用规则(可替换为你原先的 rules 文本) ——
    rules = (
        "通用指令:\n"
        "- 保持术语/专有名词一致.不随意增译/省译;\n"
        "- 数字、时间、人名、地名准确无误;\n"
        "- 若遇未登录术语.保持原文+括注(仅第一次出现)。\n"
    )

    # —— 参考例句(来自外部传入的 ref_context) ——
    ref_hint = f"参考例句(保持术语/风格一致):\n{ref_context}\n" if ref_context else ""

    # —— 目标语言 ——
    target_lang = "中文" if (lang_pair or "").startswith("英译中") else "英文"

    # —— System Prompt (系统提示) ——
    system_prompt = (
        "你是一名资深专业笔译:\n"
        "1) 准确传达术语与事实;2) 正式清晰;3) 不臆造;\n"
        "4) 专有名词保持一致;5) 在不改变术语的前提下优化通顺。\n"
        f"{lp_note}\n\n{rules}\n{kb_block}\n"
    ).strip()

    # —— User Prompt(用户提示.确保无条件赋值) ——
    user_prompt = (
        (term_hint or "") +
        (ref_hint or "") +
        f"请将以下文本翻译为{target_lang}:\n\n{(base_prompt or '').strip()}"
    )

    return system_prompt + "\n\n" + user_prompt

def detect_terms_simple(block: str, term_map: dict) -> dict:
    return {k: v for k, v in term_map.items() if k and (k in block)}
# -------- Glossary & Instruction helpers (放在 ds_translate 上方) --------
def build_term_hint(term_dict: dict, lang_pair: str, max_terms: int = 80) -> str:
    """
    将术语映射转成可读的“硬约束”规则文本.支持以下几种 term_dict 结构:
      { "contract": "合同" }
      { "contract": {"target":"合同", "pos":"NOUN", "usage_note":"法律语境"} }
      { "contract": ("合同", "NOUN") }   # 元组形式 (target, pos)
    空目标会被忽略;自动去重并最多输出 max_terms 条.避免提示过长。
    """
    lines = []
    seen = set()
    items = list(term_dict.items())[: max_terms * 2]  # 稍多取一些.过滤空后再截断

    for src, val in items:
        if src is None: 
            continue
        src = str(src).strip()
        if not src or src in seen:
            continue

        tgt, pos, note = None, None, None
        if isinstance(val, dict):
            tgt  = (val.get("target") or val.get("tgt") or "").strip()
            pos  = (val.get("pos") or "").strip() or None
            note = (val.get("usage_note") or val.get("note") or "").strip() or None
        elif isinstance(val, (list, tuple)) and len(val) >= 1:
            tgt = str(val[0]).strip()
            if len(val) >= 2:
                pos = (str(val[1]).strip() or None)
        else:
            tgt = str(val or "").strip()

        if not tgt:
            continue

        seen.add(src)
        if pos:
            line = f"- When '{src}' is a {pos}, translate it as '{tgt}'."
        else:
            line = f"- Translate '{src}' as '{tgt}'."
        if note:
            line += f" ({note})"
        lines.append(line)

    if not lines:
        return ""  # 没有可用术语就返回空串.不要干扰提示

    header = "GLOSSARY (STRICT):\n"
    return header + "\n".join(lines[:max_terms]) + "\n"


def build_instruction(lang_pair: str) -> str:
    """
    生成简洁的翻译指令。你也可以按项目风格再扩展。
    """
    lp = (lang_pair or "").replace(" ", "")
    if "中→英" in lp or "中->英" in lp or "zh" in lp.lower() and "en" in lp.lower():
        return (
            "Translate the source text from Chinese to English. "
            "Use a professional, natural style; follow the GLOSSARY (STRICT) exactly; "
            "preserve proper nouns and numbers; keep paragraph structure. "
            "Do not add explanations."
        )
    if "英→中" in lp or "英->中" in lp or "en" in lp.lower() and "zh" in lp.lower():
        return (
            "Translate the source text from English to Chinese. "
            "用专业、通顺、符合领域文体的中文表达;严格遵守上方 GLOSSARY (STRICT);"
            "专有名词、数字与计量单位保持准确;段落结构保持一致。不得添加解释。"
        )
    # 兜底
    return (
        "Translate the source text. Follow the GLOSSARY (STRICT) exactly. "
        "Keep the original structure and do not add explanations."
    )

def ds_translate(block: str, term_dict: dict, lang_pair: str, ak: str, model: str, ref_context: str = "") -> str:
    term_hint = build_term_hint(term_dict, lang_pair)  # 你现有的术语提示
    instr = build_instruction(lang_pair)   # type: ignore

    """
    使用 DeepSeek REST API 翻译一个文本块。term_dict 为 {源: 目标} 的映射.注入为强约束提示。
    """
    import requests

    if not block.strip():
        return ""

    if term_dict:
        term_lines = "\n".join([f"- {k} -> {v}" for k, v in term_dict.items()])
        term_hint = (
            "TERMINOLOGY:\n"
            "Use the following mappings EXACTLY and consistently. Do not invent alternatives.\n"
            f"{term_lines}\n"
        )
    else:
        term_hint = "TERMINOLOGY:\nEnsure consistent terminology; avoid paraphrasing fixed terms.\n"

    if lang_pair == "中译英":
        instr = "Translate the Chinese text into English with high fidelity and formal style."
    elif lang_pair == "英译中":
        instr = "将下列英文准确译为中文.语体正式、专业。"
    else:
        instr = "Translate accurately into the other language."

    system_msg = (
        "You are a senior professional translator. Prioritize accuracy, faithfulness, and consistent terminology. "
        "No hallucinations. If a term mapping is provided, follow it strictly."
    )
    
    user_msg = (
    f"{term_hint}"
    + (f"REFERENCE CONTEXT (use if relevant):\n{ref_context}\n\n" if ref_context else "")
    + "INSTRUCTION:\n" + instr + "\n"
    + "RESPONSE FORMAT:\n- Return ONLY the final translation text, no explanations, no backticks.\n\n"
    + "SOURCE:\n" + block
    )


    url = "https://api.deepseek.com/v1/chat/completions"
    headers = {"Authorization": f"Bearer {ak}", "Content-Type": "application/json"}
    payload = {
        "model": model or "deepseek-chat",
        "messages": [
            {"role": "system", "content": system_msg},
            {"role": "user", "content": user_msg},
        ],
        "temperature": 0.2,
    }

    for attempt in range(3):
        try:
            resp = requests.post(url, headers=headers, json=payload, timeout=60)
            if resp.status_code == 200:
                data = resp.json()
                return data["choices"][0]["message"]["content"].strip()
            else:
                txt = f"[DeepSeek {resp.status_code}] {resp.text}"
                if resp.status_code in (429, 500, 502, 503, 504) and attempt < 2:
                    time.sleep(1.5 * (attempt + 1))
                    continue
                return txt
        except Exception as e:
            if attempt < 2:
                time.sleep(1.5 * (attempt + 1))
                continue
            return f"[DeepSeek Request Error] {e}"
    return "[DeepSeek Error] Unknown failure."

def ds_extract_terms(text: str, ak: str, model: str, src_lang: str = "zh", tgt_lang: str = "en"):
    """
    用 DeepSeek 从文本中抽取术语对.返回 JSON 数组:
    [{"source_term":"...", "target_term":"...", "domain":"...", "strategy":"...", "example":"..."}]
    """
    import requests

    if not text.strip():
        return []

    system_msg = (
        "You are a terminology mining assistant. Extract high-value bilingual term pairs suitable for a project glossary. "
        "Return JSON array only. No extra text."
    )
    user_msg = f"""
Source language: {src_lang}
Target language: {tgt_lang}
任务:从给定文本中抽取双语术语条目.输出 JSON 数组。字段名与取值必须是中文。
字段定义:
- source_term: 源语(中文术语或专名)
- target_term: 译文(英文)
- domain: 领域.取值集合之一:["政治","经济","文化","文物","金融","法律","其他"]
- strategy: 翻译策略.取值集合之一:["直译","意译","转译","音译","省略","增译","规范化","其他"]
- example: 例句(原文中包含该术语的一句.尽量保留标点)

要求:
1) 仅输出 JSON.不要多余说明。
2) 同一术语重复时合并.选择最典型的例句。
3) 若无法判断 domain/strategy.填“其他”。

Text:
{text}
"""

    url = "https://api.deepseek.com/v1/chat/completions"
    headers = {"Authorization": f"Bearer {ak}", "Content-Type": "application/json"}
    payload = {
        "model": model or "deepseek-chat",
        "messages": [
            {"role": "system", "content": system_msg},
            {"role": "user", "content": user_msg},
        ],
        "temperature": 0.1,
    }
    try:
        r = requests.post(url, headers=headers, json=payload, timeout=60)
        r.raise_for_status()
        data = r.json()
        txt = data["choices"][0]["message"]["content"].strip()
        # 只保留 JSON 片段
        start = txt.find("[")
        end = txt.rfind("]")
        if start == -1 or end == -1:
            return []
        arr = json.loads(txt[start:end+1])
        out = []
        for o in arr:
            src = (o.get("source_term") or o.get("source") or "").strip()
            tgt = (o.get("target_term") or o.get("target") or "").strip()
            dom = (o.get("domain") or "").strip() or None
            strat = (o.get("strategy") or "").strip() or None
            ex = (o.get("example") or "").strip() or None
            if src:
                out.append({"source_term": src, "target_term": tgt, "domain": dom, "strategy": strat, "example": ex})
        return out
    except Exception:
        return []

# ========== 文件读写与导出 ==========
def read_source_file(path: str) -> str:
    if not path or not os.path.exists(path):
        return ""
    ext = os.path.splitext(path)[1].lower()
    if ext == ".txt":
        for enc in ["utf-8", "utf-8-sig", "gb18030", "gbk"]:
            try:
                with open(path, "r", encoding=enc) as f:
                    return f.read()
            except Exception:
                continue
        with open(path, "r", errors="ignore") as f:
            return f.read()
    elif ext == ".docx":
        try:
            from docx import Document
            doc = Document(path)
            return "\n".join(p.text for p in doc.paragraphs)
        except Exception:
            return ""
    elif ext == ".xlsx":
        try:
            xls = pd.ExcelFile(path)
            parts = []
            for sheet in xls.sheet_names:
                df = xls.parse(sheet)
                parts.append(df.astype(str).to_csv(sep=" ", index=False, header=False))
            return "\n".join(parts)
        except Exception:
            return ""
    else:
        # 兜底尝试文本读取
        try:
            with open(path, "r", encoding="utf-8") as f:
                return f.read()
        except Exception:
            return ""

def build_bilingual_lines(src_text: str, tgt_text: str):
    """
    用段落做对齐：
    - 每一段中文对应一段英文
    - 段内不再拆句（避免 CSV / Word 错位）
    """
    return pair_paragraphs(src_text, tgt_text)

def export_csv_bilingual(src_text: str, tgt_text: str) -> bytes:
    s, t = build_bilingual_lines(src_text, tgt_text)
    df = pd.DataFrame({"Source": s, "Target": t})
    buf = io.StringIO()
    df.to_csv(buf, index=False)
    return buf.getvalue().encode("utf-8-sig")

def export_docx_bilingual(src_text: str, tgt_text: str) -> bytes:
    try:
        from docx import Document
        from docx.oxml.ns import qn
    except Exception:
        st.error("缺少 python-docx.请先安装:pip install python-docx")
        return b""
    doc = Document()
    # 基础字体
    try:
        doc.styles['Normal'].font.name = 'Calibri'
        doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    except Exception:
        pass
    s, t = build_bilingual_lines(src_text, tgt_text)
    table = doc.add_table(rows=1, cols=2)
    hdr = table.rows[0].cells
    hdr[0].text = "Source"
    hdr[1].text = "Target"
    for a, b in zip(s, t):
        row = table.add_row().cells
        row[0].text = a
        row[1].text = b
    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()

def export_docx_inplace(src_path: str, tgt_text: str) -> bytes:
    """DOCX:在每个原段落下插入译文段(简洁版)"""
    try:
        from docx import Document
        from docx.shared import Pt
    except Exception:
        st.error("缺少 python-docx.请安装:pip install python-docx")
        return b""

    doc = Document(src_path)
    tgt_lines = [x for x in tgt_text.splitlines()]

    i = 0
    for p in doc.paragraphs:
        tr = tgt_lines[i] if i < len(tgt_lines) else ""
        if tr.strip():
            run = doc.add_paragraph().add_run(tr)
            try:
                run.italic = True
                run.font.size = Pt(10)
            except Exception:
                pass
        i += 1
    while i < len(tgt_lines):
        doc.add_paragraph(tgt_lines[i])
        i += 1

    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()
# ==== 工具:文件读取 / 分句 / 向量 / 对齐 / 索引 ====
import os, re, io, json, numpy as _np

def _lazy_import_doc_pdf():
    docx = pdfplumber = None
    try:
        import docx as _docx
        docx = _docx
    except Exception:
        pass
    try:
        import pdfplumber as _pdfplumber
        pdfplumber = _pdfplumber
    except Exception:
        pass
    return docx, pdfplumber

def read_any_text(path_or_bytes, ext):
    """返回纯文本(单语)。ext: 'txt/csv/xlsx/docx/pdf'"""
    import pandas as pd
    from pathlib import Path
    docx, pdfplumber = _lazy_import_doc_pdf()

    def _read_docx_plain(fp):
        if not docx: return ""
        doc = docx.Document(fp)
        paras = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        return "\n".join(paras)

    def _read_pdf_plain(fp):
        if not pdfplumber: return ""
        text_chunks = []
        with pdfplumber.open(fp) as pdf:
            for page in pdf.pages:
                t = (page.extract_text() or "").strip()
                if t: text_chunks.append(t)
        return "\n".join(text_chunks)

    if ext == "txt":
        if isinstance(path_or_bytes, (str, os.PathLike)):
            return open(path_or_bytes, "r", encoding="utf-8", errors="ignore").read()
        else:
            return path_or_bytes.getvalue().decode("utf-8","ignore")

    if ext == "csv":
        df = pd.read_csv(path_or_bytes if not isinstance(path_or_bytes,(str,os.PathLike)) else path_or_bytes, encoding="utf-8", keep_default_na=False)
        return "\n".join([str(x) for x in _np.ravel(df.values) if str(x).strip()])

    if ext == "xlsx":
        df = pd.read_excel(path_or_bytes if not isinstance(path_or_bytes,(str,os.PathLike)) else path_or_bytes)
        return "\n".join([str(x) for x in _np.ravel(df.values) if str(x).strip()])

    if ext == "docx":
        return _read_docx_plain(path_or_bytes)
    if ext == "pdf":
        return _read_pdf_plain(path_or_bytes)
    return ""

def read_bilingual_pairs(path_or_bytes, ext):
    """返回[(src,tgt)];支持:CSV/XLSX 两列;DOCX 表格两列;其余返回空"""
    docx, _ = _lazy_import_doc_pdf()
    pairs = []

    if ext in ("csv","xlsx"):
        df = (pd.read_csv(path_or_bytes, encoding="utf-8", keep_default_na=False)
              if ext=="csv" else pd.read_excel(path_or_bytes))
        cols = [c.strip().lower() for c in df.columns]
        # 尝试自动找两列
        if len(cols)>=2:
            c1, c2 = 0, 1
            pairs = [(str(df.iloc[i,c1]).strip(), str(df.iloc[i,c2]).strip())
                     for i in range(len(df))
                     if str(df.iloc[i,c1]).strip() or str(df.iloc[i,c2]).strip()]
    elif ext == "docx" and docx:
        doc = docx.Document(path_or_bytes)
        for tbl in doc.tables:
            # 只处理两列表
            if len(tbl.columns) >= 2:
                for r in tbl.rows:
                    c0 = r.cells[0].text.strip()
                    c1 = r.cells[1].text.strip() if len(r.cells)>1 else ""
                    if c0 or c1:
                        pairs.append((c0, c1))
    return [(s,t) for s,t in pairs if s or t]

# 段落切分
def split_paragraphs(text: str) -> list[str]:
    """
    段落切分（用于翻译 & 导出）：
    - 统一换行符
    - 以【至少一个空行】作为段落分隔
    - 段内保留句子，只去掉纯空行
    """
    text = (text or "").replace("\r\n", "\n").replace("\r", "\n")
    # 常见情况：用“一行一段”的稿子，实际上中间没有空行
    # 这种就按单行当作段落
    if "\n\n" not in text and "\n \n" not in text:
        lines = [ln.strip() for ln in text.split("\n")]
        return [ln for ln in lines if ln]

    # 正常：有空行分段
    parts = re.split(r"\n\s*\n+", text)
    paras = []
    for p in parts:
        p = p.strip()
        if not p:
            continue
        # 段内如果还有软换行，压成空格，避免导出时被拆成多行
        p = re.sub(r"\s*\n\s*", " ", p)
        paras.append(p)
    return paras

def pair_paragraphs(src_full: str, tgt_full: str) -> tuple[list[str], list[str]]:
    """
    根据全文中英，按“段落”配对：
    - 源文/译文各自做 split_paragraphs
    - 行数不一致时用空串补齐
    - 保证导出时一行中文对应一行英文
    """
    src_paras = split_paragraphs(src_full or "")
    tgt_paras = split_paragraphs(tgt_full or "")

    n = max(len(src_paras), len(tgt_paras))
    src_paras += [""] * (n - len(src_paras))
    tgt_paras += [""] * (n - len(tgt_paras))
    return src_paras, tgt_paras

# 预编译（可放全局）
_RE_WS = re.compile(r"[ \t\u00A0\u200B\u200C\u200D]+")
_RE_ZH_SENT = re.compile(r"(?<=[。！？；])\s*")           # 中文句末
_RE_EN_SENT = re.compile(r"(?<=[\.\?\!;:])\s+")          # 英文句末（放宽，不强制大写）
_RE_BLANK_PARA = re.compile(r"\n{2,}")                   # 空行分段

def _norm_text(text: str) -> str:
    t = (text or "").replace("\r\n", "\n").replace("\r", "\n").replace("\x0b", "\n")
    t = _RE_WS.sub(" ", t)
    t = re.sub(r"\n{3,}", "\n\n", t)  # 过多空行压到两个
    return t.strip()

def _is_zh(text: str) -> bool:
    # 简单判定：含有较多中文字符
    zh_hits = len(re.findall(r"[\u4e00-\u9fff]", text))
    en_hits = len(re.findall(r"[A-Za-z]", text))
    return zh_hits >= en_hits

def split_sents(
    text: str,
    lang_hint: str = "auto",
    min_char: int = 1,
    prefer_newline: bool = True,
    **kwargs,
):
    """
    统一的分句/分段函数：
    - 兼容旧调用：split_sents(text, lang="zh")
    - 支持新参数：prefer_newline=True 时，优先按换行切
    """
    # 兼容旧参数名 lang=
    lang = kwargs.get("lang", lang_hint)

    t = _norm_text(text)
    if not t:
        return []

    pieces = []

    # A) 若文本中有换行 & prefer_newline=True：先按行切，再在行内按句末细分
    if prefer_newline and ("\n" in t):
        lines = [ln.strip() for ln in t.split("\n") if ln.strip()]
        for ln in lines:
            if lang == "auto":
                cur_lang = "zh" if _is_zh(ln) else "en"
            else:
                cur_lang = lang

            if cur_lang.startswith(("zh", "cn")):
                sents = [s.strip() for s in _RE_ZH_SENT.split(ln) if s and s.strip()]
            else:
                sents = [s.strip() for s in _RE_EN_SENT.split(ln) if s and s.strip()]

            pieces.extend(sents if sents else [ln])
    else:
        # B) 没有换行或不偏好换行：整块按句末标点切
        if lang == "auto":
            cur_lang = "zh" if _is_zh(t) else "en"
        else:
            cur_lang = lang

        if cur_lang.startswith(("zh", "cn")):
            pieces = [s.strip() for s in _RE_ZH_SENT.split(t) if s and s.strip()]
        else:
            pieces = [s.strip() for s in _RE_EN_SENT.split(t) if s and s.strip()]

        if not pieces:
            pieces = [t]

    # 过滤过短片段
    return [x for x in pieces if len(x) >= min_char]

# 兼容旧函数名
split_sentences = split_sents

def smart_split_blocks(text: str, max_chars: int = 1200, lang_hint: str = "auto"):
    """
    先用 split_sents（按句/按行）切出基本单元，
    - 若总长度不超过 max_chars：每句/行单独作为一个 block；
    - 若总长度超过 max_chars：再按长度打包成大块，保证不拆句。
    """
    sents = split_sents(text, lang_hint=lang_hint, min_char=1, prefer_newline=True)
    if not sents:
        return []

    # 计算全部句子的总长度（加上少量换行）
    total_len = sum(len(s) for s in sents) + max(0, len(sents) - 1)

    # 情况一：整体不长，直接按句/行返回
    if total_len <= max_chars:
        return sents

    # 情况二：整体较长，再按长度打包
    blocks = []
    buf = ""

    for s in sents:
        s = s.strip()
        if not s:
            continue

        if not buf:
            buf = s
            continue

        if len(buf) + 1 + len(s) <= max_chars:
            buf = buf + "\n" + s
        else:
            blocks.append(buf)
            buf = s

    if buf:
        blocks.append(buf)

    return blocks

def split_blocks(text: str, max_len: int = 1200):
    blocks, curbuf = [], ""
    for line in text.splitlines(True):
        if len(curbuf) + len(line) > max_len:
            if curbuf:
                blocks.append(curbuf)
            curbuf = ""
        curbuf += line
    if curbuf:
        blocks.append(curbuf)
    return blocks
# =========================
# 术语提示 & 译后一致性检查
# =========================
def check_term_consistency(out_text: str, term_dict: dict, source_text: str = "") -> list:
    """
    译后的一致性粗检:如果源文包含术语键.但译文未出现对应值.则记录提醒。
    仅做最小代价的字符串级检查(不改动原译文)。
    返回形如 ["contract→合同", ...] 的列表;为空表示全部符合/不适用。
    """
    if not out_text or not term_dict:
        return []
    warns = []
    s = (source_text or "")[:2000]  # 限长度.防极端长文本
    out_low = out_text.lower()
    for k, v in term_dict.items():
        if not k or not v:
            continue
        # 如果源文包含该术语键(大小写忽略英文;中文直接包含)
        hit_src = (k.lower() in s.lower()) if any(ord(ch) < 128 for ch in k) else (k in s)
        if hit_src:
            # 译文是否出现目标术语(同理大小写容忍英文)
            ok = (v.lower() in out_low) if any(ord(ch) < 128 for ch in v) else (v in out_text)
            if not ok:
                warns.append(f"{k}→{v}")
    return warns

def _lazy_embedder():
    # 优先 sentence-transformers;失败退化到 TF-IDF
    try:
        from sentence_transformers import SentenceTransformer
        import numpy as np
        mdl = SentenceTransformer("sentence-transformers/paraphrase-multilingual-MiniLM-L12-v2")
        def _emb(texts):
            arr = mdl.encode(texts, normalize_embeddings=True)
            return arr.astype("float32")
        return _emb, "sbert"
    except Exception:
        from sklearn.feature_extraction.text import TfidfVectorizer
        import numpy as np
        def _emb(texts):
            vec = TfidfVectorizer(min_df=1).fit_transform(texts)
            # 归一化
            norms = _np.sqrt((vec.multiply(vec)).sum(axis=1)).A.ravel() + 1e-8
            vec = vec.multiply(1/norms[:,None])
            return vec
        return _emb, "tfidf"

from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity
import numpy as np

def align_semantic(src_sents, tgt_sents, max_jump=3):
    """简单贪心 + 滑窗的 1-1 句对齐.返回 [(src, tgt, score)]"""
    if not src_sents or not tgt_sents:
        return []

    emb, kind = _lazy_embedder()

    # === 优先使用 SBERT ===
    if kind == "sbert":
        E1 = emb(src_sents)
        E2 = emb(tgt_sents)
        sims = E1 @ E2.T  # (n, m)
    else:
        # === TF-IDF 回退:确保同一词表维度 ===
        vec = TfidfVectorizer(
            analyzer="char_wb",  # 字符 n-gram 对中英混合最稳
            ngram_range=(1, 2),
            min_df=1
        )
        combo = src_sents + tgt_sents
        X = vec.fit_transform(combo)
        n = len(src_sents)
        E1 = X[:n, :]
        E2 = X[n:, :]

        # 稀疏→相似度矩阵
        sims = cosine_similarity(E1, E2, dense_output=True)  # shape (n, m)

    # === 贪心对齐 ===
    i = j = 0
    n, m = len(src_sents), len(tgt_sents)
    pairs = []
    while i < n and j < m:
        j_min = max(0, j - max_jump)
        j_max = min(m, j + max_jump + 1)
        window = sims[i, j_min:j_max]
        if window.size == 0:
            break
        k = int(window.argmax())
        j_sel = j_min + k
        score = float(sims[i, j_sel])
        pairs.append((src_sents[i], tgt_sents[j_sel], score))
        i += 1
        j = j_sel + 1
    return pairs

# ===== 向量索引(FAISS优先.降级为 NumPy) =====
def _lazy_faiss():
    try:
        import faiss
        return faiss
    except Exception:
        return None

def save_semantic_index(index_dir, pid, texts, vectors):
    os.makedirs(index_dir, exist_ok=True)
    _np.save(os.path.join(index_dir, f"{pid}_texts.npy"), _np.array(texts, dtype=object))
    faiss = _lazy_faiss()
    if faiss is not None and isinstance(vectors, _np.ndarray):
        idx = faiss.IndexFlatIP(vectors.shape[1])
        idx.add(vectors)
        faiss.write_index(idx, os.path.join(index_dir, f"{pid}.faiss"))
    else:
        _np.save(os.path.join(index_dir, f"{pid}_vecs.npy"), vectors)

# ========== 术语库管理 ==========
def render_term_management(st, cur, conn, base_dir, key_prefix="term"):
    sk = make_sk(key_prefix)

    st.subheader("📘 术语库管理")
    sub_tabs = st.tabs(["查询与编辑", "批量导入 CSV", "统计与导出", "快速搜索", "批量挂接项目", "从历史提取术语", "分类管理"])

    # —— 查询与编辑
    with sub_tabs[0]:
        sk0 = lambda name: f"{key_prefix}_t0_{name}"

        c1, c2, c3, c4 = st.columns(4)
        with c1: kw = st.text_input("关键词(源/目标/例句)", "", key=sk("kw_example"))
        with c2: dom = st.text_input("领域", "", key=sk("dom"))
        with c3: strat = st.text_input("策略", "", key=sk("strat"))
        with c4: pid = st.text_input("项目ID过滤", "", key=sk("pid"))
        cat = st.text_input("分类(支持子串)", "", key=sk("cat"))

        # —— 兼容老库:如无 category 列则补列(已有会忽略)
        try:
            cur.execute("ALTER TABLE term_ext ADD COLUMN category TEXT;")
            conn.commit()
        except Exception:
            pass

        # 1) 以数据库真实列为准检测是否存在 category
        cols_db = [c[1].lower() for c in cur.execute("PRAGMA table_info(term_ext);").fetchall()]
        has_category = ("category" in cols_db)

        # 2) 拼 SQL(只在 DB 真的有该列时才 SELECT category)
        base_cols = "id, source_term, target_term, domain, project_id, strategy, example"
        sel_cols = base_cols + (", category" if has_category else "")
        sql = f"SELECT {sel_cols} FROM term_ext WHERE 1=1"
        params = []

        if kw:
            like = f"%{kw}%"
            sql += " AND (IFNULL(source_term,'') LIKE ? OR IFNULL(target_term,'') LIKE ? OR IFNULL(example,'') LIKE ?)"
            params += [like, like, like]

        if dom:
            sql += " AND IFNULL(domain,'') LIKE ?"
            params += [f"%{dom}%"]

        if strat:
            sql += " AND IFNULL(strategy,'') LIKE ?"
            params += [f"%{strat}%"]

        if pid and str(pid).isdigit():
            sql += " AND IFNULL(project_id,0) = ?"
            params += [int(pid)]

        if has_category and cat:
            sql += " AND IFNULL(category,'') LIKE ?"
            params += [f"%{cat}%"]

        sql += " ORDER BY source_term COLLATE NOCASE LIMIT 1000"

        # 3) 查询并构造 DataFrame(表头与实际列对齐)
        rows = cur.execute(sql, params).fetchall()
        headers = ["ID","源术语","目标术语","领域","项目ID","策略","例句"]
        if has_category:
            headers.append("分类")

        df = pd.DataFrame(rows, columns=headers)
        st.caption(f"当前查询返回:{len(df)} 条")

        # 4) 空数据就不渲染编辑器
        if df.empty:
            st.info("没有匹配的术语。")
        else:
            # 没有 sel 列则插入
            if "sel" not in df.columns:
                df.insert(0, "sel", False)

            # 动态构建列配置.只有当 DB 真有“分类”时才加入
            col_cfg = {
                "ID": st.column_config.NumberColumn("ID", disabled=True),
                "sel": st.column_config.CheckboxColumn("选择"),
                "源术语": "源术语",
                "目标术语": "目标术语",
                "领域": "领域",
                "项目ID": st.column_config.NumberColumn("项目ID", step=1, required=False),
                "策略": "策略",
                "例句": st.column_config.TextColumn("例句"),
            }
            if has_category:
                col_cfg["分类"] = st.column_config.TextColumn("分类")

            edited = st.data_editor(
                df,
                num_rows="dynamic",
                key=sk0("editor"),
                column_config=col_cfg,
            )

            c1, c2, c3 = st.columns([1, 1, 2])
            with c1:
                if st.button("💾 保存修改", type="primary", key=sk("save_terms")):
                    updated = inserted = 0
                    for _, row in edited.iterrows():
                        if pd.notna(row["ID"]):  # 更新
                            if has_category:
                                cur.execute("""
                                    UPDATE term_ext
                                    SET source_term=?, target_term=?, domain=?, project_id=?, strategy=?, example=?, category=?
                                    WHERE id=?;
                                """, (
                                    (row["源术语"] or "").strip(),
                                    (row["目标术语"] or None),
                                    (row["领域"] or None),
                                    (int(row["项目ID"]) if pd.notna(row["项目ID"]) else None),
                                    (row["策略"] or None),
                                    (row["例句"] or None),
                                    (row.get("分类") or None),
                                    int(row["ID"])
                                ))
                            else:
                                cur.execute("""
                                    UPDATE term_ext
                                    SET source_term=?, target_term=?, domain=?, project_id=?, strategy=?, example=?
                                    WHERE id=?;
                                """, (
                                    (row["源术语"] or "").strip(),
                                    (row["目标术语"] or None),
                                    (row["领域"] or None),
                                    (int(row["项目ID"]) if pd.notna(row["项目ID"]) else None),
                                    (row["策略"] or None),
                                    (row["例句"] or None),
                                    int(row["ID"])
                                ))
                            updated += cur.rowcount
                        else:  # 新增
                            if str(row["源术语"]).strip():
                                if has_category:
                                    cur.execute("""
                                        INSERT INTO term_ext (source_term, target_term, domain, project_id, strategy, example, category)
                                        VALUES (?, ?, ?, ?, ?, ?, ?)
                                    """, (
                                        (row["源术语"] or "").strip(),
                                        (row["目标术语"] or None),
                                        (row["领域"] or None),
                                        (int(row["项目ID"]) if pd.notna(row["项目ID"]) else None),
                                        (row["策略"] or None),
                                        (row["例句"] or None),
                                        (row.get("分类") or None)
                                    ))
                                else:
                                    cur.execute("""
                                        INSERT INTO term_ext (source_term, target_term, domain, project_id, strategy, example)
                                        VALUES (?, ?, ?, ?, ?, ?)
                                    """, (
                                        (row["源术语"] or "").strip(),
                                        (row["目标术语"] or None),
                                        (row["领域"] or None),
                                        (int(row["项目ID"]) if pd.notna(row["项目ID"]) else None),
                                        (row["策略"] or None),
                                        (row["例句"] or None),
                                    ))
                                inserted += 1
                    conn.commit()
                    st.success(f"✅ 已保存:更新 {updated} 条.新增 {inserted} 条。")
                    st.rerun()

                with c2:
                    cc2a, cc2b, cc2c = st.columns([1, 1, 2])
                    if "sel" not in edited.columns:
                        edited.insert(0, "sel", False)
                    if cc2a.button("全选", key=sk("sel_all")):
                        if "ID" in edited.columns:
                            all_ids = edited[pd.notna(edited["ID"])]["ID"].astype(int).tolist()
                            edited.loc[:, "sel"] = True
                            st.session_state[sk("selected_ids")] = set(all_ids)
                        st.rerun()
                    if cc2b.button("清空", key=sk("sel_clear")):
                        edited.loc[:, "sel"] = False
                        st.session_state[sk("selected_ids")] = set()
                        st.rerun()
                    if cc2c.button("🗑️ 删除已勾选", key=sk("del_sel")):
                        if "ID" in edited.columns:
                            to_delete = edited[(edited["sel"] == True) & pd.notna(edited["ID"])]["ID"].astype(int).tolist()
                        else:
                            to_delete = []
                        if not to_delete:
                            st.warning("未勾选任何记录")
                        else:
                            cur.executemany("DELETE FROM term_ext WHERE id=?", [(i,) for i in to_delete])
                            conn.commit()
                            st.success(f"🗑️ 已删除 {len(to_delete)} 条")
                            st.rerun()

                with c3:
                    proj_opts = cur.execute("SELECT id, title FROM items ORDER BY id DESC").fetchall()
                    proj_map = {"(不挂接/置空)": None, **{f"#{i} {t}": i for (i, t) in proj_opts}}

                    cc3a, cc3b = st.columns([2, 1])
                    target_proj_label = cc3a.selectbox("批量挂接到项目", list(proj_map.keys()), key=sk("bind_proj_sel"))
                    if cc3b.button("执行挂接", type="primary", key=sk("bind_apply")):
                        if "ID" in edited.columns:
                            to_update = edited[(edited["sel"] == True) & pd.notna(edited["ID"])]["ID"].astype(int).tolist()
                        else:
                            to_update = []
                        if not to_update:
                            st.warning("未勾选任何记录")
                        else:
                            pid_val = proj_map.get(target_proj_label)
                            q_marks = ",".join("?" for _ in to_update)
                            cur.execute(f"UPDATE term_ext SET project_id=? WHERE id IN ({q_marks})", (pid_val, *to_update))
                            conn.commit()
                            st.success(f"✅ 已挂接 {len(to_update)} 条到项目:{target_proj_label or '(空)'}")
                            st.rerun()
                    else:
                        st.info("暂无术语记录")

            st.markdown("### 单条新增 / 编辑")
            with st.form(sk0("term_edit")):
                col1, col2, col3 = st.columns(3)
                with col1:
                    rid_edit = st.text_input("要编辑的记录 ID(留空则新增)", "", key=sk("rid_edit"))
                    source_term = st.text_input("源语言术语(必填)*", key=sk("source_term"))
                    target_term = st.text_input("目标语言术语", key=sk("target_term"))
                with col2:
                    domain = st.text_input("领域", key=sk("domain"))
                    project_id = st.text_input("项目ID(可空)", key=sk("project_id"))
                    strategy = st.text_input("策略(直译/意译/转译/音译/省略/增译/规范化…)", key=sk("strategy"))
                with col3:
                    example = st.text_area("例句", height=80, key=sk("example"))
                    category = st.text_input("分类(可选)", key=sk("category"))

                b1, b2 = st.columns(2)
                add = b1.form_submit_button("保存(新增或更新)")
                delbtn = b2.form_submit_button("删除(按 ID)")

            if add:
                if not source_term.strip():
                    st.error("源术语必填")
                else:
                    if rid_edit and rid_edit.isdigit():
                        if _has_col("term_ext", "category"):
                            cur.execute("""
                            UPDATE term_ext
                            SET source_term=?, target_term=?, domain=?, project_id=?, strategy=?, example=?, category=?
                            WHERE id=?;
                            """, (source_term.strip(), target_term.strip() or None, domain or None,
                                int(project_id) if project_id.isdigit() else None, strategy or None, example or None,
                                category or None, int(rid_edit)))
                        else:
                            cur.execute("""
                            UPDATE term_ext
                            SET source_term=?, target_term=?, domain=?, project_id=?, strategy=?, example=?
                            WHERE id=?;
                            """, (source_term.strip(), target_term.strip() or None, domain or None,
                                int(project_id) if project_id.isdigit() else None, strategy or None, example or None,
                                int(rid_edit)))
                        conn.commit(); st.success("✅ 已更新"); st.rerun()
                    else:
                        if _has_col("term_ext", "category"):
                            cur.execute("""
                            INSERT INTO term_ext (source_term, target_term, domain, project_id, strategy, example, category)
                            VALUES (?, ?, ?, ?, ?, ?, ?)
                            """, (
                                source_term.strip(),
                                (target_term.strip() or None) if target_term else None,
                                (domain or None),
                                (int(project_id) if project_id.isdigit() else None) if project_id else None,
                                (strategy or None),
                                (example or None),
                                (category or None)
                            ))
                        else:
                            cur.execute("""
                            INSERT INTO term_ext (source_term, target_term, domain, project_id, strategy, example)
                            VALUES (?, ?, ?, ?, ?, ?)
                            """, (
                                source_term.strip(),
                                (target_term.strip() or None) if target_term else None,
                                (domain or None),
                                (int(project_id) if project_id.isdigit() else None) if project_id else None,
                                (strategy or None),
                                (example or None),
                            ))
                        conn.commit(); st.success("✅ 已新增"); st.rerun()

            if delbtn:
                if rid_edit and rid_edit.isdigit():
                    cur.execute("DELETE FROM term_ext WHERE id=?", (int(rid_edit),))
                    conn.commit(); st.success("🗑️ 已删除"); st.rerun()
                else:
                    st.error("请填写要删除的 ID")

    # —— 批量导入 CSV(增强版:列名规范化 / 动态带或不带 category / 去重或Upsert / 逐行容错)
    with sub_tabs[1]:
        sk1 = lambda n: f"{key_prefix}_t1_{n}"
        st.caption("CSV 推荐列:source_term, target_term, domain, project_id, strategy, example(可选:category / 分类)")
        up = st.file_uploader("上传 CSV 文件", type=["csv"], key=sk1("csv"))

        # 小工具:补列 + 唯一索引 + 规范函数
        def _ensure_col(table, col, type_):
            try:
                cur.execute(f"ALTER TABLE {table} ADD COLUMN {col} {type_};")
                conn.commit()
            except Exception:
                pass

        def _ensure_unique_index():
            try:
                cur.execute("""
                CREATE UNIQUE INDEX IF NOT EXISTS ux_term_proj
                ON term_ext(LOWER(TRIM(source_term)), IFNULL(project_id,-1));
                """)
                conn.commit()
            except Exception:
                pass

        def _norm_cols(df):
            # 列名:去BOM/两端空格 -> 小写 -> 替换空格为下划线 -> 中英列名映射
            df.columns = [str(c).replace("\ufeff","").strip() for c in df.columns]
            df.columns = [c.lower().replace(" ", "_") for c in df.columns]
            mapping = {
                "源术语": "source_term", "source": "source_term",
                "目标术语": "target_term", "target": "target_term",
                "领域": "domain",
                "项目id": "project_id", "项目_id": "project_id",
                "策略": "strategy",
                "例句": "example",
                "分类": "category",
            }
            df = df.rename(columns={k.lower(): v for k, v in mapping.items()})
            return df

        def _norm(v):
            if v is None: return None
            v = str(v).strip().replace("\u3000", " ")
            return v if v else None

        def _to_int(v):
            try:
                return int(v) if v is not None and str(v).strip() != "" else None
            except Exception:
                return None

        if up is not None:
            try:
                df_up = pd.read_csv(up, encoding="utf-8-sig")
            except Exception:
                df_up = pd.read_csv(up, encoding="utf-8", errors="ignore")

            df_up = _norm_cols(df_up)
            st.write("检测到列:", list(df_up.columns))
            render_table(df_up.head(10), key=sk("csv_preview"))

            # DB 侧确保有 category 列(兼容老库;若已有会忽略)
            _ensure_col("term_ext", "category", "TEXT")

            # 侧边选项
            c1, c2, c3 = st.columns(3)
            with c1:
                dedup = st.checkbox("去重(源术语+项目ID)", value=True, key=sk1("dedup"))
            with c2:
                use_upsert = st.checkbox("已存在则更新(Upsert)", value=False, key=sk1("upsert"))
            with c3:
                skip_empty = st.checkbox("跳过空译文", value=False, key=sk1("skip_empty"))

            # Upsert 需要唯一索引
            if use_upsert:
                _ensure_unique_index()

            if st.button("导入术语库", key=sk1("import_btn")):
                # 是否包含 category 列(以CSV为准;DB已有不强制CSV必须有)
                has_category_col = ("category" in df_up.columns)

                # 去重缓存(仅在非Upsert模式下使用)
                existing = set()
                if dedup and not use_upsert:
                    rows_exist = cur.execute("""
                        SELECT LOWER(TRIM(source_term)), IFNULL(project_id,-1)
                        FROM term_ext
                    """).fetchall()
                    existing = set(rows_exist)

                ins = skp = upd = 0
                errors = []

                for idx, row in df_up.iterrows():
                    src = _norm(row.get("source_term"))
                    if not src:
                        skp += 1
                        continue

                    tgt = _norm(row.get("target_term"))
                    if skip_empty and not tgt:
                        skp += 1
                        continue

                    dom = _norm(row.get("domain"))
                    pid = _to_int(row.get("project_id"))
                    stg = _norm(row.get("strategy"))
                    exa = _norm(row.get("example"))
                    cat = _norm(row.get("category")) if has_category_col else None

                    try:
                        if use_upsert:
                            # Upsert 分支(需要唯一索引)
                            if has_category_col:
                                cur.execute("""
                                INSERT INTO term_ext (source_term, target_term, domain, project_id, strategy, example, category)
                                VALUES (?,?,?,?,?,?,?)
                                ON CONFLICT(LOWER(TRIM(source_term)), IFNULL(project_id,-1))
                                DO UPDATE SET
                                    target_term=COALESCE(excluded.target_term, term_ext.target_term),
                                    domain     =COALESCE(excluded.domain,      term_ext.domain),
                                    strategy   =COALESCE(excluded.strategy,    term_ext.strategy),
                                    example    =COALESCE(excluded.example,     term_ext.example),
                                    category   =COALESCE(excluded.category,    term_ext.category);
                                """, (src, tgt, dom, pid, stg, exa, cat))
                            else:
                                cur.execute("""
                                INSERT INTO term_ext (source_term, target_term, domain, project_id, strategy, example)
                                VALUES (?,?,?,?,?,?)
                                ON CONFLICT(LOWER(TRIM(source_term)), IFNULL(project_id,-1))
                                DO UPDATE SET
                                    target_term=COALESCE(excluded.target_term, term_ext.target_term),
                                    domain     =COALESCE(excluded.domain,      term_ext.domain),
                                    strategy   =COALESCE(excluded.strategy,    term_ext.strategy),
                                    example    =COALESCE(excluded.example,     term_ext.example);
                                """, (src, tgt, dom, pid, stg, exa))
                            upd += 1  # 计为“处理成功”.不区分新旧
                        else:
                            # 非 Upsert:去重(源术语+项目ID)
                            key = (src.lower(), pid if pid is not None else -1)
                            if dedup and key in existing:
                                skp += 1
                                continue

                            if has_category_col:
                                cur.execute("""
                                    INSERT INTO term_ext (source_term, target_term, domain, project_id, strategy, example, category)
                                    VALUES (?,?,?,?,?,?,?)
                                """, (src, tgt, dom, pid, stg, exa, cat))
                            else:
                                cur.execute("""
                                    INSERT INTO term_ext (source_term, target_term, domain, project_id, strategy, example)
                                    VALUES (?,?,?,?,?,?)
                                """, (src, tgt, dom, pid, stg, exa))
                            ins += 1
                            if dedup:
                                existing.add(key)

                    except Exception as e:
                        errors.append((idx+1, src, str(e)))
                        skp += 1
                        continue

                conn.commit()

                # 结果提示
                if use_upsert:
                    st.success(f"✅ 已处理 {ins+upd} 条(其中可能含新增+更新).跳过 {skp} 条。")
                else:
                    st.success(f"✅ 新增 {ins} 条.跳过 {skp} 条。")

                if errors:
                    with st.expander("❗ 行级错误明细(不影响其他行写入)", expanded=False):
                        for i, s, e in errors:
                            st.write(f"第 {i} 行({s}):{e}")


    # —— 统计与导出
    with sub_tabs[2]:
        sk2 = lambda n: f"{key_prefix}_t2_{n}"
        st.markdown("#### 术语统计")
        df_stats = pd.read_sql_query("SELECT strategy, domain, category FROM term_ext WHERE source_term IS NOT NULL", conn)
        if df_stats.empty:
            st.info("术语库为空.请先添加或导入")
        else:
            strat_count = df_stats["strategy"].fillna("未标注").value_counts().reset_index()
            strat_count.columns = ["strategy","count"]
            render_table(strat_count, hide_index=True, key=sk2("strat_tbl"))

            dom_count = df_stats["domain"].fillna("未标注").value_counts().reset_index()
            dom_count.columns = ["domain","count"]
            render_table(dom_count, hide_index=True, key=sk2("dom_tbl"), editable=True)

        st.markdown("---")
        st.markdown("#### 导出术语表")
        df_exp = pd.read_sql_query("""
            SELECT source_term AS '源术语',
                   target_term AS '目标术语',
                   domain AS '领域',
                   strategy AS '翻译策略',
                   example AS '示例句',
                   category AS '分类'
            FROM term_ext ORDER BY source_term COLLATE NOCASE
        """, conn)

        from io import BytesIO
        buff = BytesIO()
        with pd.ExcelWriter(buff, engine="xlsxwriter") as writer:
            df_exp.to_excel(writer, index=False, sheet_name="术语库")
        st.download_button("📥 下载 Excel",
                           buff.getvalue(),
                           file_name="terms.xlsx",
                           mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                           key=sk2("dl_terms"))

    # —— 快速搜索
    with sub_tabs[3]:
        sk3 = lambda n: f"{key_prefix}_t3_{n}"
        q = st.text_input("快速搜索(前缀/子串)", "", key=sk3("q"))
        limit = st.number_input("返回上限", 1, 5000, 1000, 100, key=sk3("limit"))
        if st.button("搜索", key=sk3("q_btn")):
            if q:
                like = f"%{q}%"
                rows = cur.execute("""
                    SELECT id, source_term, target_term, domain, project_id
                    FROM term_ext
                    WHERE source_term LIKE ? OR target_term LIKE ?
                    ORDER BY id DESC
                    LIMIT ?
                """, (like, like, int(limit))).fetchall()
                render_table(pd.DataFrame(rows, columns=["ID","源术语","目标术语","领域","项目"]),
                             key=sk3("q_grid"),editable=True)
            else:
                st.warning("请输入关键词")

    # —— 批量挂接项目
    with sub_tabs[4]:
        sk4 = lambda n: f"{key_prefix}_t4_{n}"
        st.caption("将一批术语统一设置 project_id.便于项目内优先匹配")
        ids_txt = st.text_area("术语ID列表(逗号/空格/换行分隔)", key=sk4("ids"))
        pid_to = st.text_input("目标项目ID", key=sk4("pid_to"))
        if st.button("批量挂接", key=sk4("batch_btn")):
            import re
            if not pid_to.isdigit():
                st.error("项目ID需为数字")
            else:
                raw = re.split(r"[,\s]+", ids_txt.strip())
                ids = [int(x) for x in raw if x.isdigit()]
                if not ids:
                    st.warning("未识别到有效ID")
                else:
                    qmarks = ",".join(["?"]*len(ids))
                    cur.execute(f"UPDATE term_ext SET project_id=? WHERE id IN ({qmarks})", (int(pid_to), *ids))
                    conn.commit()
                    st.success(f"✅ 已挂接 {len(ids)} 条到项目 {pid_to}")

    # —— 从历史提取术语
    with sub_tabs[5]:
        sk5 = lambda n: f"{key_prefix}_t5_{n}"
        st.markdown("#### 从翻译历史记录抽取术语(DeepSeek)")
        ak, model = get_deepseek()
        if not ak:
            st.warning("未检测到 DeepSeek Key.请先在“设置”中配置。")
        else:
            mode_pick = st.radio(
                "选择来源",
                ["按项目抽取(合并多条)", "按单条记录抽取"],
                horizontal=True,
                key=sk5("ext_mode"),
            )
            if mode_pick == "按项目抽取(合并多条)":
                pid_ext = st.text_input("项目ID", key=sk5("ext_pid"))
                max_chars = st.number_input("采样最大字符数", 1000, 20000, 5000, 500, key=sk5("ext_max"))
                if st.button("开始抽取", key=sk5("ext_go_proj")):
                    if pid_ext.isdigit():
                        rows = cur.execute(
                            "SELECT src_path, output_text FROM trans_ext WHERE project_id=? ORDER BY id DESC LIMIT 10",
                            (int(pid_ext),),
                        ).fetchall()
                        buf = []
                        total = 0
                        for sp, ot in rows:
                            src = read_source_file(sp) if sp else ""
                            txt = (src + "\n" + (ot or "")).strip()
                            if not txt:
                                continue
                            if total + len(txt) > int(max_chars):
                                remain = max(0, int(max_chars) - total)
                                buf.append(txt[:remain])
                                break
                            else:
                                buf.append(txt)
                                total += len(txt)
                        big = "\n\n".join(buf)
                        res = ds_extract_terms(big, ak, model, src_lang="zh", tgt_lang="en")
                        if not res:
                            st.info("未抽取到术语或解析失败")
                        else:
                            st.success(f"抽取到 {len(res)} 条.准备批量写入……")
                            ins = 0
                            for o in res:
                                cur.execute(
                                    "INSERT INTO term_ext (source_term, target_term, domain, project_id, strategy, example) VALUES (?, ?, ?, ?, ?, ?)",
                                    (o["source_term"], o.get("target_term") or None, o.get("domain"), int(pid_ext), o.get("strategy"), o.get("example")),
                                )
                                ins += 1
                            conn.commit()
                            st.success(f"✅ 已写入术语库 {ins} 条")
                    else:
                        st.warning("请输入数字型项目ID")
            else:
                rid_ext = st.text_input("历史记录ID", key=sk5("ext_rid"))
                if st.button("开始抽取", key=sk5("ext_go_rec")):
                    if rid_ext.isdigit():
                        row = cur.execute(
                            "SELECT src_path, output_text, project_id FROM trans_ext WHERE id=?",
                            (int(rid_ext),),
                        ).fetchone()
                        if not row:
                            st.warning("未找到该记录")
                        else:
                            sp, ot, pid0 = row
                            src = read_source_file(sp) if sp else ""
                            big = (src + "\n" + (ot or "")).strip()
                            res = ds_extract_terms(big, ak, model, src_lang="zh", tgt_lang="en")
                            if not res:
                                st.info("未抽取到术语或解析失败")
                            else:
                                st.success(f"抽取到 {len(res)} 条.准备批量写入……")
                                ins = 0
                                for o in res:
                                    cur.execute(
                                        "INSERT INTO term_ext (source_term, target_term, domain, project_id, strategy, example) VALUES (?, ?, ?, ?, ?, ?)",
                                        (o["source_term"], o.get("target_term") or None, o.get("domain"), pid0, o.get("strategy"), o.get("example")),
                                    )
                                    ins += 1
                                conn.commit()
                                st.success(f"✅ 已写入术语库 {ins} 条(project_id={pid0})")

    # —— 分类管理
    with sub_tabs[6]:
        sk6 = lambda n: f"{key_prefix}_t6_{n}"
        st.markdown("#### 分类管理")
        c1, c2 = st.columns(2)
        with c1:
            ids_txt = st.text_area("按 ID 批量设置分类(逗号/空格/换行分隔)", key=sk6("cat_ids"))
            cat_to = st.text_input("要设置的分类名", key=sk6("cat_name"))
            if st.button("批量设置分类", key=sk6("cat_set_ids")):
                import re
                raw = re.split(r"[,\s]+", (ids_txt or "").strip())
                ids = [int(x) for x in raw if x.isdigit()]
                if not ids or not cat_to.strip():
                    st.warning("请填入ID列表与分类名称")
                else:
                    qmarks = ",".join(["?"] * len(ids))
                    cur.execute(f"UPDATE term_ext SET category=? WHERE id IN ({qmarks})", (cat_to.strip(), *ids))
                    conn.commit()
                    st.success(f"✅ 已设置 {len(ids)} 条为分类:{cat_to.strip()}")

        with c2:
            pid_cat = st.text_input("将某项目ID全部术语设置为分类", key=sk6("cat_pid"))
            cat2_to = st.text_input("分类名", key=sk6("cat2_name"))
            if st.button("按项目ID统一分类", key=sk6("cat_set_pid")):
                if pid_cat.isdigit() and cat2_to.strip():
                    cur.execute("UPDATE term_ext SET category=? WHERE project_id=?", (cat2_to.strip(), int(pid_cat)))
                    conn.commit()
                    st.success(f"✅ 已将项目 {pid_cat} 的术语分类设为:{cat2_to.strip()}")
                else:
                    st.warning("请填写项目ID与分类名")

# ========== Session 初始化 ==========
if "kb_embedder" not in st.session_state and KBEmbedder:
    st.session_state["kb_embedder"] = KBEmbedder(lazy=True)

# ========== 页面结构 ==========
st.title("🧠 个人翻译知识库管理系统 · 修正版03")
tabs = st.tabs(["📂 翻译项目管理", "📘 术语库管理", "📊 翻译历史", "📚 语料库管理", "⚙ 设置"])

# ========== Tab1:翻译项目管理 ==========
with tabs[0]:
    st.subheader("翻译项目管理")
    with st.form("new_project"):
        TAG_OPTIONS = ["政治", "经济", "文化", "文物", "金融", "法律"]
        SCENE_OPTIONS = ["学术", "配音稿", "正式会议"]
        use_semantic = st.checkbox("在翻译时启用语义召回参考", value=True)
        
        # === 语义召回范围选择 ===
        scope_label = "语义召回范围"
        scope_options = {
            "仅当前项目": "project",
            "同领域 + 当前项目": "domain",
            "全库": "all"
        }
        default_scope = "仅当前项目"
        sel = st.selectbox(
            scope_label,
            list(scope_options.keys()),
            index=list(scope_options.keys()).index(default_scope),
            key="scope_sel_newproj"
        )
        st.session_state["scope_newproj"] = scope_options[sel]

        c1, c2 = st.columns([3, 2])
        with c1:
            title = st.text_input("项目名称")
            tags_sel = st.multiselect("项目标签(可多选)", TAG_OPTIONS)
            scene_sel = st.selectbox("场合", SCENE_OPTIONS, index=0)
        with c2:
            translation_type = st.selectbox("翻译方式", ["使用术语库", "纯机器翻译"])
            translation_mode = st.radio("模式", ["标准模式", "术语约束模式"], horizontal=True)
            prompt_text = st.text_area(
                "翻译提示(注入模型 System Prompt)",
                placeholder="写下对 DeepSeek 的硬性/优先级指令.如:时态统一为过去式.专有名词保持原文……",
                height=120,
                key="new_proj_prompt"
            )
        # === 领域自动绑定逻辑 ===
        # 若用户选择了多个标签.则默认以第一个标签为领域
        domain_val = tags_sel[0] if tags_sel else None

        # === 提交按钮 ===
        submitted = st.form_submit_button("💾 创建项目")
        if submitted:
            if not title:
                st.error("请填写项目名称")
            else:
                try:
                    # 确保 items 表存在 domain 字段
                    cur.execute("PRAGMA table_info(items);")
                    cols = [r[1] for r in cur.fetchall()]
                    if "domain" not in cols:
                        cur.execute("ALTER TABLE items ADD COLUMN domain TEXT;")
                        conn.commit()

                    # 插入新项目（含 domain）
                    cur.execute("""
                        INSERT INTO items(title, body, tags, domain,type)
                        VALUES (?, ?, ?, ?, 'project')
                    """, (
                        title,
                        prompt_text or "",
                        ",".join(tags_sel or []),
                        domain_val
                    ))
                    conn.commit()

                    st.success(f"✅ 项目 '{title}' 已创建（领域:{domain_val or '未指定'}）")
                except Exception as e:
                    st.error(f"❌ 创建项目失败: {e}")

    rows = cur.execute("""
        SELECT
            i.id,
            i.title,
            COALESCE(i.tags,'')         AS tags,
            COALESCE(e.src_path,'')     AS src_path,
            COALESCE(i.created_at,'')   AS created_at,
            COALESCE(i.scene,'')        AS scene,
            COALESCE(i.prompt,'')       AS prompt,
            COALESCE(i.mode,'')         AS mode,
            COALESCE(i.trans_type,'')   AS trans_type
        FROM items i
        LEFT JOIN item_ext e ON e.item_id = i.id
        WHERE COALESCE(i.type,'')='project'
        ORDER BY i.id DESC
    """).fetchall()

    if not rows:
        st.info("暂无项目")
    else:
        for pid, title, tags_str, path, ct, scene, prompt_ro, mode, trans_type in rows:
            tag_display = tags_str or "无"
            file_display = os.path.basename(path) if path else "无"

            with st.expander(f"{title}｜方式:{mode or '未设'}｜标签:{tag_display}｜场合:{scene or '未填'}｜文件:{file_display}｜创建:{ct}"):
                c1, c2, c3 = st.columns([2, 2, 1])
                with c1:
                    st.selectbox("翻译方向", ["中译英", "英译中"], key=f"lang_{pid}")
                with c2:
                    max_len = st.number_input("分块长度", 600, 2000, 1200, 100, key=f"len_{pid}")
                with c3:
                    use_terms = st.checkbox("使用术语库", value=(mode == "术语约束模式"), key=f"ut_{pid}")

                st.caption(f"标签:{tag_display}")
                st.caption(f"场合:{scene or '未填写'}")
               
                # === 领域（domain）设置:跟随第一个标签 或 手动选择 ===
                # 读取当前项目的 domain / tags
                # 保底:items 表若没有 domain 列.动态补列（兼容旧库）
                cols_items = [r[1] for r in cur.execute("PRAGMA table_info(items)").fetchall()]
                if "domain" not in cols_items:
                    try:
                        cur.execute("ALTER TABLE items ADD COLUMN domain TEXT;")
                        conn.commit()
                    except Exception:
                        pass  # 并发或已有列时忽略
 
                row = cur.execute(
                    "SELECT IFNULL(domain,''), IFNULL(tags,'') FROM items WHERE id=?",
                    (pid,)
                ).fetchone()
                domain0, tags0 = (row or ["", ""])
                tags_list = [t for t in (tags0.split(",") if tags0 else []) if t]

                DOMAIN_OPTIONS = ["政治", "经济", "文化", "文物", "金融", "法律"]

                dom_mode = st.radio(
                    "领域设置方式",
                    ["跟随第一个标签", "手动选择"],
                    horizontal=True,
                    key=f"dom_mode_{pid}"
                )

                if dom_mode == "跟随第一个标签":
                    domain_val = (tags_list[0] if tags_list else (domain0 or None))
                    st.caption(f"当前领域（自动）:{domain_val or '未指定'}（由第一个标签决定）")
                else:
                    idx = DOMAIN_OPTIONS.index(domain0) if domain0 in DOMAIN_OPTIONS else 0
                    domain_val = st.selectbox(
                        "领域（手动选择）",
                        DOMAIN_OPTIONS,
                        index=idx,
                        key=f"dom_sel_{pid}"
                    )

                sync_corpus = st.checkbox(
                    "同时回填该项目下语料的领域（仅补空或原领域相同时覆盖）",
                    value=True,
                    key=f"sync_corpus_{pid}"
                )

                if st.button("💾 保存领域设置", key=f"save_dom_{pid}", type="secondary"):
                    try:
                        # 确保 items.domain 存在
                        cols_items = [r[1] for r in cur.execute("PRAGMA table_info(items)").fetchall()]
                        if "domain" not in cols_items:
                            cur.execute("ALTER TABLE items ADD COLUMN domain TEXT;")
                            conn.commit()

                        # 更新 items.domain
                        cur.execute("UPDATE items SET domain=? WHERE id=?", (domain_val, pid))
                        conn.commit()

                        # 同步语料库的 domain（优先 corpus_main.退回 corpus）
                        def _table_exists(tb: str) -> bool:
                            return bool(cur.execute(
                                "SELECT name FROM sqlite_master WHERE type='table' AND name=?;", (tb,)
                            ).fetchone())

                        corpus_tb = "corpus_main" if _table_exists("corpus_main") else ("corpus" if _table_exists("corpus") else None)
                        if sync_corpus and corpus_tb and domain_val:
                            # 确保列存在
                            cols_corpus = [r[1] for r in cur.execute(f"PRAGMA table_info({corpus_tb})").fetchall()]
                            if "domain" not in cols_corpus:
                                cur.execute(f"ALTER TABLE {corpus_tb} ADD COLUMN domain TEXT;")
                                conn.commit()

                            # 仅补空或原 domain 与 domain0 相同时覆盖.避免误伤跨领域数据
                            cur.execute(f"""
                                UPDATE {corpus_tb}
                                SET domain = ?
                                WHERE project_id = ?
                                AND (domain IS NULL OR TRIM(domain)='' OR domain = ?)
                            """, (domain_val, pid, domain0 or ""))
                            conn.commit()

                        st.success("✅ 已保存领域设置")
                        st.rerun()

                    except Exception as e:
                        st.error(f"❌ 保存失败:{e}")

                if prompt_ro:
                    try:
                        cur.execute("SELECT IFNULL(prompt, '') FROM items WHERE id=?", (pid,))
                        prompt_ro = (cur.fetchone() or [""])[0]
                    except Exception:
                        prompt_ro = ""

                    st.text_area("翻译提示(只读)", prompt_ro or "", height=120, key=f"proj_prompt_ro_{pid}")

                colf1, colf2, colf3 = st.columns([1, 1, 2])
                if path and os.path.exists(path):
                    st.caption(f"源文件:{path}")
                    if colf1.button("删除文件", key=f"del_file_{pid}"):
                        try:
                            os.remove(path)
                        except Exception:
                            pass
                        cur.execute("UPDATE item_ext SET src_path='' WHERE item_id=?", (pid,))
                        conn.commit()
                        st.success("已删除文件")
                        st.rerun()

                    if colf2.button("删除项目", key=f"del_proj_{pid}"):
                        try:
                            if path and os.path.exists(path):
                                os.remove(path)
                        except Exception:
                            pass
                        cur.execute("DELETE FROM items WHERE id=?", (pid,))
                        cur.execute("DELETE FROM item_ext WHERE item_id=?", (pid,))
                        conn.commit()
                        st.success("项目已删除")
                        st.rerun()

                else:
                    up2 = st.file_uploader("补传文件", type=["txt","docx","xlsx"], key=f"up2_{pid}")
                    if up2:
                        new_path = os.path.join(PROJECT_DIR, f"{pid}_{up2.name}")
                        with open(new_path, "wb") as f:
                            f.write(up2.read())
                        cur.execute("SELECT id FROM item_ext WHERE item_id=?", (pid,))
                        r = cur.fetchone()
                        if r:
                            cur.execute("UPDATE item_ext SET src_path=? WHERE id=?", (new_path, r[0]))
                        else:
                            cur.execute("INSERT INTO item_ext (item_id, src_path) VALUES (?, ?)", (pid, new_path))
                        conn.commit()
                        st.success("✅ 已上传并关联")
                        st.rerun()

                # —— 执行翻译
                if st.button("执行翻译", key=f"run_{pid}", type="primary"):
                    # 1) 结果缓存初始化(统一用 session_state）
                    if "all_results" not in st.session_state:
                        st.session_state["all_results"] = []
                    st.session_state["all_results"].clear()

                    # 2) 环境检查
                    ak, model = get_deepseek()
                    if not ak:
                        st.error("未检测到 DeepSeek Key.请在 `.streamlit/secrets.toml` 配置 [deepseek]")
                        st.stop()
                    if not path or not os.path.exists(path):
                        st.error("缺少源文件")
                        st.stop()

                    # 3) 读取与分段
                    src_text = read_source_file(path)
                    st.code(repr(src_text[:400]))                     # 看字符串里有没有 '\n'
                    st.write({"len": len(src_text), "nl": src_text.count("\n"), "cr": src_text.count("\r")})
                    st.write({"preview_lines": src_text.splitlines()[:3]})
                    
                    # 用统一的 split_paragraphs 做切分
                    blocks = split_paragraphs(src_text)
                    if not blocks:
                        st.error("源文件内容为空，或未识别到有效段落")
                        st.stop()

                    st.info(f"按段落切分，共 {len(blocks)} 段，开始翻译…")

                    lang_pair_val = st.session_state.get(f"lang_{pid}", "中译英")
                    use_semantic  = bool(st.session_state.get(f"use_sem_{pid}", True))
                    scope_val     = st.session_state.get(f"scope_{pid}", st.session_state.get("scope_newproj", "project"))

                    # 4) 循环翻译
                    proj_terms_all = load_terms_for_project(cur, pid)  # 一次取全(全局+项目）.循环内做命中收缩

                    def _detect_hits(block_text: str, term_map: dict[str, str]) -> dict[str, str]:
                        bt_low = (block_text or "").lower()
                        out = {}
                        for k, v in (term_map or {}).items():
                            if not k:
                                continue
                            key_low = k.lower()
                            if key_low in bt_low or k in block_text:
                                out[k] = v
                        return out

                    for i, blk in enumerate(blocks, start=1):
                        blk = str(blk or "").strip()
                        if not blk:
                            continue

                        # 术语:句内命中 + (若有）动态覆盖
                        hits = _detect_hits(blk, proj_terms_all)
                        dyn_map = dyn_map if "dyn_map" in locals() and isinstance(dyn_map, dict) else {}
                        merged_terms = {**hits, **dyn_map}

                        # —— 参考例句（仍来自语料库 corpus）——
                        if use_semantic:
                            try:
                                # 用你自己的 split_sents：中英文都能切
                                sents_blk = split_sents(
                                    blk,
                                    lang_hint="auto",
                                    min_char=3,
                                    prefer_newline=True
                                )
                            except TypeError:
                                # 如果你现在的 split_sents 还没有这些参数，就退回最简单版本
                                try:
                                    sents_blk = split_sents(blk, lang_hint="auto")
                                except Exception:
                                    sents_blk = [blk]

                            if sents_blk:
                                # 取最后 2–3 句做“局部语义焦点”
                                if len(sents_blk) > 3:
                                    focus_text = "\n".join(sents_blk[-3:])
                                else:
                                    focus_text = "\n".join(sents_blk)
                            else:
                                focus_text = blk

                            ref_context = _build_ref_context(
                                pid,
                                focus_text,
                                topk=20,
                                min_sim=0.35,
                                prefer_side="both",
                                scope=scope_val
                            )
                        else:
                            ref_context = ""


                        # 预览(可选）
                        with st.expander(f"🔎 本段({i}) 术语注入预览", expanded=False):
                            st.code((build_term_hint(merged_terms, lang_pair_val) or "")[:1200])
                        with st.expander("🔗 项目+动态术语(合并后映射）", expanded=False):
                            st.dataframe(pd.DataFrame([(k, v) for k, v in merged_terms.items()],
                                                    columns=["source_term", "target_term"]),
                                        use_container_width=True)
                        with st.expander("📚 参考例句块(注入前）", expanded=False):
                            st.text(ref_context[:1500])

                        # —— 真正调用翻译(只调用一次）——
                        out_text = ds_translate(
                            block=blk,
                            term_dict=merged_terms,        # 关键:把术语映射传进去
                            lang_pair=lang_pair_val,
                            ak=ak,
                            model=model,
                            ref_context=ref_context
                        )

                        # 记录结果(统一用 session_state）
                        st.session_state["all_results"].append(out_text)
                        st.write(f"✅ 第 {i} 段完成")

                        # 译后一致性检查(仅检查.不触发二次翻译）
                        violated = check_term_consistency(out_text, merged_terms, blk)
                        if violated:
                            st.warning("以下术语未在译文中出现(建议人工核对）: " + ".".join(violated))

                    # ===== 循环结束后:汇总输出 =====

                    # 结果与源段落兜底与对齐
                    all_results_safe = list(st.session_state.get("all_results", []))
                    blocks_src_safe  = list(blocks if 'blocks' in locals() else [])
                    if len(blocks_src_safe) != len(all_results_safe):
                        n = min(len(blocks_src_safe), len(all_results_safe))
                        blocks_src_safe  = blocks_src_safe[:n]
                        all_results_safe = all_results_safe[:n]

                    # 文本汇总
                    final_text = "\n\n".join(all_results_safe)

                    # 一致性报告(语义+术语）
                    try:
                        term_map_report = proj_terms_all  # 直接用刚才加载的“全局+项目”术语
                        df_rep = semantic_consistency_report(
                            project_id=pid,
                            blocks_src=blocks_src_safe,
                            blocks_tgt=all_results_safe,
                            term_map=term_map_report,
                            topk=3,
                            thr=0.70
                        )
                        st.markdown("### 🔎 译后一致性报告(语义+术语)")
                        st.dataframe(df_rep, use_container_width=True)
                    except Exception as e:
                        st.caption(f"(一致性报告生成失败: {e})")

                    # 下载按钮(TXT）
                    proj_title = (title if 'title' in locals() and title else f"project_{pid}")
                    st.download_button(
                        "⬇️ 下载翻译结果 (TXT)",
                        final_text or "",
                        file_name=f"{proj_title}_翻译结果.txt",
                        mime="text/plain",
                        key=f"dl_txt_{pid}"
                    )

                    # 写入历史 trans_ext
                    try:
                        # 兼容取值（都做了兜底.防止未定义）
                        src_path = path if ('path' in locals() and path) else None
                        mode_val = mode if ('mode' in locals() and mode) else "标准模式"
                        lang_pair_val = lang_pair if ('lang_pair' in locals() and lang_pair) else "自动"

                        # 计算段数:建议以「源文本」统计；若无则退回最终译文
                        src_for_seg = src_text if ('src_text' in locals() and src_text) else final_text
                        seg_count = len(split_sents(src_for_seg, lang_hint="auto"))

                        cur.execute("""
                            INSERT INTO trans_ext (
                                project_id, src_path, lang_pair, mode, output_text,
                                stats_json, segments, term_hit_total, created_at
                            )
                            VALUES (?, ?, ?, ?, ?, ?, ?, ?, datetime('now'))
                        """, (
                            pid,             # 项目ID
                            src_path,        # 源文件路径（可为空）
                            lang_pair_val,   # 语对
                            mode_val,        # 模式
                            final_text,      # 输出文本（最终译文）
                            None,            # 统计JSON（占位）
                            seg_count,       # 段数（修复:不用 blocks_src_safe）
                            None             # 术语命中数（占位.可后续填真实值）
                        ))
                        conn.commit()
                        st.success("📝 已写入翻译历史")
                    except Exception as e:
                        st.warning(f"写入翻译历史失败: {e}")



# ========== Tab2:术语库管理 ==========
with tabs[1]:
    render_term_management(st, cur, conn, BASE_DIR, key_prefix="term")

# ========== Tab3:翻译历史(增强版) ==========
with tabs[2]:
    st.subheader("📊 翻译历史记录(可写入语料 / 抽取术语 / 下载对照 / 删除)")

    rows = cur.execute("""
        SELECT id, project_id, lang_pair,
               substr(IFNULL(output_text,''),1,120) AS prev, created_at
        FROM trans_ext
        ORDER BY datetime(created_at) DESC
        LIMIT 200
    """).fetchall()

    if not rows:
        st.info("暂无历史记录。")
    else:
        for rid, pid, lp, prev, ts in rows:
            # 项目标题(做语料标题/展示)
            ttl_row = cur.execute("SELECT IFNULL(title,'') FROM items WHERE id=?", (pid,)).fetchone()
            proj_title = (ttl_row or [""])[0] or f"project#{pid}"

            with st.expander(f"#{rid}｜项目 {pid}｜{proj_title}｜{lp}｜{ts}", expanded=False):
                # 译文全文 & 源文件路径
                det = cur.execute("SELECT output_text, src_path FROM trans_ext WHERE id=?", (rid,)).fetchone()
                tgt_full, src_path = det or ("", "")
                st.code(prev or "", language="text")
                st.text_area("译文全文", tgt_full or "", height=220, key=f"hist_full_{rid}")

                # 尝试读取原文(如果当时保存了源文件路径)
                try:
                    src_full = read_source_file(src_path) if src_path else ""
                except Exception:
                    src_full = ""

                with st.expander("原文预览(若上传了源文件)", expanded=False):
                    st.text_area("原文全文", src_full or "(未保存/未上传源文件)", height=160, key=f"hist_src_{rid}")

                c1, c2, c3, c4, c5 = st.columns(5)

                # 1) 添加进语料库
                with c1:
                    if st.button("➕ 添加进语料库", key=f"hist_add_corpus_{rid}"):
                        cur.execute("""
                            INSERT INTO corpus (title, project_id, lang_pair, src_text, tgt_text, note, created_at)
                            VALUES (?, ?, ?, ?, ?, ?, datetime('now'))
                        """, (f"{proj_title} · history#{rid}", pid, lp, src_full or None, tgt_full or "", f"from trans_ext#{rid}",))
                        conn.commit()
                        st.success("✅ 已写入语料库")

                # 2) 提取术语(走你现有的 DeepSeek 抽取函数)
                with c2:
                    if st.button("🧠 提取术语", key=f"hist_extract_terms_{rid}"):
                        ak, model = get_deepseek()
                        if not ak:
                            st.warning("未检测到 DeepSeek Key(请到“设置”页配置)")
                        else:
                            # 合并原文+译文.提高候选质量
                            big = ((src_full or "") + "\n" + (tgt_full or "")).strip()
                            res = ds_extract_terms(big, ak, model, src_lang="zh", tgt_lang="en")
                            if not res:
                                st.info("未抽取到术语或解析失败")
                            else:
                                ins = 0
                                for o in res:
                                    cur.execute("""
                                        INSERT INTO term_ext (source_term, target_term, domain, project_id, strategy, example)
                                        VALUES (?, ?, ?, ?, ?, ?)
                                    """, (
                                        o.get("source_term") or "",
                                        (o.get("target_term") or None),
                                        (o.get("domain") or None),
                                        pid,
                                        (o.get("strategy") or "history-extract"),
                                        (o.get("example") or None),
                                    ))
                                    ins += 1
                                conn.commit()
                                st.success(f"✅ 已写入术语库 {ins} 条")

                # 3) 下载双语对照(CSV / DOCX)
                with c3:
                    if st.button("⬇️ CSV 对照", key=f"hist_dl_bicsv_btn_{rid}"):
                        if not src_full:
                            st.warning("找不到原文(未上传源文件).无法生成 CSV 对照")
                        else:
                            try:
                                csv_name = f"bilingual_history_{rid}.csv"
                                csv_bytes = export_csv_bilingual((src_full, tgt_full),
                                    list(zip(
                                        [s for s in src_full.splitlines() if s.strip()],
                                        [t for t in tgt_full.splitlines() if t.strip()]
                                    )),
                                    filename=f"bilingual_history_{rid}.csv"
                                )
                            except TypeError:
                                # 如果你的导出函数是 text→bytes 版本
                                csv_name = f"bilingual_history_{rid}.csv"
                                csv_bytes = export_csv_bilingual(src_full, tgt_full)
                            st.download_button("下载 CSV", data=csv_bytes,
                                               file_name=csv_name, mime="text/csv",
                                               key=f"hist_dl_bicsv_{rid}")

                with c4:
                    if st.button("⬇️ DOCX 对照", key=f"hist_dl_bidocx_btn_{rid}"):
                        if not src_full:
                            st.warning("找不到原文(未上传源文件).无法生成 DOCX 对照")
                        else:
                            try:
                                docx_path = export_docx_bilingual(
                                    list(zip(
                                        [s for s in src_full.splitlines() if s.strip()],
                                        [t for t in tgt_full.splitlines() if t.strip()]
                                    )),
                                    filename=f"bilingual_history_{rid}.docx"
                                )
                                with open(docx_path, "rb") as f:
                                    data_docx = f.read()
                            except TypeError:
                                # 如果你的导出函数是 text→bytes 版本
                                data_docx = export_docx_bilingual(src_full, tgt_full)
                            st.download_button("下载 DOCX", data=data_docx,
                                               file_name=f"bilingual_history_{rid}.docx",
                                               mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                               key=f"hist_dl_bidocx_{rid}")

                # 4) 🗑 删除本条历史(安全确认)
                with c5:
                    with st.expander("🗑 删除本条历史(不可恢复)", expanded=False):
                        st.warning("此操作将永久删除该条 trans_ext 记录(不影响已写入语料库/术语表的数据)。")
                        ok = st.checkbox(f"我确认删除 #{rid}", key=f"hist_del_ck_{rid}")
                        if st.button("确认删除", key=f"hist_del_btn_{rid}") and ok:
                            cur.execute("DELETE FROM trans_ext WHERE id=?", (rid,))
                            conn.commit()
                            st.success("已删除.请刷新页面查看结果。")
                            st.stop()  # 终止本次渲染.避免在已删除数据上继续操作

                # 原有的“下载译文 TXT”
                st.download_button("下载译文 (TXT)", tgt_full or "",
                                   file_name=f"history_{rid}.txt",
                                   mime="text/plain",
                                   key=f"hist_dl_txt_{rid}")

# ========== Tab4:语料库管理 ==========
def render_corpus_manager(st, cur, conn, pid_prefix="corpus"):
    st.header("📚 语料库管理")
    sk = make_sk(pid_prefix)

with tabs[3]:
    render_corpus_manager(st, cur, conn)

    # 语料库语义索引用同一个缓存目录
    index_dir = INDEX_DIR  # INDEX_DIR 前面全局已经定义 = Path(BASE_DIR) / ".cache_index"
    sub = st.tabs(["新建语料", "浏览/检索", "使用与导出"])
    # -------- 新建语料 --------
    with sub[0]:
        st.subheader("📥 上传 / 对齐 / 入库")

        colA, colB = st.columns(2)
        with colA:
            one_file = st.file_uploader("① 单个文件(DOCX 表格对照 / 单语 DOCX/TXT/PDF)",
                                        type=["docx", "txt", "pdf"], key="up_one")
        with colB:
            two_src = st.file_uploader("② 原文文件(可选:与 ③ 搭配做对齐)",
                                    type=["docx", "txt", "csv", "pdf"], key="up_src")
            two_tgt = st.file_uploader("③ 译文文件(可选:与 ② 搭配做对齐)",
                                    type=["docx", "txt", "csv", "pdf"], key="up_tgt")

        st.divider()
        meta1, meta2, meta3 = st.columns([2,1,1])
        with meta1:
            title = st.text_input("语料标题", value="未命名语料")
        with meta2:
            lp = st.selectbox("方向", ["自动", "中译英", "英译中"])
        with meta3:
            pid_val = st.text_input("项目ID(可留空)")
        pid = int(pid_val) if pid_val.strip().isdigit() else None

        ins = 0
        pairs = []
        src_text = tgt_text = ""
        preview_df = None

        # ========== 路径 A:单个文件 ==========
        if one_file is not None and (two_src is None and two_tgt is None):
            ext = (one_file.name.split(".")[-1] or "").lower()
            bio = io.BytesIO(one_file.getvalue())

            if ext == "docx":
                # 先尝试“表格对照”
                tables = read_docx_tables_info(io.BytesIO(bio.getvalue()))
                if tables:
                    st.caption("检测到 DOCX 表格.优先作为双语对照导入。")
                    # 简单起见.默认第 0 张表的第 0/1 列;你也可以加入下拉选择
                    pairs = extract_pairs_from_docx_table(io.BytesIO(bio.getvalue()),
                                                        table_index=0, src_col=0, tgt_col=1,
                                                        ffill=True, drop_empty_both=True, dedup=True)
                    if not pairs:
                        # 无法抽到对照 → 当作单语文本
                        src_text = read_docx_text(io.BytesIO(bio.getvalue()))
                else:
                    # 没有表格 → 单语文本
                    src_text = read_docx_text(io.BytesIO(bio.getvalue()))

            elif ext == "txt":
                src_text = read_txt(bio)

            elif ext == "pdf":
                src_text = read_pdf_text(io.BytesIO(bio.getvalue()))

        # ========== 路径 B:两个文件(原文 + 译文)==========
        elif two_src is not None and two_tgt is not None:
            def read_any(f):
                e = (f.name.split(".")[-1] or "").lower()
                b = io.BytesIO(f.getvalue())
                if e == "docx":  return read_docx_text(b)
                if e == "txt":   return read_txt(b)
                if e == "csv":   # 假定首列文本
                    try:
                        df = pd.read_csv(b)
                        return "\n".join(df.iloc[:,0].astype(str).fillna(""))
                    except Exception:
                        return ""
                if e == "pdf":   return read_pdf_text(b)
                return ""
            src_text = read_any(two_src)
            tgt_text = read_any(two_tgt)

        # ========== 预览与决定入库方式 ==========
        # 情况 1:有 pairs(来自 DOCX 表格)
        if pairs:
            st.success(f"解析到 {len(pairs)} 对(DOCX 表格)")
            preview_df = pd.DataFrame(pairs[:200], columns=["源句","目标句"])

        # 情况 2:没有 pairs.但拿到了 src/tgt 文本 → 切句/对齐
        elif src_text and tgt_text:
            sents_src = split_sents(src_text, "zh" if lp.startswith("中") else "auto")
            sents_tgt = split_sents(tgt_text, "en" if lp.startswith("英") else "auto")
            st.caption(f"将对齐:src={len(sents_src)}  tgt={len(sents_tgt)}")
            if st.button("🔎 执行语义对齐", key="do_align"):
                pairs_aligned = align_semantic(sents_src, sents_tgt, max_jump=5)
                st.info(f"对齐得到 {len(pairs_aligned)} 对")
                pairs = [(s,t) for (s,t,score) in pairs_aligned]
                if pairs:
                    preview_df = pd.DataFrame(pairs[:200], columns=["源句","目标句"])

        # 情况 3:只有单语文本(PDF/DOCX/TXT)
        elif src_text and not tgt_text:
            sents_src = split_sents(src_text, "zh" if lp.startswith("中") else "auto")
            st.info(f"检测到单语文本.共 {len(sents_src)} 句;将仅写入 src_text。")
            preview_df = pd.DataFrame([[s, ""] for s in sents_src[:200]], columns=["源句","目标句"])
            pairs = [(s, "") for s in sents_src]

        # 预览
        if preview_df is not None:
            st.dataframe(preview_df, use_container_width=True)

        # —— 按钮 + 选项:导入语料库 | 同时建立向量
        c_imp, c_opt, c_build = st.columns([1,1,1])
        do_import = c_imp.button("📥 写入语料库", type="primary", key=sk("write_pairs_btn"))
        do_build_opt = c_opt.checkbox("同时建立向量", value=True, key=sk("build_vec_opt"))
        only_build_now = c_build.button("🧠 仅建立向量(不导入)", key=sk("only_build"))

        # —— 小工具:把 [(s,t,score)] / [(s,t)] 统一成 [(s,t)]
        def normalize_pairs_to2(pairs):
            out = []
            for p in (pairs or []):
                if isinstance(p, (list, tuple)) and len(p) >= 2:
                    out.append((p[0], p[1]))
            return out

        # 4) 写入语料库 +(可选)建索引
        if pairs and do_import:
            pairs2 = normalize_pairs_to2(pairs)
            ins = 0
            for s, t in pairs2:
                s = (s or "").strip()
                t = (t or "").strip()
                if not (s or t):
                    continue
                cur.execute("""
                    INSERT INTO corpus(title, project_id, lang_pair, src_text, tgt_text, note, created_at)
                    VALUES (?, ?, ?, ?, ?, ?, datetime('now'))
                """, (title or (one_file.name if one_file else two_src.name),
                    pid, lp, s or None, t or None, "auto-import"))
                ins += 1
            conn.commit()
            st.success(f"✅ 已写入语料库 {ins} 条。")

            # —— 可选:建立向量(只用目标句.也可改为源句.按你的检索习惯)
            if do_build_opt:
                texts = [t for _, t in pairs2 if t]   # ← 用 pairs2.避免三元组解包错误
                if texts:
                    try:
                        emb, kind = _lazy_embedder()  # 你现有的懒加载(SBERT/TF-IDF)
                        # 导入后/按钮点击时:
                        res = build_project_vector_index(int(pid) if pid is not None else 0,
                                                        use_src=True, use_tgt=True)
                        st.success(f"🧠 向量索引已更新:新增 {res['added']}.总量 {res['total']}")
                    except Exception as e:
                        st.warning(f"索引未更新:{e}")

        elif (src_text and not tgt_text) and do_import:
            # —— 单语:按句切分写入 src_text.tgt_text 置空
            sents = split_sents(src_text, "zh" if lp.startswith("中") else "en")
            ins = 0
            for s in sents:
                s = (s or "").strip()
                if not s: continue
                cur.execute("""
                    INSERT INTO corpus(title, project_id, lang_pair, src_text, tgt_text, note, created_at)
                    VALUES (?, ?, ?, ?, NULL, ?, datetime('now'))
                """, (title or (two_src.name if two_src else "mono"), pid, lp, s, "mono"))
                ins += 1
            conn.commit()
            st.success(f"✅ 已写入语料库 {ins} 条。")

            if do_build_opt and sents:
                try:
                    emb, kind = _lazy_embedder()
                    V = emb(sents)
                    save_semantic_index(str(index_dir), pid or "global", sents, V)
                    st.success(f"🧠 向量索引已更新({len(sents)} 条)。")
                except Exception as e:
                    st.warning(f"索引未更新:{e}")

        if only_build_now:
            # 任选“按项目建索引”或“全库建索引”
            st.info("正在读取语料并建立向量索引…")
            rows = cur.execute("SELECT id, IFNULL(src_text,''), IFNULL(tgt_text,'') FROM corpus").fetchall()
            texts = []
            for _, s, t in rows:
                txt = (t or s).strip()   # 按你的检索习惯:优先用译文;也可换成 s
                if txt:
                    texts.append(txt)
            if not texts:
                st.warning("没有可向量化的语料。")
            else:
                try:
                    emb, kind = _lazy_embedder()
                    V = emb(texts)
                    save_semantic_index(str(index_dir), pid or "global", texts, V)
                    st.success(f"🧠 已重建向量索引({len(texts)} 条)。")
                except Exception as e:
                    st.error(f"重建失败:{e}")

        # 5) 语义检索(对当前项目或全局)
        st.subheader("🔎 语义检索(例句召回)")
        colq1, colq2 = st.columns([3,1])
        with colq1:
            q = st.text_input("输入要检索的短语/句子", key="corpus_sem_q")
        with colq2:
            topk = st.number_input("TopK", min_value=1, max_value=50, value=5)

        if st.button("🔍 语义检索", type="primary"):
            hits = search_semantic(pid, q, topk=int(topk), scope="project")
            for sc, meta, txt in hits:
                st.write(f"得分: {sc:.3f}")
                st.write(f"来源: {meta}")
                st.write(txt)
                st.markdown("---")

    # -------- 浏览/检索 --------
    with sub[1]:
        st.subheader("🔎 浏览/检索")
        k1, k2, k3 = st.columns([2,1,1])
        with k1:
            kw = st.text_input("关键词(标题/备注/译文)", "", key=sk("kw"))
        with k2:
            lp_filter = st.selectbox("方向过滤", ["全部","中译英","英译中","自动"], key=sk("lp_filter"))
        with k3:
            limit = st.number_input("条数", min_value=10, max_value=1000, value=200, step=10, key=sk("limit"))

        sql = "SELECT id, title, IFNULL(project_id,''), IFNULL(lang_pair,''), substr(IFNULL(tgt_text,''),1,80), created_at FROM corpus WHERE 1=1"
        params = []
        if kw.strip():
            like = f"%{kw.strip()}%"
            sql += " AND (title LIKE ? OR IFNULL(note,'') LIKE ? OR IFNULL(tgt_text,'') LIKE ?)"
            params += [like, like, like]
        if lp_filter != "全部":
            sql += " AND IFNULL(lang_pair,'') = ?"; params += [lp_filter]
        sql += " ORDER BY id DESC LIMIT ?"; params += [int(limit)]
        rows = cur.execute(sql, params).fetchall()

        if not rows:
            st.info("暂无数据或未命中检索条件")
        else:
            for rid, t, pid, lpv, prev, ts in rows:
                with st.expander(f"#{rid}｜{t}｜{lpv or '—'}｜{ts}"):
                    st.caption(f"关联项目:{pid or '(无)'}")
                    st.code(prev or "", language="text")
                    c1, c2, c3, c4 = st.columns(4)
                    with c1:
                        if st.button("查看全文", key=sk(f"view_{rid}")):
                            det = cur.execute("SELECT src_text, tgt_text FROM corpus WHERE id=?", (rid,)).fetchone()
                            st.text_area("源文", det[0] or "(未保存)", height=180, key=sk(f"cor_src_{rid}"))
                            st.text_area("译文", det[1] or "", height=220, key=sk(f"cor_tgt_{rid}"))
                    with c2:
                        if st.button("标记为参考", key=sk(f"ref_{rid}")):
                            st.session_state.setdefault("corpus_refs", set())
                            st.session_state["corpus_refs"].add(rid)
                            st.success("✅ 已加入参考集合(右侧“使用与导出”查看)")
                    with c3:
                        if st.button("导出TXT", key=sk(f"cor_txt_{rid}")):
                            det = cur.execute("SELECT tgt_text FROM corpus WHERE id=?", (rid,)).fetchone()
                            st.download_button("下载译文TXT", det[0] or "", file_name=f"corpus_{rid}.txt", mime="text/plain", key=sk(f"cor_txt_dl_{rid}"))
                    with c4:
                        if st.button("删除", key=sk(f"del_{rid}")):
                            cur.execute("DELETE FROM corpus WHERE id=?", (rid,))
                            conn.commit()
                            st.warning("🗑️ 已删除.刷新后生效")
                            st.rerun()

    # -------- 使用与导出 --------
    with sub[2]:
        st.subheader("🧩 使用与导出")
        ids = list(st.session_state.get("corpus_refs", []))
        st.caption(f"已选参考数:{len(ids)}")
        if ids:
            qmarks = ",".join(["?"] * len(ids))
            dets = cur.execute(f"SELECT id, title, lang_pair, IFNULL(src_text,''), IFNULL(tgt_text,'') FROM corpus WHERE id IN ({qmarks})", ids).fetchall()
            merged_demo = "\n\n---\n\n".join([f"\n\n源文:\n{src}\n\n译文:\n{tgt}" for (i,t,lp,src,tgt) in dets])
            st.text_area("合并预览", merged_demo, height=240, key=sk("merge_preview"))

            cxa, cxb = st.columns(2)
            with cxa:
                if st.button("清空参考集合", key=sk("clear_refs")):
                    st.session_state["corpus_refs"] = set()
                    st.info("已清空")
            with cxb:
                st.download_button("⬇️ 导出合并TXT", merged_demo, file_name="corpus_refs_merged.txt", mime="text/plain", key=sk("merge_dl"))

            st.markdown("---")
            st.subheader("🔗 用作后续翻译参考(Few-shot/示例上下文)")
            st.caption("勾选此项后.系统会在调用 DeepSeek 时自动注入这些参考片段(以“参考示例”块加入 system prompt)。")

            use_as_fewshot = st.checkbox("启用参考示例注入", value=True, key="cor_use_ref")
            if st.button("保存参考注入开关", key=sk("save_switch")):
                st.session_state["cor_use_ref"] = bool(use_as_fewshot)
                st.success("✅ 已保存(对后续新翻译生效)")
        else:
            st.info("暂无选中参考。请到【浏览/检索】页勾选“标记为参考”。")
        
        st.markdown("---")
        st.subheader("🔎 向量索引(语义召回)")
        pid_opts = cur.execute("SELECT id, title FROM items ORDER BY id DESC").fetchall()
        pid_map = {f"#{i} {t}": i for (i, t) in pid_opts}
        proj_sel = st.selectbox("选择项目以构建/更新索引", ["(请选择)"] + list(pid_map.keys()), key=sk("vec_proj"))
        if st.button("⚙️ 构建/更新向量索引", key=sk("build_vec")):
            if proj_sel != "(请选择)":
                res = build_project_vector_index(pid_map[proj_sel], use_src=True, use_tgt=True)
                st.success(f"索引已更新:新增 {res['added']}.总量 {res['total']}")
            else:
                st.warning("请先选择项目")

        q_demo = st.text_area("试搜一句话(将以语义相似检索参考)", "", height=80, key=sk("q_demo"))
        topk = st.number_input("Top-K", 1, 10, 5, key=sk("q_topk"))
        if st.button("🔍 语义召回测试", key=sk("q_vec")):
            if proj_sel != "(请选择)" and q_demo.strip():
                hits = semantic_retrieve(pid_map[proj_sel], q_demo.strip(), topk=int(topk))
                if not hits:
                    st.info("索引为空或未命中。请先构建索引。")
                else:
                    for sc, m, txt in hits:
                        st.write(f"**{m['side']}** | {m['title']} | 相似度:{sc:.2f}")
                        st.code(txt, language="text")
            else:
                st.warning("请选择项目并输入查询句")

        with st.expander("🧪 索引/语料健康检查", expanded=False):
            if proj_sel == "(请选择)":
                st.info("先在上面选择一个项目。")
            else:
                _pid = pid_map[proj_sel]
                # 1) 该项目语料条数
                try:
                    cnt = cur.execute("SELECT COUNT(*) FROM corpus WHERE project_id=?", (_pid,)).fetchone()[0]
                except Exception as e:
                    st.error(f"查询 corpus 失败:{e}")
                    cnt = None

                st.write(f"项目 {_pid} 的语料条数:**{cnt}**")

                # 2) 是否存在索引文件与映射条数
                try:
                    from pathlib import Path
                    import json
                    idx_dir = INDEX_DIR  # 即 Path(BASE_DIR) / ".cache_index"
                    f_map = idx_dir / f"vecmap_{_pid}.json"
                    f_faiss = idx_dir / f"faiss_{_pid}.bin"
                    f_npy = idx_dir / f"vectors_{_pid}.npy"

                    if f_map.exists():
                        data = json.loads(f_map.read_text(encoding="utf-8") or "[]")
                        st.write(f"映射文件:{f_map.name}（条数:**{len(data)}**）")
                    else:
                        st.warning("未找到映射文件 vecmap_*.json")

                    st.write(f"FAISS 索引文件存在:{f_faiss.exists()}")
                    st.write(f"NPY 向量文件存在:{f_npy.exists()}")

                except Exception as e:
                    st.error(f"索引文件检查失败:{e}")

# ========== Tab5:设置 ==========
with tabs[4]:
    st.subheader("⚙ DeepSeek Key 配置说明")
    st.markdown("""
在 `.streamlit/secrets.toml` 中加入:
```
[deepseek]
api_key = "你的KEY"
model = "deepseek-chat"
```
    """)
    ak, model = get_deepseek()
    if ak:
        st.success(f"已检测到 DeepSeek Key(模型:{model})")
    else:
        st.warning("未检测到 DeepSeek Key")
